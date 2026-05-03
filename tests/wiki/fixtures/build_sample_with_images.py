"""一次性脚本：生成 sample_with_images.docx 测试 fixture。

生成的文档结构：
  para 0: "第一段：介绍"
  para 1: 图1 (red 32x32 PNG)
  para 2: "第二段：补充"
  para 3: 图2 (blue 32x32 PNG)
  para 4: "结尾段"
"""
import io
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image

def make_image_bytes(color: tuple) -> bytes:
    img = Image.new('RGB', (32, 32), color=color)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()

def main():
    doc = Document()
    doc.add_paragraph('第一段：介绍这份文档的目的。')
    p_img1 = doc.add_paragraph()
    p_img1.add_run().add_picture(io.BytesIO(make_image_bytes((255, 0, 0))), width=Inches(0.5))
    doc.add_paragraph('第二段：补充说明上面那张红色示意图。')
    p_img2 = doc.add_paragraph()
    p_img2.add_run().add_picture(io.BytesIO(make_image_bytes((0, 0, 255))), width=Inches(0.5))
    doc.add_paragraph('第三段：结尾。')

    out = Path(__file__).parent / 'sample_with_images.docx'
    doc.save(out)
    print(f'wrote {out}')

if __name__ == '__main__':
    main()
