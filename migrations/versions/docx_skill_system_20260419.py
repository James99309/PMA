"""add skill_type/skill_body to cli_skills and create cli_skill_versions

Revision ID: docx_skill_system_20260419
Revises: c1i5ki110001
Create Date: 2026-04-19 10:00:00.000000

"""
from alembic import op
import sqlalchemy as sa

revision = 'docx_skill_system_20260419'
down_revision = 'c1i5ki110001'
branch_labels = None
depends_on = None

DOCX_SKILL_BODY = r'''from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def init_doc(db_type="SP8D"):
    doc = Document()
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None: rPr = OxmlElement('w:rPr'); rPrDefault.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    for k,v in [('w:ascii','Arial'),('w:eastAsia','Microsoft YaHei'),('w:hAnsi','Arial'),('w:cs','Arial')]:
        rFonts.set(qn(k), v)
    for lvl, color, before, after in [('Heading 1','1F4E79','320','160'), ('Heading 2','2E75B6','240','120')]:
        h = doc.styles[lvl]
        rPr = h.element.find(qn('w:rPr'))
        if rPr is None: rPr = OxmlElement('w:rPr'); h.element.append(rPr)
        for tag in ['w:rFonts','w:b','w:color','w:sz','w:szCs']:
            for old in rPr.findall(qn(tag)): rPr.remove(old)
        rf = OxmlElement('w:rFonts')
        for k,v in [('w:ascii','Microsoft YaHei'),('w:eastAsia','Microsoft YaHei'),('w:hAnsi','Microsoft YaHei'),('w:cs','Microsoft YaHei')]:
            rf.set(qn(k), v)
        rPr.insert(0, rf)
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
        pPr = h.element.find(qn('w:pPr'))
        if pPr is None: pPr = OxmlElement('w:pPr'); h.element.insert(0, pPr)
        for old in pPr.findall(qn('w:spacing')): pPr.remove(old)
        sp = OxmlElement('w:spacing'); sp.set(qn('w:before'),before); sp.set(qn('w:after'),after)
        pPr.append(sp)
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.0)
    _add_header_footer(doc, db_type)
    return doc

COLOR = {
    'primary':'1F4E79','secondary':'2E75B6','h1':'1F4E79','h2':'2E75B6',
    'border':'DDDDDD','separator':'1F4E79','meta':'888888','dim':'555555',
    'dim_light':'AAAAAA','red':'C00000','green':'1E7B1E','row_alt':'F2F2F2',
    'row_sum':'EBF3FB','score_bad':'FFF0F0','score_good':'F0FFF0','score_na':'F5F5F5',
    'cell_bad':'FFDAD6','cell_good':'E2EFDA','cell_warn':'FFF2CC',
}

def _set_run_font(run, ascii_f, east_f=None, sz=None, bold=False, color=None):
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')): rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), ascii_f); rf.set(qn('w:hAnsi'), ascii_f); rf.set(qn('w:cs'), ascii_f)
    if east_f: rf.set(qn('w:eastAsia'), east_f)
    if east_f and east_f != ascii_f: rf.set(qn('w:hint'), 'eastAsia')
    rPr.insert(0, rf)
    if sz:
        for old in rPr.findall(qn('w:sz')): rPr.remove(old)
        for old in rPr.findall(qn('w:szCs')): rPr.remove(old)
        s = OxmlElement('w:sz'); s.set(qn('w:val'), str(sz)); rPr.append(s)
        sc = OxmlElement('w:szCs'); sc.set(qn('w:val'), str(sz)); rPr.append(sc)
    for old in rPr.findall(qn('w:b')): rPr.remove(old)
    if bold: rPr.append(OxmlElement('w:b'))
    if color:
        for old in rPr.findall(qn('w:color')): rPr.remove(old)
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)

def add_mixed(para, text, sz=11, color=None, bold=False):
    for seg in re.split(r'([\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\u2014\u2013]+)', text):
        if not seg: continue
        has_cn = any('\u4e00' <= c <= '\u9fff' for c in seg)
        run = para.add_run(seg)
        if has_cn: _set_run_font(run, 'Microsoft YaHei', 'Microsoft YaHei', sz, bold, color)
        else:       _set_run_font(run, 'Calibri', 'Arial', sz, bold, color)

def add_h1(doc, text):
    para = doc.add_heading(level=1); para.clear()
    add_mixed(para, text, sz=18)
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:pBdr')): pPr.remove(old)
    pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), COLOR['separator'])
    pBdr.append(bot); pPr.append(pBdr)

def add_h2(doc, text):
    para = doc.add_heading(level=2); para.clear()
    add_mixed(para, text, sz=16)

def add_body(doc, text, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    add_mixed(para, text, sz=11, color=color)
    return para

def add_score(doc, text, sentiment='bad'):
    cfg = {
        'bad':     (COLOR['red'],   COLOR['score_bad']),
        'good':    (COLOR['green'], COLOR['score_good']),
        'neutral': ('888888',       COLOR['score_na']),
    }
    bc, fill = cfg[sentiment]
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'16')
    left.set(qn('w:space'),'4'); left.set(qn('w:color'), bc)
    pBdr.append(left); pPr.append(pBdr)
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), fill); pPr.append(shd)
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'),'120'); pPr.append(ind)
    add_mixed(para, text, sz=11, color=bc)
    return para

def add_separator(doc):
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(3)
    sep.paragraph_format.space_after = Pt(20)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), COLOR['separator'])
    pBdr.append(bot); pPr.append(pBdr)

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def _set_cell_borders(cell, color='DDDDDD'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'), color)
        tcB.append(el)
    tcPr.append(tcB)

def _set_cell_padding(cell, top=100, bottom=100, left=120, right=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')): tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def _apply_cell(doc, cell, bg, is_label=False):
    _set_cell_bg(cell, bg)
    _set_cell_borders(cell, COLOR['border'])
    _set_cell_padding(cell)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:vAlign')): tcPr.remove(old)
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'),'center'); tcPr.append(va)
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_label else WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            has_cn = any('\u4e00' <= c <= '\u9fff' for c in run.text)
            ea = 'Microsoft YaHei' if has_cn else 'Arial'
            asc = 'Microsoft YaHei' if has_cn else 'Calibri'
            _set_run_font(run, asc, ea, 11, bold=is_label, color='FFFFFF' if is_label else None)

def _set_table_full_width(table, widths_cm):
    tbl = table._tbl; tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')): tblPr.remove(old)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'9638'); tblW.set(qn('w:type'),'dxa')
    tblPr.append(tblW)
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None: tblGrid = OxmlElement('w:tblGrid'); tbl.insert(1, tblGrid)
    for gc in tblGrid.findall(qn('w:gridCol')): tblGrid.remove(gc)
    for w in widths_cm:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tblGrid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(widths_cm[i]*567))); tcW.set(qn('w:type'),'dxa')
                tcPr.append(tcW)

def fill_table(table, data):
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            if r < len(table.rows) and c < len(table.columns):
                table.rows[r].cells[c].text = val

def style_info_table(doc, table, widths_cm):
    table.style = doc.styles['Normal Table']
    _set_table_full_width(table, widths_cm)
    for row in table.rows:
        for c_i, cell in enumerate(row.cells):
            is_label = (c_i % 2 == 0)
            _apply_cell(doc, cell, COLOR['primary'] if is_label else 'FFFFFF', is_label)

def style_data_table(doc, table, widths_cm):
    table.style = doc.styles['Normal Table']
    _set_table_full_width(table, widths_cm)
    for r_i, row in enumerate(table.rows):
        for cell in row.cells:
            if r_i == 0: _apply_cell(doc, cell, COLOR['primary'], True)
            else:
                bg = COLOR['row_alt'] if (r_i-1)%2==0 else 'FFFFFF'
                _apply_cell(doc, cell, bg, False)

def set_cell_color(cell, color_key):
    _set_cell_bg(cell, COLOR[color_key])
    _set_cell_borders(cell, COLOR['border'])

def _add_field(para, field_name):
    for el in [
        ('w:fldChar', {'w:fldCharType': 'begin'}),
        ('w:instrText', None, f' {field_name} '),
        ('w:fldChar', {'w:fldCharType': 'end'}),
    ]:
        run = OxmlElement('w:r')
        node = OxmlElement(el[0])
        if el[1]:
            for k,v in el[1].items(): node.set(qn(k),v)
        if len(el) > 2: node.text = el[2]
        run.append(node); para._p.append(run)

def _add_header_footer(doc, db_type="SP8D"):
    is_cn = (db_type == "SP8D")
    company = '和源通信（上海）股份有限公司' if is_cn else 'Evertac Solutions'
    section = doc.sections[0]
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    hp = section.header.paragraphs[0]
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(0); hp.paragraph_format.space_after = Pt(2)
    run = hp.add_run(company)
    if is_cn: _set_run_font(run,'Microsoft YaHei','Microsoft YaHei',16,color='888888')
    else:      _set_run_font(run,'Calibri','Arial',16,color='888888')
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), COLOR['border'])
    pBdr.append(bot); pPr.append(pBdr)
    fp = section.footer.paragraphs[0]
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(2); fp.paragraph_format.space_after = Pt(0)
    pPr2 = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr'); top = OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'6')
    top.set(qn('w:space'),'1'); top.set(qn('w:color'), COLOR['border'])
    pBdr2.append(top); pPr2.append(pBdr2)
    if is_cn:
        r=fp.add_run('第 '); _set_run_font(r,'Microsoft YaHei','Microsoft YaHei',16,color='AAAAAA')
        _add_field(fp,'PAGE')
        r=fp.add_run(' 页 / 共 '); _set_run_font(r,'Microsoft YaHei','Microsoft YaHei',16,color='AAAAAA')
        _add_field(fp,'NUMPAGES')
        r=fp.add_run(' 页'); _set_run_font(r,'Microsoft YaHei','Microsoft YaHei',16,color='AAAAAA')
    else:
        r=fp.add_run('Page '); _set_run_font(r,'Calibri','Arial',16,color='AAAAAA')
        _add_field(fp,'PAGE')
        r=fp.add_run(' of '); _set_run_font(r,'Calibri','Arial',16,color='AAAAAA')
        _add_field(fp,'NUMPAGES')

def add_doc_header(doc, company_tag, person_name, role, period, date, db_type="SP8D"):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(5)
    for seg in re.split(r'([\u4e00-\u9fff]+)', company_tag):
        if not seg: continue
        r = p.add_run(seg)
        has_cn = any('\u4e00'<=c<='\u9fff' for c in seg)
        if has_cn: _set_run_font(r,'Microsoft YaHei','Microsoft YaHei',15,color=COLOR['meta'])
        else:       _set_run_font(r,'Calibri','Arial',15,color=COLOR['meta'])
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(4)
    add_mixed(p2, f'{person_name} · {period}绩效评估报告', sz=22, color=COLOR['h1'])
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(3)
    add_mixed(p3, f'{role} | 评估期间：{period} | 编制日期：{date}', sz=11, color=COLOR['meta'])
    add_separator(doc)
'''


def upgrade():
    conn = op.get_bind()

    # 1. 给 cli_skills 加 skill_type 和 skill_body
    for col, col_type, default in [
        ('skill_type', 'VARCHAR(20)', "'sql'"),
        ('skill_body', 'TEXT', 'NULL'),
    ]:
        exists = conn.execute(sa.text(
            "SELECT EXISTS (SELECT 1 FROM information_schema.columns "
            "WHERE table_name='cli_skills' AND column_name=:col)"
        ), {'col': col}).scalar()
        if not exists:
            conn.execute(sa.text(
                f"ALTER TABLE cli_skills ADD COLUMN {col} {col_type} DEFAULT {default}"
            ))

    # 2. 创建 cli_skill_versions 表
    exists = conn.execute(sa.text(
        "SELECT EXISTS (SELECT 1 FROM information_schema.tables WHERE table_name='cli_skill_versions')"
    )).scalar()
    if not exists:
        op.create_table(
            'cli_skill_versions',
            sa.Column('id', sa.Integer(), autoincrement=True, nullable=False),
            sa.Column('skill_id', sa.Integer(), nullable=False),
            sa.Column('version', sa.Integer(), nullable=False),
            sa.Column('skill_body', sa.Text(), nullable=True),
            sa.Column('change_note', sa.String(500), nullable=True),
            sa.Column('created_by', sa.Integer(), nullable=True),
            sa.Column('created_at', sa.DateTime(),
                      server_default=sa.text('CURRENT_TIMESTAMP'), nullable=False),
            sa.ForeignKeyConstraint(['skill_id'], ['cli_skills.id'], ondelete='CASCADE'),
            sa.ForeignKeyConstraint(['created_by'], ['users.id'], ondelete='SET NULL'),
            sa.PrimaryKeyConstraint('id'),
        )
        op.create_index('ix_cli_skill_versions_skill_id', 'cli_skill_versions', ['skill_id'])

    # 3. 插入 docx skill（不存在时）
    existing = conn.execute(
        sa.text("SELECT id FROM cli_skills WHERE name='pma_docx_style'")
    ).scalar()
    if not existing:
        conn.execute(sa.text("""
            INSERT INTO cli_skills
                (name, title, description, parameters, queries, output_format,
                 skill_type, skill_body, scope, is_active, created_at, updated_at)
            VALUES
                ('pma_docx_style',
                 'PMA Word 文档样式库',
                 'Word 文档生成的标准样式函数库。生成绩效报告、汇报文档等 Word 文件时调用。支持 SP8D（中文）和 OVS（英文）两套公司标识。',
                 '[]'::jsonb,
                 '[]'::jsonb,
                 NULL,
                 'docx',
                 :body,
                 'global',
                 true,
                 CURRENT_TIMESTAMP,
                 CURRENT_TIMESTAMP)
        """), {'body': DOCX_SKILL_BODY})

        # 记录初始版本
        skill_id = conn.execute(
            sa.text("SELECT id FROM cli_skills WHERE name='pma_docx_style'")
        ).scalar()
        conn.execute(sa.text("""
            INSERT INTO cli_skill_versions (skill_id, version, skill_body, change_note, created_at)
            VALUES (:sid, 1, :body, '初始版本 v1.0 — 从参考文档提取样式', CURRENT_TIMESTAMP)
        """), {'sid': skill_id, 'body': DOCX_SKILL_BODY})


def downgrade():
    op.drop_index('ix_cli_skill_versions_skill_id', table_name='cli_skill_versions')
    op.drop_table('cli_skill_versions')
    op.drop_column('cli_skills', 'skill_body')
    op.drop_column('cli_skills', 'skill_type')
