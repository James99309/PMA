#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成日历系统培训手册Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_image_placeholder(doc, description):
    """添加图片占位符"""
    p = doc.add_paragraph()
    run = p.add_run(f"[图片占位符]")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(f"\n{description}")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 添加边框效果
    p_fmt = p.paragraph_format
    p_fmt.left_indent = Inches(0.2)

def create_manual():
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(11)

    # ==================== 封面 ====================
    title = doc.add_heading('日历系统培训手册', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('帮助管理行程和日志记录')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # ==================== 目录 ====================
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 系统概述',
        '2. 进入日历页面',
        '3. 日历视图说明',
        '4. 创建工作计划',
        '5. 管理行程',
        '6. 编写日志',
        '7. AI智能分析功能',
        '8. 权限与共享',
        '9. 常见问题',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_page_break()

    # ==================== 1. 系统概述 ====================
    doc.add_heading('1. 系统概述', level=1)

    doc.add_paragraph(
        '日历系统是PMA系统中用于管理工作行程和日志记录的核心模块。通过日历系统，您可以：'
    )

    features = [
        '📅 规划和管理每日/每周/每月的工作安排',
        '🔗 将工作计划关联到项目、客户和联系人',
        '📝 记录工作执行情况和行动记录',
        '📊 自动汇总每日工作成果',
        '🤖 获取AI智能分析和业绩建议',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('工作类型分类', level=2)
    doc.add_paragraph('系统支持9大类、32种工作类型，以不同颜色区分：')

    # 工作类型表格
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'

    headers = ['类别', '颜色', '包含类型']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    work_types = [
        ('通用', '灰色', '会议、出差、内部培训、其他'),
        ('行销', '绿色', '拜访客户、售前支持、商务洽谈、客户维护'),
        ('市场', '橙色', '市场策划、品牌推广、活动执行'),
        ('服务', '蓝色', '现场运维、服务响应、技术支持、故障处理'),
        ('行政', '紫色', '行政事务、办公管理、资产管理'),
        ('人事', '粉色', '人事事务、招聘面试、员工关系、绩效管理'),
        ('财务', '青色', '财务工作、报销审核、账务处理'),
        ('产品', '橘色', '产品调研、需求分析、产品规划'),
        ('供应链', '天蓝', '采购管理、库存管理、物流协调、品质跟踪'),
    ]

    for i, (cat, color, types) in enumerate(work_types, 1):
        table.rows[i].cells[0].text = cat
        table.rows[i].cells[1].text = color
        table.rows[i].cells[2].text = types

    doc.add_page_break()

    # ==================== 2. 进入日历页面 ====================
    doc.add_heading('2. 进入日历页面', level=1)

    doc.add_heading('入口位置', level=2)
    doc.add_paragraph(
        '在系统右上角，找到语言选择按钮左侧的日历图标，点击即可进入日历页面。'
    )

    add_image_placeholder(doc, '描述：显示系统顶部导航栏，用红框标注日历图标的位置（位于语言切换按钮左侧）')

    doc.add_heading('页面结构', level=2)
    doc.add_paragraph('进入日历页面后，您将看到：')

    structure_items = [
        '页面标题区：显示"工作日历"标题和副标题',
        '操作按钮区：账户选择器（管理员/部门负责人可见）、"添加工作"按钮',
        '日历主体区：显示日历内容',
    ]
    for i, item in enumerate(structure_items, 1):
        doc.add_paragraph(f'{i}. {item}')

    add_image_placeholder(doc, '描述：显示完整的日历页面，包含顶部标题区、操作按钮区和月视图日历')

    # ==================== 3. 日历视图说明 ====================
    doc.add_heading('3. 日历视图说明', level=1)
    doc.add_paragraph('日历系统提供三种视图模式，通过日历右上角的按钮组切换：')

    doc.add_heading('3.1 月视图 (Month)', level=2)
    month_features = [
        '默认视图，显示整月的工作安排',
        '每个日期格子显示当天的工作项',
        '有工作项的日期，日期数字会以蓝色加粗显示',
        '超过显示限制时，会显示"另外X个"链接',
    ]
    for f in month_features:
        doc.add_paragraph(f, style='List Bullet')

    add_image_placeholder(doc, '描述：月视图界面，显示当月日历格子，部分日期有不同颜色的工作事项，日期数字为蓝色加粗')

    doc.add_heading('3.2 周视图 (Week)', level=2)
    week_features = [
        '显示一周的详细时间安排',
        '左侧显示时间刻度（00:00 - 24:00）',
        '可以更精确地查看每天的时间分配',
        '适合查看具体时间段的会议和行程',
    ]
    for f in week_features:
        doc.add_paragraph(f, style='List Bullet')

    add_image_placeholder(doc, '描述：周视图界面，左侧显示时间轴，顶部显示周一到周日，中间区域显示各时间段的工作安排')

    doc.add_heading('3.3 日视图 (Day)', level=2)
    day_features = [
        '显示单日的详细时间安排',
        '提供最精细的时间管理视图',
        '适合规划当天具体的时间安排',
    ]
    for f in day_features:
        doc.add_paragraph(f, style='List Bullet')

    add_image_placeholder(doc, '描述：日视图界面，左侧显示从早到晚的时间刻度，右侧显示当天各时间段的工作项')

    doc.add_heading('视图切换方法', level=2)
    doc.add_paragraph('点击日历右上角的按钮组，选择对应的视图：月、周、日')

    add_image_placeholder(doc, '描述：日历右上角的视图切换按钮组，包含"月"、"周"、"日"三个按钮，当前选中的按钮高亮显示')

    doc.add_page_break()

    # ==================== 4. 创建工作计划 ====================
    doc.add_heading('4. 创建工作计划', level=1)

    doc.add_heading('4.1 打开创建窗口', level=2)
    doc.add_paragraph('有两种方式创建新的工作计划：')

    doc.add_paragraph('方式一：点击"添加工作"按钮', style='List Bullet')
    doc.add_paragraph('点击页面右上角的蓝色"添加工作"按钮，系统会打开创建工作项的模态框')

    doc.add_paragraph('方式二：点击日历日期', style='List Bullet')
    doc.add_paragraph('在日历上直接点击想要创建行程的日期，系统会自动将该日期填入开始日期')

    add_image_placeholder(doc, '描述：页面右上角的蓝色"添加工作"按钮，带有加号图标')

    doc.add_heading('4.2 填写工作信息', level=2)

    doc.add_paragraph('必填信息：').bold = True

    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'
    headers2 = ['字段', '说明', '示例']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    table2.rows[1].cells[0].text = '标题'
    table2.rows[1].cells[1].text = '工作项的名称，简明扼要'
    table2.rows[1].cells[2].text = '拜访XX客户'
    table2.rows[2].cells[0].text = '开始日期'
    table2.rows[2].cells[1].text = '计划执行的日期'
    table2.rows[2].cells[2].text = '2026-01-03'

    add_image_placeholder(doc, '描述：工作项创建弹窗，显示标题输入框、开始日期选择器、结束日期选择器')

    doc.add_paragraph()
    doc.add_paragraph('可选信息：').bold = True

    optional_fields = [
        ('结束日期', '用于跨天任务，如出差3天'),
        ('工作类型', '从9大类32种类型中选择'),
        ('全天事件', '勾选后不需要填写具体时间'),
        ('开始/结束时间', '非全天事件需要填写'),
        ('预计时长', '预估需要的工作时长（小时）'),
    ]

    table3 = doc.add_table(rows=len(optional_fields)+1, cols=2)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = '字段'
    table3.rows[0].cells[1].text = '说明'
    for cell in table3.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    for i, (field, desc) in enumerate(optional_fields, 1):
        table3.rows[i].cells[0].text = field
        table3.rows[i].cells[1].text = desc

    add_image_placeholder(doc, '描述：工作类型下拉选择器展开状态，显示9个分组（通用、行销、市场等）及其子类型')

    doc.add_heading('4.3 添加参与人员（共享行程）', level=2)
    doc.add_paragraph('如果这个行程有其他同事参与：')

    share_steps = [
        '在"共享给"区域，点击虚线圆形的加号按钮',
        '在弹出的用户选择器中选择参与人员',
        '选中的用户头像会显示在共享区域',
        '保存后，参与人员的日历中也会显示该行程（以虚线边框标识）',
    ]
    for i, step in enumerate(share_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：共享区域显示已选的用户头像，以及一个虚线圆形的添加按钮')
    add_image_placeholder(doc, '描述：用户选择器弹窗，显示可选的团队成员列表，支持搜索')

    p = doc.add_paragraph()
    run = p.add_run('注意事项：')
    run.bold = True

    notes = [
        '共享给其他人的行程，他们只能查看，不能编辑',
        '被共享的行程在日历上会以虚线边框显示，便于区分',
        '只有行程的创建者可以修改共享人员',
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    doc.add_heading('4.4 关联项目和客户', level=2)
    doc.add_paragraph('将工作计划关联到具体的项目或客户，便于后续的工作记录追踪：')

    doc.add_paragraph('关联项目：')
    project_steps = [
        '在"关联项目"输入框中输入项目名称',
        '系统会自动搜索匹配的项目',
        '从下拉列表中选择目标项目',
    ]
    for i, step in enumerate(project_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：关联项目输入框，输入关键词后显示匹配项目的下拉列表，每个项目显示项目名称和客户名称')

    doc.add_paragraph('关联客户：')
    customer_steps = [
        '在"关联客户"输入框中输入公司名称',
        '系统会自动搜索匹配的客户',
        '选择客户后，会额外显示"联系人"选择框',
        '可以进一步选择具体的联系人',
    ]
    for i, step in enumerate(customer_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：关联客户已选择，下方显示联系人下拉选择框，可选择该客户的联系人')

    doc.add_heading('4.5 填写行动记录', level=2)
    doc.add_paragraph('在"行动记录"文本框中，记录本次工作的具体内容：')

    record_items = [
        '会议的讨论要点',
        '客户拜访的沟通内容',
        '技术支持的处理情况',
        '其他需要记录的信息',
    ]
    for item in record_items:
        doc.add_paragraph(item, style='List Bullet')

    add_image_placeholder(doc, '描述：行动记录多行文本输入框，placeholder提示"记录工作内容、沟通情况..."')

    p = doc.add_paragraph()
    run = p.add_run('重要提示：')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run('这些行动记录会在工作项完成后，自动关联到相应的客户和项目，成为历史记录。')

    doc.add_heading('4.6 保存工作计划', level=2)
    doc.add_paragraph('填写完成后：')
    save_steps = [
        '点击右下角的"创建"按钮保存',
        '新创建的行程会立即显示在日历上',
        '如果有共享人员，他们的日历也会同步显示',
    ]
    for i, step in enumerate(save_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：模态框底部的按钮区，显示"取消"和"创建"两个按钮')

    doc.add_page_break()

    # ==================== 5. 管理行程 ====================
    doc.add_heading('5. 管理行程', level=1)

    doc.add_heading('5.1 查看行程详情', level=2)
    doc.add_paragraph('点击日历上的任意行程，可以查看或编辑：')
    doc.add_paragraph('自己的行程：打开编辑模态框，可以修改所有内容', style='List Bullet')
    doc.add_paragraph('他人共享的行程：打开只读查看模态框，仅可查看', style='List Bullet')

    add_image_placeholder(doc, '描述：共享行程的只读查看界面，显示行程标题、时间、关联项目客户等信息，带有"共享"标签')

    doc.add_heading('5.2 编辑行程', level=2)
    doc.add_paragraph('在行程编辑模态框中，可以修改以下内容：')
    edit_items = ['标题、描述', '日期和时间', '工作类型', '关联的项目、客户、联系人', '共享人员']
    for item in edit_items:
        doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('注意：')
    run.bold = True
    p.add_run('只有行程的创建者可以编辑，被共享的人员无法编辑。')

    doc.add_heading('5.3 拖拽调整行程', level=2)
    doc.add_paragraph('日历支持通过拖拽快速调整行程：')

    doc.add_paragraph('调整日期：在月视图中，直接拖动行程到目标日期')
    doc.add_paragraph('调整时间：在周视图/日视图中，拖动行程上下移动可调整开始时间；拖动行程底部边缘可调整结束时间')

    add_image_placeholder(doc, '描述：周视图中，一个行程正在被拖拽，显示拖拽时的视觉效果和目标位置')

    doc.add_heading('5.4 标记行程完成', level=2)
    doc.add_paragraph('当行程的结束时间已过后，您可以标记该行程为完成状态：')

    doc.add_paragraph('方法一：快速完成')
    quick_steps = [
        '点击已到期的行程打开编辑框',
        '点击左下角的"标记完成"按钮',
        '行程会直接标记为已完成状态',
    ]
    for i, step in enumerate(quick_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：编辑模态框底部左侧显示"标记完成"绿色按钮和"删除"红色按钮')

    doc.add_paragraph('方法二：详细完成（记录实际情况）')
    detail_steps = [
        '点击行程打开编辑框',
        '在弹出的完成确认窗口中填写：',
    ]
    for i, step in enumerate(detail_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    doc.add_paragraph('实际时长：填写实际花费的时间（小时）', style='List Bullet')
    doc.add_paragraph('执行备注：记录执行过程中的情况', style='List Bullet')

    add_image_placeholder(doc, '描述：完成工作项弹窗，显示工作项摘要信息、实际时长输入框、执行备注文本框，底部有"标记取消"和"标记完成"按钮')

    doc.add_paragraph('完成后的效果：')
    effects = [
        '行程在日历上会显示删除线效果',
        '行程状态变为"已完成"',
        '行程内容将被锁定，无法再修改',
        '相关记录会自动关联到客户和项目',
    ]
    for e in effects:
        doc.add_paragraph(e, style='List Bullet')

    add_image_placeholder(doc, '描述：月视图中一个已完成的行程，显示删除线效果，颜色变淡')

    doc.add_heading('5.5 取消行程', level=2)
    doc.add_paragraph('如果行程不再需要执行：')
    cancel_steps = [
        '点击行程打开编辑框',
        '在完成确认窗口中点击左下角的"标记取消"',
        '行程会标记为已取消状态',
    ]
    for i, step in enumerate(cancel_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph('取消后的效果：行程显示删除线，颜色变为半透明，无法再编辑。')

    doc.add_heading('5.6 删除行程', level=2)
    doc.add_paragraph('如果需要彻底删除一个行程：')
    delete_steps = [
        '点击行程打开编辑框',
        '点击左下角的红色"删除"按钮',
        '确认删除后，行程将从日历上移除',
    ]
    for i, step in enumerate(delete_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    p = doc.add_paragraph()
    run = p.add_run('注意：')
    run.bold = True
    p.add_run('删除操作是软删除，数据在系统中会保留用于统计分析。')

    doc.add_page_break()

    # ==================== 6. 编写日志 ====================
    doc.add_heading('6. 编写日志', level=1)

    doc.add_heading('6.1 发起日志', level=2)
    doc.add_paragraph('在完成一天的工作后，您可以编写当日工作日志：')
    log_steps = [
        '在月视图中，点击当天的日期数字',
        '系统会打开日志编辑模态框',
    ]
    for i, step in enumerate(log_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：月视图中，鼠标悬停在日期数字上，显示可点击状态')

    doc.add_heading('6.2 填写日志内容', level=2)
    doc.add_paragraph('日志编辑界面包含一个大的文本输入区域：')
    fill_steps = [
        '在文本框中描述当天的工作内容',
        '可以记录工作进度、遇到的问题、心得体会等',
        '最多可输入5000字',
    ]
    for i, step in enumerate(fill_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：日志编辑弹窗，显示日期和"的日志"标题，下方是多行文本输入框，底部显示字数计数和"预览"、"取消"按钮')

    doc.add_heading('6.3 预览日志', level=2)
    doc.add_paragraph('填写完成后，点击"预览"按钮查看日志汇总效果。预览界面会自动汇总显示：')

    doc.add_paragraph('统计卡片')
    stats = [
        '计划工作项：当天安排的工作总数',
        '已完成：已完成的工作数量',
        '总工时：累计工作时长',
        '未完成：待完成的工作数量',
    ]
    for s in stats:
        doc.add_paragraph(s, style='List Bullet')

    add_image_placeholder(doc, '描述：日志预览界面顶部的4个统计卡片，分别显示计划工作项、已完成、总工时、未完成的数值')

    doc.add_paragraph('已完成工作：')
    doc.add_paragraph('显示当天已完成的所有工作项，按工作类型颜色标注，显示实际工时和执行备注', style='List Bullet')

    add_image_placeholder(doc, '描述：已完成工作区域，每个工作项显示类型标签（带颜色）、标题、工时，有的显示执行备注')

    doc.add_paragraph('未完成工作：显示计划但未完成的工作项，便于追踪待办事项')

    doc.add_paragraph('行动记录：系统会自动统计当天在各模块的操作记录')

    table4 = doc.add_table(rows=6, cols=2)
    table4.style = 'Table Grid'
    table4.rows[0].cells[0].text = '模块'
    table4.rows[0].cells[1].text = '记录内容'
    for cell in table4.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    modules = [
        ('客户', '新创建和更新的客户'),
        ('联系人', '新创建和更新的联系人'),
        ('项目', '新创建和更新的项目'),
        ('报价单', '新创建和更新的报价单'),
        ('订单', '新创建和更新的订单'),
    ]
    for i, (mod, content) in enumerate(modules, 1):
        table4.rows[i].cells[0].text = mod
        table4.rows[i].cells[1].text = content

    add_image_placeholder(doc, '描述：行动记录区域，按模块分组显示，每个模块用图标标识，后面跟随新建和更新的记录标签')

    doc.add_heading('6.4 提交日志', level=2)
    doc.add_paragraph('预览确认无误后：')
    submit_steps = [
        '点击"提交日志"按钮',
        '日志将被正式提交并锁定',
        '提交后的日志无法再修改',
    ]
    for i, step in enumerate(submit_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：预览界面底部显示"返回编辑"和"提交日志"两个按钮')

    doc.add_paragraph('提交后的效果：')
    submit_effects = [
        '日志状态变为"已提交"',
        '日志内容被锁定',
        '管理层可以查看下属的已提交日志',
    ]
    for e in submit_effects:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_heading('6.5 修改日志', level=2)
    doc.add_paragraph('在日志提交之前，可以随时修改：')
    modify_steps = [
        '在预览界面点击"返回编辑"',
        '修改日志内容',
        '再次预览确认',
    ]
    for i, step in enumerate(modify_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    p = doc.add_paragraph()
    run = p.add_run('注意：')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run('日志一旦提交，将无法修改。')

    doc.add_page_break()

    # ==================== 7. AI智能分析功能 ====================
    doc.add_heading('7. AI智能分析功能', level=1)

    doc.add_heading('7.1 功能概述', level=2)
    doc.add_paragraph('系统集成了AI智能分析功能，可以：')
    ai_features = [
        '📊 分析您的业绩数据和KPI完成情况',
        '📈 识别业绩亮点和趋势',
        '⚠️ 提醒项目风险和费用异常',
        '💡 给出个性化的改进建议',
        '🔮 预测下月业绩走势',
    ]
    for f in ai_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('7.2 启动AI分析', level=2)

    doc.add_paragraph('自动触发：')
    doc.add_paragraph('登录系统时，如果有可用的AI分析，系统会自动弹出报告', style='List Bullet')

    doc.add_paragraph('手动触发：')
    manual_steps = [
        '在仪表盘页面右上角，找到分析图标（带有图表的图标）',
        '如果图标上有红点，表示有新的分析数据',
        '点击图标打开AI分析报告',
    ]
    for i, step in enumerate(manual_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：仪表盘页面右上角，分析图标带有红色小圆点提示，表示有新的分析数据')

    doc.add_heading('7.3 AI分析报告内容', level=2)
    doc.add_paragraph('打开AI分析报告后，您将看到以下内容：')

    doc.add_paragraph('综合评分')
    doc.add_paragraph('右上角显示综合评分（1-100分），评分颜色：')
    score_colors = [
        '绿色：80分以上（优秀）',
        '蓝色：60-79分（良好）',
        '橙色：40-59分（一般）',
        '红色：40分以下（需改进）',
    ]
    for c in score_colors:
        doc.add_paragraph(c, style='List Bullet')

    add_image_placeholder(doc, '描述：报告头部右侧的圆形评分徽章，显示数字评分，根据分数显示不同颜色')

    doc.add_paragraph('整体评价：简洁的整体业绩评价文字')
    add_image_placeholder(doc, '描述：灰色背景的整体评价文字区域')

    doc.add_paragraph('业绩亮点：列出您当前表现突出的方面')
    add_image_placeholder(doc, '描述：业绩亮点区域，每个亮点前有绿色对勾图标')

    doc.add_paragraph('本月分析：包含表现评价、环比趋势、同比趋势')
    add_image_placeholder(doc, '描述：青色背景的本月分析卡片，显示表现评价、环比趋势、同比趋势')

    doc.add_paragraph('项目分析：漏斗健康度、签约趋势、风险提醒')
    add_image_placeholder(doc, '描述：项目分析区域，显示漏斗健康度和签约趋势，如有风险会用警告图标标注')

    doc.add_paragraph('费用分析：费用效率、投入产出、预警信息')
    add_image_placeholder(doc, '描述：费用分析区域，显示效率和投入产出分析，如有预警用橙色警告图标标注')

    doc.add_paragraph('改进建议：针对您的具体情况，给出3-5条可执行的改进建议')
    add_image_placeholder(doc, '描述：蓝色背景的改进建议卡片，用编号列出具体建议')

    doc.add_paragraph('下月预测：基于历史数据和当前趋势，预测下月的业绩情况')
    add_image_placeholder(doc, '描述：蓝色背景的下月预测卡片，带有趋势箭头图标')

    doc.add_heading('7.4 管理者团队分析', level=2)
    doc.add_paragraph('如果您是M级管理者（部门负责人或更高级别），AI报告会额外提供：')

    doc.add_paragraph('三个分析标签：')
    tabs = [
        '我的分析：个人业绩分析',
        '团队整体：团队汇总分析',
        '成员分析：查看每个团队成员的分析',
    ]
    for i, tab in enumerate(tabs, 1):
        doc.add_paragraph(f'{i}. {tab}')

    add_image_placeholder(doc, '描述：AI报告顶部显示三个标签按钮：我的分析、团队整体、成员分析')

    doc.add_paragraph('团队整体分析：')
    doc.add_paragraph('点击"团队整体"标签，首次需要点击"生成团队分析"按钮触发分析')
    team_content = [
        '团队整体表现评价',
        '团队KPI完成情况汇总',
        '团队成员表现对比',
        '团队改进建议',
    ]
    for c in team_content:
        doc.add_paragraph(c, style='List Bullet')

    add_image_placeholder(doc, '描述：团队分析内容截图')

    doc.add_paragraph('成员分析：')
    doc.add_paragraph('点击"成员分析"标签，显示团队成员列表，每个成员显示评分徽章，点击成员可查看其详细分析报告')

    add_image_placeholder(doc, '描述：成员分析列表截图')

    doc.add_heading('7.5 刷新分析', level=2)
    doc.add_paragraph('AI分析会自动缓存以节省资源。如需获取最新分析：')
    refresh_steps = [
        '在报告底部找到"刷新分析"链接',
        '点击后系统会重新调用AI生成新的分析',
        '首次生成或刷新可能需要10-20秒',
    ]
    for i, step in enumerate(refresh_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_image_placeholder(doc, '描述：报告底部显示"数据来自 XXXX-XX-XX"和蓝色的"刷新分析"链接')

    doc.add_heading('7.6 关闭报告', level=2)
    doc.add_paragraph('报告底部提供两个操作按钮：')
    doc.add_paragraph('今日不再显示：关闭报告，今天不会再自动弹出', style='List Bullet')
    doc.add_paragraph('我知道了：标记为已读，正常关闭', style='List Bullet')

    add_image_placeholder(doc, '描述：报告底部显示"今日不再显示"（次要按钮）和"我知道了"（主要蓝色按钮）')

    doc.add_page_break()

    # ==================== 8. 权限与共享 ====================
    doc.add_heading('8. 权限与共享', level=1)

    doc.add_heading('8.1 日历查看权限', level=2)

    table5 = doc.add_table(rows=4, cols=2)
    table5.style = 'Table Grid'
    table5.rows[0].cells[0].text = '角色'
    table5.rows[0].cells[1].text = '可查看的日历'
    for cell in table5.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    permissions = [
        ('普通员工', '自己的日历 + 共享给自己的行程'),
        ('部门负责人', '自己的 + 同部门成员的日历'),
        ('CEO/管理员', '所有用户的日历'),
    ]
    for i, (role, perm) in enumerate(permissions, 1):
        table5.rows[i].cells[0].text = role
        table5.rows[i].cells[1].text = perm

    doc.add_heading('8.2 账户选择器', level=2)
    doc.add_paragraph('管理员、CEO和部门负责人在日历页面会看到账户选择器下拉框：')
    selector_items = [
        '默认显示"我的日历"',
        '展开后可选择查看其他成员的日历',
        '查看他人日历时，"添加工作"按钮会被禁用',
    ]
    for i, item in enumerate(selector_items, 1):
        doc.add_paragraph(f'{i}. {item}')

    add_image_placeholder(doc, '描述：页面右上角的账户选择下拉框，显示"我的日历"选项和团队成员列表')

    doc.add_heading('8.3 共享行程的标识', level=2)
    doc.add_paragraph('在日历上，共享给您的行程会以虚线边框显示，方便区分：')
    doc.add_paragraph('自己创建的行程：实线边框', style='List Bullet')
    doc.add_paragraph('他人共享的行程：虚线边框', style='List Bullet')

    add_image_placeholder(doc, '描述：日历上两个行程对比，一个实线边框（自己的），一个虚线边框（共享的）')

    doc.add_heading('8.4 日志查看权限', level=2)

    table6 = doc.add_table(rows=4, cols=2)
    table6.style = 'Table Grid'
    table6.rows[0].cells[0].text = '角色'
    table6.rows[0].cells[1].text = '可查看的日志'
    for cell in table6.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    log_permissions = [
        ('普通员工', '只能查看自己的日志'),
        ('部门负责人', '自己的 + 直属下属的日志'),
        ('CEO/管理员', '所有用户的日志'),
    ]
    for i, (role, perm) in enumerate(log_permissions, 1):
        table6.rows[i].cells[0].text = role
        table6.rows[i].cells[1].text = perm

    doc.add_page_break()

    # ==================== 9. 常见问题 ====================
    doc.add_heading('9. 常见问题', level=1)

    faqs = [
        ('Q1: 为什么我看不到"标记完成"按钮？',
         '"标记完成"按钮只有在行程的结束时间已过后才会显示。如果是全天事件，需要等到当天结束；如果设置了具体结束时间，需要等到该时间过后。'),

        ('Q2: 为什么我无法编辑某个行程？',
         '可能的原因：\n1. 该行程是他人共享给您的，您只有查看权限\n2. 该行程已经标记完成或取消，已被锁定\n3. 您正在查看他人的日历（账户选择器非"我的日历"）'),

        ('Q3: 日志提交后可以修改吗？',
         '不可以。日志一旦提交，将被永久锁定。请在提交前仔细预览确认。如确需修改，请联系系统管理员。'),

        ('Q4: AI分析报告多久更新一次？',
         'AI分析每天会自动更新一次。您也可以通过点击"刷新分析"手动触发更新。'),

        ('Q5: 为什么我看不到AI分析功能？',
         '可能的原因：\n1. 系统未配置AI服务（需要管理员配置API密钥）\n2. 您的账户未启用AI分析功能\n3. 您之前选择了"今日不再显示"'),

        ('Q6: 如何取消共享给某人的行程？',
         '1. 打开行程编辑窗口\n2. 在"共享给"区域，将鼠标悬停在要移除的用户头像上\n3. 点击出现的X按钮移除该用户\n4. 保存行程'),

        ('Q7: 日历上的不同颜色代表什么？',
         '不同颜色代表不同的工作类型分类，便于快速识别。详见本手册第1章的"工作类型分类"表格。'),

        ('Q8: 拖拽调整行程时有什么限制？',
         '1. 只能拖拽自己创建的行程\n2. 已完成/已取消的行程无法拖拽\n3. 拖拽到的目标日期不能早于今天'),
    ]

    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph(a)
        doc.add_paragraph()

    # ==================== 附录 ====================
    doc.add_heading('附录：快捷操作速查表', level=1)

    table7 = doc.add_table(rows=11, cols=2)
    table7.style = 'Table Grid'
    table7.rows[0].cells[0].text = '操作'
    table7.rows[0].cells[1].text = '方法'
    for cell in table7.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    shortcuts = [
        ('创建新行程', '点击"添加工作"按钮 或 点击日历日期'),
        ('编辑行程', '点击日历上的行程'),
        ('调整日期', '月视图拖拽行程到目标日期'),
        ('调整时间', '周/日视图上下拖拽行程'),
        ('标记完成', '编辑行程 → 点击"标记完成"'),
        ('删除行程', '编辑行程 → 点击"删除"'),
        ('写日志', '点击日期数字'),
        ('切换视图', '点击右上角 月/周/日 按钮'),
        ('查看他人日历', '使用账户选择器（需相应权限）'),
        ('刷新AI分析', '报告底部点击"刷新分析"'),
    ]
    for i, (op, method) in enumerate(shortcuts, 1):
        table7.rows[i].cells[0].text = op
        table7.rows[i].cells[1].text = method

    doc.add_paragraph()
    doc.add_paragraph()

    # 页脚信息
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('文档版本：1.0  |  最后更新：2026年1月')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 保存文档
    output_path = '/Users/nijie/Downloads/日历系统培训手册_完整版.docx'
    doc.save(output_path)
    print(f'文档已保存到: {output_path}')
    return output_path

if __name__ == '__main__':
    create_manual()
