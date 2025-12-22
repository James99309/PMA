#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
销售人员操作手册生成器 V2
包含详细步骤截图和区域标注
"""
import os
import sys
import asyncio
import shutil
from datetime import datetime
from typing import List, Dict, Optional

# 路径修正
def get_project_root():
    current = os.path.dirname(os.path.abspath(__file__))
    while current != '/':
        if os.path.exists(os.path.join(current, 'app')) and \
           os.path.exists(os.path.join(current, 'run.py')):
            return current
        current = os.path.dirname(current)
    raise RuntimeError("无法找到项目根目录")

PROJECT_ROOT = get_project_root()
sys.path.insert(0, PROJECT_ROOT)

from scripts.tools.screenshot_tool import ScreenshotTool
from scripts.tools.image_annotator import ImageAnnotator
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class SalesTutorialGeneratorV2:
    """销售操作手册生成器 V2"""

    def __init__(self, base_url: str, username: str, password: str):
        self.base_url = base_url
        self.username = username
        self.password = password
        self.screenshot_dir = os.path.join(PROJECT_ROOT, 'docs/screenshots/tutorial_v2')
        self.output_dir = os.path.join(PROJECT_ROOT, 'docs')
        self.screenshots = {}
        self.tool = None

        # 清空并重建截图目录
        if os.path.exists(self.screenshot_dir):
            shutil.rmtree(self.screenshot_dir)
        os.makedirs(self.screenshot_dir, exist_ok=True)

    async def init_browser(self):
        """初始化浏览器"""
        self.tool = ScreenshotTool(
            base_url=self.base_url,
            username=self.username,
            password=self.password,
            output_dir=self.screenshot_dir,
            headless=True,
            language="zh",
            mask_data=True
        )
        await self.tool.start()
        logged_in = await self.tool.login()
        if not logged_in:
            raise Exception("登录失败")
        print("浏览器初始化完成")

    async def close_browser(self):
        """关闭浏览器"""
        if self.tool:
            await self.tool.close()

    async def take_screenshot(self, name: str, path: str, wait_time: int = 2000) -> str:
        """截取页面截图"""
        filepath = await self.tool.screenshot(
            path=path,
            filename=name,
            wait_time=wait_time
        )
        self.screenshots[name] = filepath
        print(f"  截图: {name}")
        return filepath

    async def get_element_position(self, selector: str) -> dict:
        """动态获取元素的精确位置"""
        try:
            element = await self.tool.page.query_selector(selector)
            if element:
                box = await element.bounding_box()
                if box:
                    return {
                        'x': int(box['x']),
                        'y': int(box['y']),
                        'width': int(box['width']),
                        'height': int(box['height']),
                        'found': True
                    }
        except Exception as e:
            print(f"    获取元素位置失败 [{selector}]: {e}")
        return {'found': False}

    async def get_elements_positions(self, selectors: List[Dict]) -> List[Dict]:
        """批量获取多个元素的位置，用于精确标注"""
        results = []
        for item in selectors:
            selector = item['selector']
            # 支持多个备选选择器（逗号分隔）
            selector_list = [s.strip() for s in selector.split(',')]
            position = None

            for sel in selector_list:
                position = await self.get_element_position(sel)
                if position.get('found'):
                    break

            if position and position.get('found'):
                results.append({
                    'type': 'rect',
                    'x': position['x'],
                    'y': position['y'],
                    'width': position['width'],
                    'height': position['height'],
                    'label': item.get('label', ''),
                    'label_position': item.get('label_position', 'top'),
                    'color': item.get('color')
                })
                print(f"    ✓ 找到: {item.get('label', selector)} at ({position['x']}, {position['y']})")
            else:
                print(f"    ✗ 未找到: {item.get('label', selector)}")

        return results

    async def take_screenshot_with_action(
        self,
        name: str,
        path: str,
        actions: List[Dict] = None,
        wait_time: int = 1500
    ) -> str:
        """截取带操作的截图"""
        filepath = await self.tool.screenshot(
            path=path,
            filename=name,
            actions=actions,
            wait_time=wait_time
        )
        self.screenshots[name] = filepath
        print(f"  截图: {name}")
        return filepath

    def annotate_image(
        self,
        name: str,
        annotations: List[Dict]
    ) -> str:
        """标注图片"""
        if name not in self.screenshots:
            return None

        input_path = self.screenshots[name]
        output_path = input_path.replace('.png', '_marked.png')

        annotator = ImageAnnotator(input_path)

        for ann in annotations:
            ann_type = ann.get('type')

            if ann_type == 'rect':
                annotator.draw_rect(
                    ann['x'], ann['y'],
                    ann['width'], ann['height'],
                    label=ann.get('label'),
                    label_position=ann.get('label_position', 'top'),
                    color=ann.get('color')
                )
            elif ann_type == 'step':
                annotator.draw_step_marker(
                    ann['x'], ann['y'],
                    ann['number'],
                    description=ann.get('description')
                )
            elif ann_type == 'arrow':
                annotator.draw_arrow(
                    ann['start_x'], ann['start_y'],
                    ann['end_x'], ann['end_y']
                )
            elif ann_type == 'text':
                annotator.draw_text(
                    ann['x'], ann['y'],
                    ann['text'],
                    font_size=ann.get('font_size', 'medium')
                )
            elif ann_type == 'circle':
                annotator.draw_circle(
                    ann['x'], ann['y'],
                    radius=ann.get('radius', 20),
                    number=ann.get('number')
                )

        annotator.save(output_path)
        self.screenshots[name + '_marked'] = output_path
        print(f"  标注: {name}")
        return output_path

    # ==================== 截图采集方法 ====================

    async def capture_login_screenshots(self):
        """采集登录相关截图"""
        print("\n[登录模块截图]")

        # 需要先登出才能截取登录页
        # 这里我们用已有的截图或跳过

    async def capture_dashboard_screenshots(self):
        """采集仪表台截图"""
        print("\n[仪表台模块截图]")

        # 仪表台主页 - 使用动态定位
        await self.capture_and_annotate('dashboard_main', '/', [
            {'selector': 'nav, aside, [class*="sidebar"]', 'label': '导航菜单', 'label_position': 'right'},
            {'selector': 'input[type="search"], input[placeholder*="搜索"], .search-input', 'label': '全局搜索', 'label_position': 'bottom'},
            {'selector': '[class*="language"], button:has-text("中文"), button:has-text("English")', 'label': '语言切换', 'label_position': 'bottom'},
        ])

        # 额外截取一张展开子菜单的导航截图
        await self.tool.page.goto(f"{self.base_url}/", wait_until='networkidle')
        await asyncio.sleep(1)

        # 尝试展开"业务管理"子菜单
        menu_item = await self.tool.page.query_selector('button:has-text("业务管理"), a:has-text("业务管理"), [class*="menu"] >> text=业务管理')
        if menu_item:
            await menu_item.click()
            await asyncio.sleep(0.5)

        filepath = await self.tool.screenshot(path='/', filename='nav_menu_expanded', wait_time=500)
        self.screenshots['nav_menu_expanded'] = filepath
        # 标注导航区域
        self.annotate_image('nav_menu_expanded', [
            {'type': 'rect', 'x': 20, 'y': 90, 'width': 170, 'height': 200, 'label': '① 业务管理子菜单'},
        ])
        print("  📸 截图+标注: nav_menu_expanded")

    async def capture_and_annotate(self, name: str, path: str, highlights: List[Dict]) -> str:
        """
        采用更可靠的截图+后期标注方式
        1. 先截取干净的截图
        2. 获取元素位置
        3. 用ImageAnnotator添加标注
        """
        # 导航到页面
        await self.tool.page.goto(f"{self.base_url}{path}", wait_until='networkidle')
        await asyncio.sleep(1.5)

        # 获取元素位置信息
        annotations = []
        step_num = 1

        for h in highlights:
            selectors = [s.strip() for s in h['selector'].split(',')]
            found = False

            for sel in selectors:
                try:
                    element = await self.tool.page.query_selector(sel)
                    if element:
                        box = await element.bounding_box()
                        if box:
                            annotations.append({
                                'type': 'rect',
                                'x': int(box['x']),
                                'y': int(box['y']),
                                'width': int(box['width']),
                                'height': int(box['height']),
                                'label': f"① {h['label']}" if step_num == 1 else f"② {h['label']}" if step_num == 2 else f"③ {h['label']}" if step_num == 3 else f"④ {h['label']}",
                                'label_position': h.get('label_position', 'top')
                            })
                            print(f"    ✓ [{step_num}] {h['label']} ({int(box['x'])}, {int(box['y'])})")
                            found = True
                            break
                except Exception as e:
                    continue

            if not found:
                print(f"    ✗ [{step_num}] {h['label']} (未找到)")

            step_num += 1

        # 截取干净的截图
        filepath = await self.tool.screenshot(path=path, filename=name, wait_time=500)
        self.screenshots[name] = filepath

        # 使用ImageAnnotator添加标注
        if annotations:
            self.annotate_image(name, annotations)
            print(f"  📸 截图+标注: {name}")
        else:
            print(f"  📸 截图: {name} (无标注)")

        return filepath

    async def capture_customer_screenshots(self):
        """采集客户管理截图"""
        print("\n[客户管理模块截图]")

        # 1. 客户列表页
        await self.capture_and_annotate('customer_list', '/customer/', [
            {'selector': '.grid.grid-cols-2.md\\:grid-cols-4, .grid > div:first-child', 'label': '统计卡片', 'label_position': 'bottom'},
            {'selector': '#filterForm, form', 'label': '筛选区域', 'label_position': 'bottom'},
            {'selector': 'button[onclick*="openAdd"], .bg-primary', 'label': '添加客户按钮', 'label_position': 'left'},
            {'selector': '#companyTable, table', 'label': '客户数据列表', 'label_position': 'top'},
        ])

        # 2. 点击添加客户，截取新建表单
        await self.tool.page.goto(f"{self.base_url}/customer/", wait_until='networkidle')
        await asyncio.sleep(1)
        # 点击添加客户按钮打开模态框
        add_btn = await self.tool.page.query_selector('button[onclick*="openAdd"], button:has-text("添加客户")')
        if add_btn:
            await add_btn.click()
            await asyncio.sleep(0.8)
            # 截取新建表单
            filepath = await self.tool.screenshot(path='/customer/', filename='customer_add_form', wait_time=500)
            self.screenshots['customer_add_form'] = filepath
            # 标注表单元素
            self.annotate_image('customer_add_form', [
                {'type': 'rect', 'x': 450, 'y': 180, 'width': 500, 'height': 45, 'label': '① 公司名称（必填）'},
                {'type': 'rect', 'x': 450, 'y': 240, 'width': 500, 'height': 45, 'label': '② 企业类型'},
                {'type': 'rect', 'x': 450, 'y': 300, 'width': 500, 'height': 45, 'label': '③ 所属行业'},
                {'type': 'rect', 'x': 450, 'y': 360, 'width': 500, 'height': 45, 'label': '④ 国家/地区'},
            ])
            print("  📸 截图+标注: customer_add_form")
            # 关闭模态框
            close_btn = await self.tool.page.query_selector('[class*="modal"] button[onclick*="close"], .modal button:has-text("取消")')
            if close_btn:
                await close_btn.click()
                await asyncio.sleep(0.5)

        # 3. 客户详情页
        detail_url = await self.get_first_record_url('/customer/', 'customer')
        if detail_url:
            await self.capture_and_annotate('customer_detail', detail_url, [
                {'selector': 'h1.text-2xl, h1, .text-2xl.font-bold', 'label': '客户名称', 'label_position': 'bottom'},
                {'selector': '.card, .bg-white.rounded-xl, [class*="card"]:first-of-type', 'label': '基本信息卡片', 'label_position': 'right'},
            ])

    async def get_first_record_url(self, list_path: str, module: str) -> Optional[str]:
        """获取列表页第一条记录的详情URL"""
        try:
            await self.tool.page.goto(f"{self.base_url}{list_path}", wait_until='networkidle')
            await asyncio.sleep(1)

            # 针对Tailwind表格的选择器
            selectors = [
                f'table tbody tr:first-child a[href*="/{module}/"]',
                f'table tbody tr:first-child a[href*="view"]',
                f'[class*="table"] a[href*="/{module}/"]:first-of-type',
                f'a[href*="/{module}/"][href*="view"]',
            ]

            for selector in selectors:
                link = await self.tool.page.query_selector(selector)
                if link:
                    href = await link.get_attribute('href')
                    if href and '/add' not in href and '/edit' not in href:
                        print(f"    找到详情链接: {href}")
                        return href

        except Exception as e:
            print(f"    获取{module}详情URL失败: {e}")
        return None

    async def capture_project_screenshots(self):
        """采集项目管理截图"""
        print("\n[项目管理模块截图]")

        # 1. 项目列表页
        await self.capture_and_annotate('project_list', '/project/', [
            {'selector': '.grid.grid-cols-2, .grid > div:first-child', 'label': '项目统计', 'label_position': 'bottom'},
            {'selector': '#filterForm, form', 'label': '筛选条件', 'label_position': 'bottom'},
            {'selector': 'button[onclick*="openProject"], a[href*="/project/add"]', 'label': '创建项目', 'label_position': 'left'},
            {'selector': 'table, #projectTable', 'label': '项目列表', 'label_position': 'top'},
        ])

        # 2. 新增项目页面
        await self.capture_and_annotate('project_add', '/project/add', [
            {'selector': '#name, input[name="name"], input[placeholder*="项目"]', 'label': '项目名称', 'label_position': 'right'},
            {'selector': '#customer_id, #customer_search, [id*="customer"]', 'label': '关联客户', 'label_position': 'right'},
            {'selector': '#project_type, select[name="project_type"]', 'label': '项目类型', 'label_position': 'right'},
            {'selector': 'button[type="submit"], .btn-primary, button:has-text("保存")', 'label': '保存按钮', 'label_position': 'top'},
        ])

        # 3. 项目详情页
        detail_url = await self.get_first_record_url('/project/', 'project')
        if detail_url:
            await self.capture_and_annotate('project_detail', detail_url, [
                {'selector': 'h1, .text-2xl.font-bold, [class*="title"]', 'label': '项目名称', 'label_position': 'bottom'},
                {'selector': '.bg-white.rounded-xl, .card:first-of-type', 'label': '项目信息', 'label_position': 'right'},
            ])

    async def capture_quotation_screenshots(self):
        """采集报价管理截图"""
        print("\n[报价管理模块截图]")

        # 1. 报价列表页
        await self.capture_and_annotate('quotation_list', '/quotation/quotations', [
            {'selector': '.grid.grid-cols-2, .grid > div:first-child', 'label': '报价统计', 'label_position': 'bottom'},
            {'selector': '#filterForm, form', 'label': '筛选区域', 'label_position': 'bottom'},
            {'selector': 'a[href*="quotation"][href*="add"], button[onclick*="Quotation"]', 'label': '新建报价', 'label_position': 'left'},
            {'selector': 'table', 'label': '报价列表', 'label_position': 'top'},
        ])

        # 2. 报价详情页
        detail_url = await self.get_first_record_url('/quotation/quotations', 'quotation')
        if detail_url:
            await self.capture_and_annotate('quotation_detail', detail_url, [
                {'selector': 'h1, .text-2xl.font-bold', 'label': '报价单标题', 'label_position': 'bottom'},
                {'selector': 'table, .product-table', 'label': '产品明细表', 'label_position': 'top'},
            ])

    async def capture_expense_screenshots(self):
        """采集报销管理截图"""
        print("\n[报销管理模块截图]")

        # 1. 报销列表页
        await self.capture_and_annotate('expense_list', '/expense/', [
            {'selector': '.grid.grid-cols-2, .grid > div:first-child', 'label': '报销统计', 'label_position': 'bottom'},
            {'selector': '#filterForm, form', 'label': '筛选区域', 'label_position': 'bottom'},
            {'selector': 'a[href*="expense"][href*="add"], button[onclick*="Expense"]', 'label': '新建报销', 'label_position': 'left'},
            {'selector': 'table', 'label': '报销列表', 'label_position': 'top'},
        ])

        # 2. 报销详情页
        detail_url = await self.get_first_record_url('/expense/', 'expense')
        if detail_url:
            await self.capture_and_annotate('expense_detail', detail_url, [
                {'selector': 'h1, .text-2xl.font-bold', 'label': '报销单标题', 'label_position': 'bottom'},
                {'selector': '.bg-white.rounded-xl, .card:first-of-type', 'label': '报销信息', 'label_position': 'right'},
            ])

    async def capture_all_screenshots(self):
        """采集所有截图"""
        print("=" * 60)
        print("开始采集教程截图...")
        print("=" * 60)

        try:
            await self.init_browser()

            await self.capture_dashboard_screenshots()
            await self.capture_customer_screenshots()
            await self.capture_project_screenshots()
            await self.capture_quotation_screenshots()
            await self.capture_expense_screenshots()

            print("\n" + "=" * 60)
            print(f"截图采集完成！共 {len(self.screenshots)} 张")
            print("=" * 60)

        finally:
            await self.close_browser()

    # ==================== Word 文档生成 ====================

    def generate_word_document(self):
        """生成 Word 文档"""
        print("\n开始生成 Word 文档...")

        doc = Document()
        self._setup_document(doc)

        # 封面
        self._add_cover(doc)

        # 目录
        self._add_toc(doc)

        # 第一章：快速入门
        self._add_chapter_1_quick_start(doc)

        # 第二章：仪表台
        self._add_chapter_2_dashboard(doc)

        # 第三章：客户管理
        self._add_chapter_3_customer(doc)

        # 第四章：项目管理
        self._add_chapter_4_project(doc)

        # 第五章：报价管理
        self._add_chapter_5_quotation(doc)

        # 第六章：报销管理
        self._add_chapter_6_expense(doc)

        # 第七章：通用功能 - 筛选与搜索
        self._add_chapter_7_filter_search(doc)

        # 第八章：通用功能 - 审批流程
        self._add_chapter_8_approval(doc)

        # 附录
        self._add_appendix(doc)

        # 保存
        output_path = os.path.join(self.output_dir, 'PMA销售人员操作手册_详细版.docx')
        doc.save(output_path)
        print(f"\n文档已保存: {output_path}")
        return output_path

    def _setup_document(self, doc):
        """设置文档基本样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 设置标题样式
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            style.font.name = '微软雅黑'
            style.font.bold = True
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _add_cover(self, doc):
        """添加封面"""
        for _ in range(5):
            doc.add_paragraph()

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('PMA 系统')
        run.bold = True
        run.font.size = Pt(42)

        doc.add_paragraph()

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('销售人员操作手册')
        run.bold = True
        run.font.size = Pt(28)

        subtitle2 = doc.add_paragraph()
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle2.add_run('（详细版）')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(128, 128, 128)

        for _ in range(6):
            doc.add_paragraph()

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f'版本: V2.0\n')
        info.add_run(f'生成日期: {datetime.now().strftime("%Y年%m月%d日")}\n')
        info.add_run('文档类型: 标准操作手册')

        doc.add_page_break()

    def _add_toc(self, doc):
        """添加目录"""
        doc.add_heading('目录', level=1)

        toc = [
            ('第一章 快速入门', [
                '1.1 系统登录',
                '1.2 界面总览',
                '1.3 基本操作指南'
            ]),
            ('第二章 仪表台', [
                '2.1 界面说明',
                '2.2 数据统计',
                '2.3 待办事项'
            ]),
            ('第三章 客户管理', [
                '3.1 客户列表',
                '3.2 新增客户',
                '3.3 编辑客户',
                '3.4 客户详情',
                '3.5 联系人管理',
                '3.6 拜访记录'
            ]),
            ('第四章 项目管理', [
                '4.1 项目列表',
                '4.2 创建项目',
                '4.3 项目阶段管理',
                '4.4 项目跟进'
            ]),
            ('第五章 报价管理', [
                '5.1 报价列表',
                '5.2 创建报价',
                '5.3 发起批价',
                '5.4 报价审批流程'
            ]),
            ('第六章 报销管理', [
                '6.1 报销列表',
                '6.2 新建报销',
                '6.3 上传票据',
                '6.4 报销详情'
            ]),
            ('第七章 筛选与搜索（通用功能）', [
                '7.1 搜索功能使用',
                '7.2 筛选条件说明',
                '7.3 组合筛选技巧',
                '7.4 重置与清除'
            ]),
            ('第八章 审批流程（通用功能）', [
                '8.1 审批状态说明',
                '8.2 提交审批',
                '8.3 撤回审批',
                '8.4 审批操作（审批人）',
                '8.5 审批流程图'
            ]),
            ('附录', [
                '常见问题',
                '快捷键说明',
                '技术支持'
            ])
        ]

        for chapter, sections in toc:
            p = doc.add_paragraph()
            run = p.add_run(chapter)
            run.bold = True
            run.font.size = Pt(12)

            for section in sections:
                p = doc.add_paragraph()
                p.add_run('    ' + section)

        doc.add_page_break()

    def _add_image(self, doc, name: str, caption: str = None, width: float = 6.0):
        """添加图片"""
        # 优先使用标注版
        marked_name = name + '_marked'
        if marked_name in self.screenshots:
            path = self.screenshots[marked_name]
        elif name in self.screenshots:
            path = self.screenshots[name]
        else:
            doc.add_paragraph(f'[图片缺失: {name}]')
            return

        try:
            doc.add_picture(path, width=Inches(width))
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        except Exception as e:
            doc.add_paragraph(f'[图片加载失败: {name}]')

    def _add_step_table(self, doc, steps: List[Dict]):
        """添加步骤表格"""
        table = doc.add_table(rows=len(steps) + 1, cols=3)
        table.style = 'Table Grid'

        # 表头
        headers = table.rows[0].cells
        headers[0].text = '步骤'
        headers[1].text = '操作'
        headers[2].text = '说明'

        for cell in headers:
            cell.paragraphs[0].runs[0].bold = True

        # 数据行
        for i, step in enumerate(steps, 1):
            row = table.rows[i].cells
            row[0].text = str(step.get('step', i))
            row[1].text = step.get('action', '')
            row[2].text = step.get('description', '')

        doc.add_paragraph()

    def _add_tip_box(self, doc, title: str, content: str, tip_type: str = 'info'):
        """添加提示框"""
        icons = {
            'info': 'ℹ️',
            'warning': '⚠️',
            'tip': '💡',
            'important': '❗'
        }

        icon = icons.get(tip_type, 'ℹ️')

        p = doc.add_paragraph()
        run = p.add_run(f'{icon} {title}')
        run.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph(content)
        p.paragraph_format.left_indent = Inches(0.3)

        doc.add_paragraph()

    # ==================== 各章节内容 ====================

    def _add_chapter_1_quick_start(self, doc):
        """第一章：快速入门"""
        doc.add_heading('第一章 快速入门', level=1)

        doc.add_paragraph(
            '欢迎使用 PMA 系统！这份手册会帮你快速上手。'
            '不用担心，操作起来很简单，跟着步骤走就行。'
        )

        # 1.1 系统登录
        doc.add_heading('1.1 登录系统', level=2)

        doc.add_paragraph(
            '首先，打开浏览器访问 PMA 系统。推荐使用 Chrome 或 Edge 浏览器，体验最佳。'
        )

        doc.add_paragraph(
            '登录步骤：\n'
            '1. 在地址栏输入系统网址\n'
            '2. 输入管理员分配给你的用户名\n'
            '3. 输入密码\n'
            '4. 点击「登录」按钮'
        )

        self._add_tip_box(doc, '账号安全小提示',
            '• 首次登录后建议立即修改密码\n'
            '• 密码输错5次会被锁定30分钟，别急慢慢输\n'
            '• 忘记密码？找管理员帮你重置',
            tip_type='warning'
        )

        # 1.2 界面总览
        doc.add_heading('1.2 认识主界面', level=2)

        doc.add_paragraph(
            '登录成功后，你会看到系统主界面。别被满屏的信息吓到，'
            '其实就几个主要区域，熟悉了就很好用：'
        )

        self._add_image(doc, 'dashboard_main', '图1-1 系统主界面')

        doc.add_paragraph(
            '① 左侧导航菜单\n'
            '这里是所有功能的入口。点击主菜单会展开子菜单，显示具体的功能模块。\n\n'
            '② 顶部搜索栏\n'
            '想快速找某个客户或项目？直接在这里搜索，输入关键词回车就行。\n\n'
            '③ 语言切换\n'
            '系统支持中英文切换，点击就能切换。'
        )

        # 导航菜单展开截图
        self._add_image(doc, 'nav_menu_expanded', '图1-2 展开的导航菜单')

        doc.add_paragraph(
            '点击「业务管理」等主菜单，会展开子菜单，显示：\n'
            '• 客户管理 — 管理客户信息和联系人\n'
            '• 项目管理 — 跟踪销售项目进展\n'
            '• 报价管理 — 创建和管理报价单\n'
            '• 报销管理 — 提交费用报销申请'
        )

        # 1.3 基本操作
        doc.add_heading('1.3 几个通用操作', level=2)

        doc.add_paragraph(
            '系统的大部分页面操作都很相似，掌握这几个基本操作就能自如使用了：'
        )

        doc.add_paragraph(
            '🔍 搜索和筛选\n'
            '每个列表页顶部都有搜索框和筛选条件，帮你快速找到需要的数据。\n\n'
            '📊 排序\n'
            '点击表格的列标题，可以按那一列排序。再点一次切换升序/降序。\n\n'
            '👆 查看详情\n'
            '列表里的名称通常是蓝色链接，点击就能打开详情页。\n\n'
            '➕ 新增数据\n'
            '页面右上角通常有「添加」或「新建」按钮，点击开始创建新记录。\n\n'
            '✏️ 编辑和删除\n'
            '进入详情页后，可以找到编辑按钮修改信息。'
        )

        doc.add_page_break()

    def _add_chapter_2_dashboard(self, doc):
        """第二章：仪表台"""
        doc.add_heading('第二章 仪表台', level=1)

        doc.add_paragraph(
            '登录系统后看到的第一个页面就是仪表台。'
            '它就像你的工作台面，把最重要的信息都摆在眼前。'
        )

        # 2.1 界面说明
        doc.add_heading('2.1 你的工作仪表盘', level=2)

        doc.add_paragraph(
            '每天打开系统，先扫一眼仪表台，就能了解今天要处理什么事：'
        )

        self._add_image(doc, 'dashboard_main', '图2-1 仪表台界面')

        # 2.2 数据统计
        doc.add_heading('2.2 关键数据一览', level=2)

        doc.add_paragraph(
            '仪表台顶部的统计卡片告诉你几个关键数字：'
        )

        doc.add_paragraph(
            '📊 客户数据\n'
            '你负责多少客户？最近新增了几个？活跃客户有多少？\n\n'
            '📁 项目进度\n'
            '手上有多少个进行中的项目？分别处于什么阶段？\n\n'
            '💰 报价情况\n'
            '有多少报价在等审批？本月报价总额是多少？\n\n'
            '📈 业绩概览\n'
            '个人业绩完成得怎么样？距离目标还有多远？'
        )

        self._add_tip_box(doc, '养成好习惯',
            '每天上班先看一眼仪表台，了解自己的工作状态。'
            '特别是待审批的事项和即将到期的跟进，别错过了。',
            tip_type='tip'
        )

        # 2.3 待办事项
        doc.add_heading('2.3 需要你关注的事项', level=2)

        doc.add_paragraph(
            '仪表台还会提醒你一些需要处理的事情：\n\n'
            '• 待处理的审批 — 有人提交的审批在等你处理\n'
            '• 需要跟进的项目 — 超过一定时间没更新的项目\n'
            '• 即将过期的报价 — 报价有效期快到了，该跟进客户了\n'
            '• 最近的拜访记录 — 团队最近的客户沟通动态'
        )

        doc.add_page_break()

    def _add_chapter_3_customer(self, doc):
        """第三章：客户管理"""
        doc.add_heading('第三章 客户管理', level=1)

        doc.add_paragraph(
            '作为销售人员，客户管理是你日常工作中最常用的功能。'
            '在这里，你可以记录所有客户信息、管理联系人、追踪拜访记录，'
            '让你对每个客户的情况了然于胸。'
        )

        # 3.1 客户列表
        doc.add_heading('3.1 查看你的客户', level=2)

        doc.add_paragraph(
            '进入系统后，点击左侧菜单的「客户管理」，就能看到你负责的所有客户。'
            '页面设计得很直观：'
        )

        self._add_image(doc, 'customer_list', '图3-1 客户列表页面')

        doc.add_paragraph('让我们从上到下认识这个页面：')

        doc.add_paragraph(
            '• 顶部的统计卡片 — 一眼就能看到你有多少客户、活跃客户占比等关键数据\n'
            '• 中间的筛选区域 — 当客户变多时，帮你快速找到想要的那个\n'
            '• 下方的客户列表 — 显示所有客户的基本信息，点击名称就能查看详情'
        )

        # 3.2 筛选功能
        doc.add_heading('3.2 快速找到目标客户', level=2)

        doc.add_paragraph(
            '随着客户越来越多，逐个翻找会很费时。这时候筛选功能就派上用场了。'
        )

        doc.add_paragraph(
            '比如说，你想找「北京地区的代理商客户」，只需要：\n'
            '1. 在「地区」下拉框选择「北京」\n'
            '2. 在「企业类型」下拉框选择「代理」\n'
            '列表会立刻刷新，只显示符合条件的客户。'
        )

        self._add_tip_box(doc, '小技巧',
            '筛选条件可以叠加使用。如果想清除所有筛选，点击「重置」按钮即可恢复显示全部客户。',
            tip_type='tip'
        )

        # 3.3 新增客户
        doc.add_heading('3.3 添加新客户', level=2)

        doc.add_paragraph(
            '遇到新的潜在客户？把他们加到系统里吧。'
            '点击页面右上角的「添加客户」按钮，会弹出新建表单：'
        )

        self._add_image(doc, 'customer_add_form', '图3-2 新建客户表单')

        doc.add_paragraph(
            '必填信息：\n'
            '• 公司名称 — 建议填写客户的完整注册名称，方便后续核对\n'
            '• 企业类型 — 是终端用户、经销商还是代理商？选择合适的类型\n\n'
            '建议填写：\n'
            '• 所属行业 — 便于后续按行业分析客户分布\n'
            '• 联系地址 — 先选择国家和地区，再填详细地址\n'
            '• 主要联系人 — 记录关键对接人的电话和邮箱'
        )

        self._add_tip_box(doc, '温馨提示',
            '信息填得越完整，后续跟进时越方便。不过如果暂时只知道公司名称，也可以先创建，之后再补充。',
            tip_type='info'
        )

        # 3.4 客户详情
        doc.add_heading('3.4 了解客户的全貌', level=2)

        doc.add_paragraph(
            '在客户列表中点击任意一个客户名称，就能进入该客户的详情页。'
            '这里汇集了关于这个客户的所有信息：'
        )

        self._add_image(doc, 'customer_detail', '图3-2 客户详情页面')

        doc.add_paragraph(
            '• 基本信息 — 公司名称、类型、行业、地址等基础资料\n'
            '• 联系人列表 — 这家公司你认识的所有人，以及他们的联系方式\n'
            '• 关联项目 — 和这个客户相关的销售项目进展\n'
            '• 拜访记录 — 每次沟通的内容摘要，方便回顾历史'
        )

        doc.add_paragraph(
            '当你需要拜访某个客户前，先来这里看看之前的沟通记录，做到心中有数。'
        )

        # 3.5 联系人管理
        doc.add_heading('3.5 管理客户联系人', level=2)

        doc.add_paragraph(
            '一家公司可能有多个对接人：采购负责人、技术负责人、老板...'
            '把他们都记录下来，沟通时才能找对人。'
        )

        doc.add_paragraph(
            '添加联系人很简单：\n'
            '1. 进入客户详情页\n'
            '2. 在「联系人」区域点击「添加」\n'
            '3. 填写姓名、职位、电话、邮箱\n'
            '4. 保存即可'
        )

        self._add_tip_box(doc, '实用建议',
            '给联系人添加备注，记录一些有用的信息。比如："张总喜欢上午10点后沟通"、"李工对价格比较敏感"等，这些细节会让你的跟进更有针对性。',
            tip_type='tip'
        )

        # 3.6 拜访记录
        doc.add_heading('3.6 记录每次沟通', level=2)

        doc.add_paragraph(
            '好记性不如烂笔头。每次和客户沟通后，花一两分钟记录下来，'
            '下次跟进时就不用绞尽脑汁回忆上次聊了什么。'
        )

        doc.add_paragraph(
            '添加拜访记录的方法：\n'
            '1. 进入客户详情页\n'
            '2. 找到「拜访记录」区域，点击「添加」\n'
            '3. 选择沟通方式：电话、拜访、邮件、线上会议\n'
            '4. 写下这次沟通的主要内容：谈了什么、客户反馈如何\n'
            '5. 设置下次跟进计划：什么时候、做什么事\n'
            '6. 保存'
        )

        self._add_tip_box(doc, '记录技巧',
            '记录不用太长，抓住要点即可。比如：\n'
            '• "客户对A产品感兴趣，要求下周提供报价"\n'
            '• "对方说预算有限，考虑B方案"\n'
            '• "约好下周三下午现场拜访"\n'
            '这样的记录简洁明了，日后查看一目了然。',
            tip_type='tip'
        )

        doc.add_page_break()

    def _add_chapter_4_project(self, doc):
        """第四章：项目管理"""
        doc.add_heading('第四章 项目管理', level=1)

        doc.add_paragraph(
            '当一个客户有了明确的购买意向，就可以创建一个「项目」来跟踪这笔生意的进展。'
            '项目管理帮助你清晰地知道每个商机走到了哪一步，还差什么才能成交。'
        )

        # 4.1 项目列表
        doc.add_heading('4.1 查看你的项目', level=2)

        doc.add_paragraph(
            '点击左侧菜单的「项目管理」，你会看到自己负责的所有项目。'
            '列表按更新时间排序，最近跟进的项目排在最前面。'
        )

        self._add_image(doc, 'project_list', '图4-1 项目列表页面')

        doc.add_paragraph(
            '列表中的关键信息一目了然：\n'
            '• 项目名称 — 点击可以查看完整详情\n'
            '• 项目阶段 — 告诉你这个项目目前进展到哪一步了\n'
            '• 报价金额 — 这个项目的预期成交金额\n'
            '• 活动状态 — 项目是否还在推进中'
        )

        # 4.2 创建项目
        doc.add_heading('4.2 发现商机？立即创建项目', level=2)

        doc.add_paragraph(
            '当你跟客户聊到具体的采购计划时，就是创建项目的好时机。'
            '点击右上角的「创建项目」，开始记录这个商机：'
        )

        doc.add_paragraph(
            '必填信息：\n'
            '• 项目名称 — 取一个能概括项目的名字，比如"XX公司2024年设备采购"\n'
            '• 关联客户 — 选择这个项目属于哪个客户\n'
            '• 项目类型 — 是直销还是渠道？这会影响后续的审批流程\n\n'
            '建议填写：\n'
            '• 项目阶段 — 目前走到哪一步了\n'
            '• 预计金额 — 大概能做多少钱\n'
            '• 预计成交时间 — 客户计划什么时候决定'
        )

        self._add_image(doc, 'project_add', '图4-2 创建项目页面')

        # 4.3 项目阶段
        doc.add_heading('4.3 理解项目阶段', level=2)

        doc.add_paragraph(
            '项目从发现到成交，会经历几个阶段。及时更新阶段，'
            '既方便自己掌握进度，也让领导了解你的业务情况。'
        )

        doc.add_paragraph(
            '常见的项目阶段：\n\n'
            '① 商机发现\n'
            '刚刚接触，客户表达了初步兴趣，但还不确定具体需求。\n\n'
            '② 需求确认\n'
            '深入沟通后，明确了客户要买什么、预算多少、时间节点。\n\n'
            '③ 方案/报价\n'
            '已经提供了技术方案或报价，等待客户反馈。\n\n'
            '④ 商务洽谈\n'
            '进入价格谈判、合同条款协商阶段。\n\n'
            '⑤ 签约成交\n'
            '合同签订，项目成功！'
        )

        self._add_tip_box(doc, '阶段更新建议',
            '每次和客户有实质性进展后，记得更新项目阶段。比如：\n'
            '• 客户确认了购买清单 → 从"商机发现"进入"需求确认"\n'
            '• 发送了正式报价 → 进入"方案/报价"阶段\n'
            '• 开始讨论付款条款 → 进入"商务洽谈"阶段',
            tip_type='tip'
        )

        # 4.4 项目跟进
        doc.add_heading('4.4 持续跟进你的项目', level=2)

        doc.add_paragraph(
            '创建项目只是开始，持续跟进才能推动成交。'
            '养成定期回顾项目的习惯：'
        )

        doc.add_paragraph(
            '• 每周看一遍自己的项目列表，哪些需要这周跟进？\n'
            '• 超过两周没动静的项目，主动联系客户问问进展\n'
            '• 重要项目在系统里记录每次沟通，方便回顾和汇报'
        )

        self._add_tip_box(doc, '避免项目"躺平"',
            '有些项目可能因为各种原因暂停了。定期清理一下：\n'
            '• 确实暂停的 → 更新状态为"暂停"，记录原因\n'
            '• 已经没戏的 → 更新为"丢失"，总结教训\n'
            '• 保持列表整洁，才能把精力放在有希望的项目上',
            tip_type='info'
        )

        doc.add_page_break()

    def _add_chapter_5_quotation(self, doc):
        """第五章：报价管理"""
        doc.add_heading('第五章 报价管理', level=1)

        doc.add_paragraph(
            '谈到价格，客户总会问"你们报价多少？"'
            '报价管理就是帮你创建规范的报价单、申请特价审批、跟踪报价状态的地方。'
        )

        # 5.1 报价列表
        doc.add_heading('5.1 查看报价记录', level=2)

        doc.add_paragraph(
            '点击左侧菜单的「报价管理」，可以看到你做过的所有报价。'
            '系统会显示每个报价的状态：草稿、待审批、已通过、已发送等。'
        )

        self._add_image(doc, 'quotation_list', '图5-1 报价列表页面')

        doc.add_paragraph(
            '通过筛选功能，你可以快速找到：\n'
            '• 还在草稿状态、需要继续完善的报价\n'
            '• 正在等待审批的报价\n'
            '• 某个客户的历史报价记录'
        )

        # 5.2 创建报价
        doc.add_heading('5.2 给客户做一份报价', level=2)

        doc.add_paragraph(
            '客户问价了？来创建一份报价单吧。点击「新建报价」，按提示填写：'
        )

        doc.add_paragraph(
            '基本信息：\n'
            '• 选择客户 — 这份报价给哪个客户\n'
            '• 关联项目 — 如果有对应的项目，关联起来方便管理\n'
            '• 报价有效期 — 这个价格有效到什么时候\n\n'
            '产品明细：\n'
            '• 添加产品 — 选择客户询价的产品\n'
            '• 填写数量 — 客户要多少\n'
            '• 设定单价 — 给这个客户的报价\n'
            '• 系统会自动计算总金额'
        )

        self._add_tip_box(doc, '报价技巧',
            '• 报价前先了解客户预算，避免报价偏离太远\n'
            '• 备注里写清楚付款条款、交货期、运费说明等\n'
            '• 如果价格需要特批，先保存草稿，再发起批价申请',
            tip_type='tip'
        )

        # 5.3 发起批价
        doc.add_heading('5.3 价格需要审批？发起批价', level=2)

        doc.add_paragraph(
            '有时候为了拿下订单，需要给客户一些特别的折扣。'
            '如果你没有权限直接给出这个价格，就需要发起「批价」申请。'
        )

        doc.add_paragraph(
            '批价流程是这样的：\n'
            '1. 先创建报价单，把产品、数量、希望给客户的价格都填好\n'
            '2. 保存为草稿（别急着提交）\n'
            '3. 在报价详情页，点击「发起批价」\n'
            '4. 写清楚为什么需要这个价格（竞争情况、客户背景等）\n'
            '5. 提交，等待上级审批'
        )

        self._add_tip_box(doc, '批价申请要点',
            '审批人会看你的申请说明来判断是否批准，所以：\n'
            '• 说明竞争对手的报价情况\n'
            '• 解释这个客户的重要性和未来潜力\n'
            '• 如果是战略客户，说明长期价值\n'
            '信息越充分，审批越快',
            tip_type='info'
        )

        # 5.4 审批流程
        doc.add_heading('5.4 报价的生命周期', level=2)

        doc.add_paragraph(
            '一份报价从创建到最终结果，会经历不同的状态：'
        )

        doc.add_paragraph(
            '📝 草稿 — 刚创建，还没提交，可以随意修改\n\n'
            '⏳ 待审批 — 已提交批价申请，等待上级审批\n\n'
            '✅ 已批准 — 价格被批准了，可以正式发给客户\n\n'
            '❌ 已驳回 — 价格没批准，需要根据意见调整后重新申请\n\n'
            '📤 已发送 — 报价单已经发给客户了\n\n'
            '⌛ 已失效 — 超过有效期，如果客户还要，需要重新报价'
        )

        self._add_tip_box(doc, '跟进报价',
            '报价发出去不是结束，别忘了跟进客户反馈：\n'
            '• 发送后3天内主动问问客户收到没有\n'
            '• 了解客户对价格的看法\n'
            '• 及时在系统里更新报价状态',
            tip_type='tip'
        )

        doc.add_page_break()

    def _add_chapter_6_expense(self, doc):
        """第六章：报销管理"""
        doc.add_heading('第六章 报销管理', level=1)

        doc.add_paragraph(
            '销售工作少不了各种费用支出：拜访客户的差旅费、招待客户的餐费、'
            '日常的交通费...这些费用都可以通过报销系统申请报销。'
        )

        # 6.1 报销列表
        doc.add_heading('6.1 查看你的报销单', level=2)

        doc.add_paragraph(
            '点击左侧菜单的「报销管理」，可以看到你提交过的所有报销单。'
            '每张单子的审批状态一目了然：正在审批中、已通过、已打款等。'
        )

        self._add_image(doc, 'expense_list', '图6-1 报销列表页面')

        doc.add_paragraph(
            '列表显示的关键信息：\n'
            '• 报销单号 — 系统自动编号，有问题时报这个号\n'
            '• 总金额 — 这张报销单的总费用\n'
            '• 审批状态 — 看看走到哪一步了\n'
            '• 关联客户/项目 — 这笔费用是为哪个客户或项目花的'
        )

        # 6.2 新建报销
        doc.add_heading('6.2 提交一笔报销', level=2)

        doc.add_paragraph(
            '有发票要报销？点击「新建报销单」，开始填写：'
        )

        doc.add_paragraph(
            '第一步：选择费用类型\n'
            '• 差旅费 — 出差的机票、火车票、住宿、餐补\n'
            '• 招待费 — 请客户吃饭、商务宴请\n'
            '• 交通费 — 市内打车、地铁公交\n'
            '• 其他 — 办公用品、通讯费等\n\n'
            '第二步：填写费用明细\n'
            '• 费用日期 — 这笔费用发生在哪天\n'
            '• 金额 — 发票上的金额\n'
            '• 事由 — 简单说明为什么花这笔钱\n\n'
            '第三步：上传发票照片\n'
            '• 拍照或上传扫描件\n'
            '• 确保发票清晰、完整'
        )

        self._add_tip_box(doc, '关联客户/项目',
            '如果这笔费用是为了某个客户或项目产生的，记得关联一下。'
            '这样做有两个好处：\n'
            '• 便于统计每个客户/项目的投入成本\n'
            '• 领导审批时能看到花费的背景',
            tip_type='tip'
        )

        # 6.3 费用类型
        doc.add_heading('6.3 常见费用报销说明', level=2)

        doc.add_paragraph(
            '不同类型的费用，报销要求可能不一样：'
        )

        doc.add_paragraph(
            '🚄 差旅费\n'
            '出差产生的费用，通常包括往返交通、住宿、餐饮补贴。\n'
            '需要提供：车票/机票、住宿发票、行程单\n\n'
            '🍽️ 招待费\n'
            '招待客户的餐饮费用。\n'
            '需要说明：招待谁、什么事由、参加人数\n\n'
            '🚕 市内交通\n'
            '拜访客户的打车费、停车费等。\n'
            '需要提供：打车发票或电子行程单\n\n'
            '📱 通讯费\n'
            '工作相关的电话费、网络费。\n'
            '需要提供：运营商开具的发票'
        )

        # 6.4 审批流程
        doc.add_heading('6.4 报销单审批流程', level=2)

        doc.add_paragraph(
            '报销单提交后，会经过这样的审批流程：'
        )

        doc.add_paragraph(
            '1️⃣ 提交申请\n'
            '你填写完报销单并提交\n\n'
            '2️⃣ 上级审批\n'
            '你的直属领导会看一下费用是否合理\n\n'
            '3️⃣ 财务复核\n'
            '财务会检查发票是否合规\n\n'
            '4️⃣ 打款\n'
            '审批通过后，财务会把钱打到你的账户'
        )

        self._add_tip_box(doc, '报销小贴士',
            '• 发票抬头必须是公司全称，否则没法报\n'
            '• 尽量在费用发生后一个月内提交，拖太久可能报不了\n'
            '• 餐费发票如果金额较大，要写清楚招待对象\n'
            '• 提交前再检查一遍，发票照片要清晰\n'
            '• 如果被驳回，看清驳回原因，修改后重新提交',
            tip_type='warning'
        )

        doc.add_page_break()

    def _add_chapter_7_filter_search(self, doc):
        """第七章：筛选与搜索（通用功能）"""
        doc.add_heading('第七章 筛选与搜索', level=1)

        doc.add_paragraph(
            '当数据越来越多，逐个翻找既费时又容易遗漏。'
            '系统的筛选和搜索功能可以帮你快速定位想要的数据。'
            '这些功能在客户、项目、报价、报销等所有列表页面都是通用的。'
        )

        # 7.1 搜索功能使用
        doc.add_heading('7.1 快速搜索', level=2)

        doc.add_paragraph(
            '每个列表页面顶部都有一个搜索框。'
            '想找什么？直接输入关键词，按回车或点搜索按钮。'
        )

        doc.add_paragraph(
            '搜索框很智能：\n'
            '• 输入公司名称的一部分就能找到 — 比如输入"科技"能找到所有名称中包含"科技"的客户\n'
            '• 输入联系人姓名也能搜到对应的客户\n'
            '• 输入项目编号、报销单号都可以快速定位'
        )

        self._add_tip_box(doc, '搜索小技巧',
            '不确定完整名称？没关系，输入你记得的几个关键字就行。'
            '比如只记得客户名称里有"智能"两个字，直接搜"智能"就能找到。',
            tip_type='tip'
        )

        # 7.2 筛选条件说明
        doc.add_heading('7.2 用筛选缩小范围', level=2)

        doc.add_paragraph(
            '除了搜索，你还可以用筛选条件来过滤数据。'
            '筛选条件通常在列表上方，以下拉框的形式出现。'
        )

        doc.add_paragraph(
            '常用的筛选维度：\n\n'
            '📊 按状态筛选\n'
            '比如只看"待审批"的报销单，或者只看"进行中"的项目。\n\n'
            '👤 按负责人筛选\n'
            '想看某个同事的客户？选择他的名字就行。\n\n'
            '🏢 按类型筛选\n'
            '比如只看"代理商"类型的客户，或者"差旅费"类型的报销。\n\n'
            '🌍 按地区筛选\n'
            '比如只看北京地区的客户。'
        )

        # 7.3 组合筛选技巧
        doc.add_heading('7.3 组合使用筛选条件', level=2)

        doc.add_paragraph(
            '筛选条件可以叠加使用，帮你更精准地找到目标数据。'
        )

        doc.add_paragraph(
            '举个例子：你想找"上海地区的代理商客户"，就可以：\n'
            '1. 在地区筛选中选择"上海"\n'
            '2. 在企业类型中选择"代理"\n'
            '3. 系统会自动显示同时满足这两个条件的客户'
        )

        self._add_tip_box(doc, '筛选逻辑',
            '多个筛选条件是"并且"的关系，数据需要同时满足所有条件才会显示。'
            '如果发现结果为空，可能是条件太严格了，试着减少一两个条件。',
            tip_type='info'
        )

        # 7.4 重置与清除
        doc.add_heading('7.4 清除筛选条件', level=2)

        doc.add_paragraph(
            '筛选完想看全部数据？有两个方法：\n\n'
            '方法一：点击「重置」按钮\n'
            '一键清除所有筛选条件，恢复显示全部数据。\n\n'
            '方法二：逐个清除\n'
            '把某个下拉框选回"全部"选项，只取消这一个条件。'
        )

        doc.add_page_break()

    def _add_chapter_8_approval(self, doc):
        """第八章：审批流程（通用功能）"""
        doc.add_heading('第八章 审批流程', level=1)

        doc.add_paragraph(
            '项目立项、报价批价、费用报销...这些业务都需要走审批流程。'
            '这一章帮你了解审批是怎么运作的，不管是提交审批还是审批别人的申请。'
        )

        # 8.1 审批状态说明
        doc.add_heading('8.1 看懂审批状态', level=2)

        doc.add_paragraph(
            '每条需要审批的数据都有一个状态，告诉你它目前走到哪一步了：'
        )

        doc.add_paragraph(
            '📝 草稿\n'
            '刚创建，还没提交审批。这时候可以随便改。\n\n'
            '⏳ 待审批\n'
            '已经提交了，等着审批人来处理。\n\n'
            '🔄 审批中\n'
            '需要多级审批的流程，已经有人审批过了，还在继续往下走。\n\n'
            '✅ 已通过\n'
            '恭喜！审批通过了，可以正式执行了。\n\n'
            '❌ 已驳回\n'
            '审批没通过，看看审批人写的意见，改好后重新提交。\n\n'
            '↩️ 已撤回\n'
            '你自己撤回来了，变回草稿状态。'
        )

        # 8.2 提交审批
        doc.add_heading('8.2 如何提交审批', level=2)

        doc.add_paragraph(
            '信息都填好了，准备提交审批？记得先检查一遍：\n'
            '• 必填项都填了吗？\n'
            '• 金额算对了吗？\n'
            '• 需要上传的附件都传了吗？'
        )

        doc.add_paragraph(
            '确认无误后：\n'
            '1. 打开详情页\n'
            '2. 找到「提交审批」按钮（通常在右上角或页面底部）\n'
            '3. 点击后会弹出确认框，确认提交\n'
            '4. 提交成功！系统会自动通知审批人'
        )

        self._add_tip_box(doc, '提交后的注意事项',
            '• 提交后就不能直接修改了，如果发现有错，需要先撤回\n'
            '• 审批人收到通知后会来处理，耐心等待\n'
            '• 可以在详情页看到审批进度',
            tip_type='info'
        )

        # 8.3 撤回审批
        doc.add_heading('8.3 提交错了？可以撤回', level=2)

        doc.add_paragraph(
            '提交之后发现填错了信息？别担心，只要审批还没完成，你就可以撤回。'
        )

        doc.add_paragraph(
            '撤回方法：\n'
            '1. 打开那条记录的详情页\n'
            '2. 找到「撤回」按钮\n'
            '3. 确认撤回\n'
            '4. 状态变回"草稿"，可以继续修改了'
        )

        self._add_tip_box(doc, '撤回的限制',
            '• 只有你自己提交的才能撤回\n'
            '• 如果审批已经完成（通过或驳回），就不能撤回了\n'
            '• 撤回后修改完，记得重新提交',
            tip_type='warning'
        )

        # 8.4 审批操作（审批人）
        doc.add_heading('8.4 轮到你审批了怎么办', level=2)

        doc.add_paragraph(
            '如果你是某个流程的审批人，会收到系统通知。'
            '来看看怎么处理别人提交的审批请求：'
        )

        doc.add_paragraph(
            '1️⃣ 收到通知\n'
            '系统会提醒你有待审批的事项。\n\n'
            '2️⃣ 打开看看\n'
            '进入详情页，仔细看看提交的内容对不对。\n\n'
            '3️⃣ 做出决定\n'
            '• 没问题就点「同意」\n'
            '• 有问题就点「驳回」，并写清楚驳回原因\n\n'
            '4️⃣ 完成\n'
            '系统会通知提交人你的审批结果。'
        )

        self._add_tip_box(doc, '审批小建议',
            '• 驳回的时候务必写清楚原因，让对方知道该改什么\n'
            '• 有疑问可以先和提交人沟通，不要急着驳回\n'
            '• 尽快处理待审批事项，别让同事干等着',
            tip_type='tip'
        )

        # 8.5 审批流程图
        doc.add_heading('8.5 完整的审批流程', level=2)

        doc.add_paragraph(
            '一个典型的审批流程是这样的：'
        )

        doc.add_paragraph(
            '正常流程：\n'
            '草稿 → 提交 → 待审批 → 主管审批 → （可能还有更高层级） → 通过！\n\n'
            '被驳回了怎么办：\n'
            '已驳回 → 根据意见修改 → 重新提交 → 再次审批'
        )

        doc.add_paragraph(
            '几点说明：\n'
            '• 不同类型的业务，审批层级可能不一样\n'
            '• 金额越大，可能需要越高层级的审批\n'
            '• 每一步审批都会留下记录，谁什么时候审批的，写了什么意见'
        )

        doc.add_page_break()

    def _add_appendix(self, doc):
        """附录"""
        doc.add_heading('附录', level=1)

        # 常见问题
        doc.add_heading('常见问题解答', level=2)

        doc.add_paragraph(
            '使用系统过程中可能会遇到一些问题，这里整理了最常被问到的几个：'
        )

        doc.add_paragraph(
            '❓ 忘记密码了怎么办？\n'
            '联系系统管理员帮你重置。新密码会发到你的邮箱。\n\n'
            '❓ 想修改自己的个人信息\n'
            '点击页面右上角的用户头像，选择「个人设置」就可以修改了。\n\n'
            '❓ 报销单提交多久能批下来？\n'
            '一般1-3个工作日。如果比较急，可以直接联系审批人催一下。\n\n'
            '❓ 怎么把数据导出来？\n'
            '大部分列表页面都有「导出」按钮，可以导出Excel格式。\n\n'
            '❓ 手机上能用吗？\n'
            '可以的！用手机浏览器打开系统网址，界面会自动适配手机屏幕。\n\n'
            '❓ 数据安全吗？\n'
            '放心，系统采用加密传输，数据会定期备份，不会丢失。'
        )

        # 快捷操作
        doc.add_heading('提高效率的小技巧', level=2)

        doc.add_paragraph(
            '掌握这些小技巧，操作起来更顺手：\n\n'
            '⌨️ 搜索快捷键\n'
            '大部分页面按 Ctrl+F（Mac 用 Cmd+F）可以快速打开搜索框。\n\n'
            '↩️ 回车确认\n'
            '填完表单后，按回车键等于点击确认按钮。\n\n'
            '❌ Esc 取消\n'
            '弹窗打开后，按 Esc 键可以快速关闭。\n\n'
            '➡️ Tab 切换\n'
            '填写表单时，按 Tab 键可以跳到下一个输入框，不用每次点鼠标。'
        )

        # 技术支持
        doc.add_heading('遇到问题找谁？', level=2)

        doc.add_paragraph(
            '如果在使用过程中遇到问题或有建议，可以通过以下方式联系我们：\n\n'
            '📧 系统管理员邮箱：admin@company.com\n'
            '📞 IT支持热线：内线 8888\n'
            '🕐 服务时间：周一至周五 9:00-18:00\n\n'
            '描述问题时，如果能附上截图和你的操作步骤，我们能更快帮你解决！'
        )

    async def run(self):
        """运行生成器"""
        # 1. 采集截图
        await self.capture_all_screenshots()

        # 2. 生成文档
        output_path = self.generate_word_document()

        return output_path


async def main():
    """主函数"""
    generator = SalesTutorialGeneratorV2(
        base_url="http://localhost:5011",
        username="admin",
        password="1505562299AaBb"
    )

    output_path = await generator.run()

    if output_path:
        print("\n" + "=" * 60)
        print("销售操作手册（详细版）生成完成！")
        print(f"文件位置: {output_path}")
        print("=" * 60)

        # 打开文件
        os.system(f'open "{output_path}"')


if __name__ == '__main__':
    asyncio.run(main())
