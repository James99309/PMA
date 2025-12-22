#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
销售人员操作手册生成器
自动截图 + 生成 Word 文档
"""
import os
import sys
import asyncio
from datetime import datetime

# 路径修正
def get_project_root():
    current = os.path.dirname(os.path.abspath(__file__))
    while current != '/':
        if os.path.exists(os.path.join(current, 'app')) and \
           os.path.exists(os.path.join(current, 'run.py')):
            return current
        current = os.path.dirname(current)
    raise RuntimeError("无法找到项目根目录")

PROJECT_ROOT = get_project_root()
sys.path.insert(0, PROJECT_ROOT)

from scripts.tools.screenshot_tool import ScreenshotTool
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


class SalesTutorialGenerator:
    """销售操作手册生成器"""

    def __init__(self, base_url: str, username: str, password: str):
        self.base_url = base_url
        self.username = username
        self.password = password
        self.screenshot_dir = os.path.join(PROJECT_ROOT, 'docs/screenshots/tutorial')
        self.output_dir = os.path.join(PROJECT_ROOT, 'docs')
        self.screenshots = {}

        os.makedirs(self.screenshot_dir, exist_ok=True)

    async def capture_all_screenshots(self):
        """截取所有需要的页面"""
        tool = ScreenshotTool(
            base_url=self.base_url,
            username=self.username,
            password=self.password,
            output_dir=self.screenshot_dir,
            headless=True,  # 无头模式，加快速度
            language="zh",
            mask_data=True
        )

        try:
            print("=" * 60)
            print("开始截取教程截图...")
            print("=" * 60)

            await tool.start()
            logged_in = await tool.login()
            if not logged_in:
                print("登录失败！")
                return False

            # 定义需要截取的页面
            pages = [
                # 登录页面（需要先登出）
                # ('login', '/auth/login', '登录页面'),

                # 仪表台
                ('dashboard', '/', '仪表台'),

                # 客户管理
                ('customer_list', '/customer/', '客户列表'),
                ('customer_add', '/customer/1/view', '客户详情'),  # 查看客户

                # 项目管理
                ('project_list', '/project/', '项目列表'),

                # 报价管理
                ('quotation_list', '/quotation/quotations', '报价列表'),

                # 报销管理
                ('expense_list', '/expense/', '报销列表'),
            ]

            for name, path, desc in pages:
                print(f"\n截取: {desc} ({path})")
                try:
                    filepath = await tool.screenshot(
                        path=path,
                        filename=name,
                        wait_time=2000
                    )
                    self.screenshots[name] = filepath
                    print(f"  ✓ 已保存: {filepath}")
                except Exception as e:
                    print(f"  ✗ 截取失败: {e}")

            print("\n" + "=" * 60)
            print(f"截图完成！共 {len(self.screenshots)} 张")
            print("=" * 60)

            return True

        finally:
            await tool.close()

    def generate_word_document(self):
        """生成 Word 文档"""
        print("\n开始生成 Word 文档...")

        doc = Document()

        # 设置中文字体
        self._set_chinese_font(doc)

        # 封面
        self._add_cover(doc)

        # 目录页
        self._add_toc(doc)

        # 第一章：快速入门
        self._add_chapter_quick_start(doc)

        # 第二章：仪表台
        self._add_chapter_dashboard(doc)

        # 第三章：客户管理
        self._add_chapter_customer(doc)

        # 第四章：项目管理
        self._add_chapter_project(doc)

        # 第五章：报价管理
        self._add_chapter_quotation(doc)

        # 第六章：报销管理
        self._add_chapter_expense(doc)

        # 附录
        self._add_appendix(doc)

        # 保存文档
        output_path = os.path.join(self.output_dir, 'PMA销售人员操作手册.docx')
        doc.save(output_path)
        print(f"\n文档已保存: {output_path}")

        return output_path

    def _set_chinese_font(self, doc):
        """设置中文字体"""
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _add_cover(self, doc):
        """添加封面"""
        # 添加空行
        for _ in range(6):
            doc.add_paragraph()

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('PMA 系统')
        run.bold = True
        run.font.size = Pt(36)
        run.font.name = '微软雅黑'

        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('销售人员操作手册')
        run.bold = True
        run.font.size = Pt(24)
        run.font.name = '微软雅黑'

        # 空行
        for _ in range(4):
            doc.add_paragraph()

        # 版本信息
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f'版本: V1.0\n生成日期: {datetime.now().strftime("%Y年%m月%d日")}')
        run.font.size = Pt(12)

        # 分页
        doc.add_page_break()

    def _add_toc(self, doc):
        """添加目录"""
        doc.add_heading('目录', level=1)

        toc_items = [
            ('第一章 快速入门', 1),
            ('1.1 系统登录', 2),
            ('1.2 界面概览', 2),
            ('1.3 常用操作', 2),
            ('第二章 仪表台', 1),
            ('2.1 数据概览', 2),
            ('2.2 待办事项', 2),
            ('第三章 客户管理', 1),
            ('3.1 客户列表', 2),
            ('3.2 新增客户', 2),
            ('3.3 客户详情', 2),
            ('3.4 联系人管理', 2),
            ('第四章 项目管理', 1),
            ('4.1 项目列表', 2),
            ('4.2 创建项目', 2),
            ('4.3 项目跟进', 2),
            ('第五章 报价管理', 1),
            ('5.1 报价列表', 2),
            ('5.2 创建报价', 2),
            ('5.3 发起批价', 2),
            ('第六章 报销管理', 1),
            ('6.1 报销列表', 2),
            ('6.2 新建报销', 2),
            ('6.3 审批流程', 2),
            ('附录', 1),
        ]

        for item, level in toc_items:
            p = doc.add_paragraph()
            if level == 1:
                run = p.add_run(item)
                run.bold = True
            else:
                p.add_run('    ' + item)

        doc.add_page_break()

    def _add_chapter_quick_start(self, doc):
        """第一章：快速入门"""
        doc.add_heading('第一章 快速入门', level=1)

        # 1.1 系统登录
        doc.add_heading('1.1 系统登录', level=2)

        doc.add_paragraph('打开浏览器，输入系统地址进入登录页面。')

        # 登录步骤
        doc.add_paragraph('登录步骤：', style='List Number')
        steps = [
            '输入用户名（通常为您的工号或邮箱）',
            '输入密码',
            '点击"登录"按钮',
            '首次登录建议修改默认密码'
        ]
        for step in steps:
            doc.add_paragraph(step, style='List Bullet')

        # 注意事项
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('注意事项：')
        run.bold = True

        notes = [
            '连续输错密码5次，账户将被锁定30分钟',
            '如忘记密码，请联系系统管理员重置',
            '建议使用 Chrome 或 Edge 浏览器访问系统'
        ]
        for note in notes:
            doc.add_paragraph('• ' + note)

        # 1.2 界面概览
        doc.add_heading('1.2 界面概览', level=2)

        doc.add_paragraph('登录成功后，您将看到系统主界面，主要包含以下区域：')

        # 添加仪表台截图
        if 'dashboard' in self.screenshots:
            doc.add_picture(self.screenshots['dashboard'], width=Inches(6))
            doc.add_paragraph('图1-1 系统主界面', style='Caption')

        # 界面说明表格
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = '区域'
        headers[1].text = '说明'

        data = [
            ('左侧导航栏', '包含所有功能模块的入口，点击可展开子菜单'),
            ('顶部搜索栏', '全局搜索功能，可快速查找客户、项目等'),
            ('右上角用户区', '显示当前用户，可切换语言、修改密码、退出登录'),
            ('主内容区', '显示当前模块的具体内容'),
        ]
        for i, (area, desc) in enumerate(data, 1):
            row = table.rows[i].cells
            row[0].text = area
            row[1].text = desc

        # 1.3 常用操作
        doc.add_heading('1.3 常用操作', level=2)

        doc.add_paragraph('作为销售人员，您日常主要使用以下功能：')

        operations = [
            ('客户管理', '记录和管理客户信息、联系人、拜访记录'),
            ('项目管理', '跟踪销售项目进度、关键节点'),
            ('报价管理', '创建报价单、发起批价申请'),
            ('报销管理', '提交差旅费、招待费等报销申请'),
        ]

        table = doc.add_table(rows=len(operations)+1, cols=2)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = '功能模块'
        headers[1].text = '主要用途'

        for i, (module, usage) in enumerate(operations, 1):
            row = table.rows[i].cells
            row[0].text = module
            row[1].text = usage

        doc.add_page_break()

    def _add_chapter_dashboard(self, doc):
        """第二章：仪表台"""
        doc.add_heading('第二章 仪表台', level=1)

        doc.add_paragraph('仪表台是您登录系统后的首页，提供业务数据的快速概览和待办事项提醒。')

        # 2.1 数据概览
        doc.add_heading('2.1 数据概览', level=2)

        if 'dashboard' in self.screenshots:
            doc.add_picture(self.screenshots['dashboard'], width=Inches(6))
            doc.add_paragraph('图2-1 仪表台数据概览', style='Caption')

        doc.add_paragraph('仪表台显示以下关键数据：')

        items = [
            '客户统计：全部客户数、活跃客户数、本月新增',
            '项目统计：进行中项目数、各阶段项目分布',
            '报价统计：待审批报价、本月报价总额',
            '工作记录：最近的客户拜访、跟进记录'
        ]
        for item in items:
            doc.add_paragraph('• ' + item)

        # 2.2 待办事项
        doc.add_heading('2.2 待办事项', level=2)

        doc.add_paragraph('待办事项区域显示需要您处理的工作：')

        items = [
            '待审批的报价单',
            '待跟进的项目',
            '即将到期的客户拜访',
            '待处理的报销单'
        ]
        for item in items:
            doc.add_paragraph('• ' + item)

        doc.add_page_break()

    def _add_chapter_customer(self, doc):
        """第三章：客户管理"""
        doc.add_heading('第三章 客户管理', level=1)

        doc.add_paragraph('客户管理模块用于记录和维护所有客户信息，包括公司信息、联系人、拜访记录等。')

        # 3.1 客户列表
        doc.add_heading('3.1 客户列表', level=2)

        doc.add_paragraph('点击左侧菜单"客户管理"进入客户列表页面。')

        if 'customer_list' in self.screenshots:
            doc.add_picture(self.screenshots['customer_list'], width=Inches(6))
            doc.add_paragraph('图3-1 客户列表页面', style='Caption')

        doc.add_paragraph('客户列表功能说明：')

        features = [
            ('筛选功能', '可按负责人、企业类型、行业、地区、状态筛选客户'),
            ('搜索功能', '在搜索框输入客户名称快速查找'),
            ('排序功能', '点击表头可按该列排序'),
            ('导出功能', '可将客户列表导出为Excel文件'),
        ]

        table = doc.add_table(rows=len(features)+1, cols=2)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = '功能'
        headers[1].text = '说明'

        for i, (feat, desc) in enumerate(features, 1):
            row = table.rows[i].cells
            row[0].text = feat
            row[1].text = desc

        # 3.2 新增客户
        doc.add_heading('3.2 新增客户', level=2)

        doc.add_paragraph('点击"添加客户"按钮，进入新增客户页面。')

        doc.add_paragraph('新增客户步骤：')
        steps = [
            '点击右上角"添加客户"按钮',
            '填写客户基本信息（公司名称为必填项）',
            '选择企业类型、行业分类',
            '填写联系地址、备注等信息',
            '点击"保存"完成创建'
        ]
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f'{i}. {step}')

        # 必填字段说明
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('必填字段：')
        run.bold = True
        doc.add_paragraph('• 公司名称')
        doc.add_paragraph('• 企业类型')

        # 3.3 客户详情
        doc.add_heading('3.3 客户详情', level=2)

        doc.add_paragraph('点击客户名称可查看客户详细信息。')

        if 'customer_add' in self.screenshots:
            doc.add_picture(self.screenshots['customer_add'], width=Inches(6))
            doc.add_paragraph('图3-2 客户详情页面', style='Caption')

        doc.add_paragraph('客户详情页包含：')
        items = [
            '基本信息：公司名称、类型、行业、地址等',
            '联系人列表：该客户的所有联系人',
            '项目关联：与该客户相关的项目',
            '拜访记录：历史拜访和跟进记录',
            '操作记录：客户信息的修改历史'
        ]
        for item in items:
            doc.add_paragraph('• ' + item)

        # 3.4 联系人管理
        doc.add_heading('3.4 联系人管理', level=2)

        doc.add_paragraph('每个客户可以添加多个联系人。')

        doc.add_paragraph('添加联系人步骤：')
        steps = [
            '进入客户详情页',
            '找到"联系人"区域',
            '点击"添加联系人"按钮',
            '填写联系人信息（姓名、职位、电话、邮箱等）',
            '点击"保存"'
        ]
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f'{i}. {step}')

        doc.add_page_break()

    def _add_chapter_project(self, doc):
        """第四章：项目管理"""
        doc.add_heading('第四章 项目管理', level=1)

        doc.add_paragraph('项目管理模块用于跟踪销售项目的全生命周期，从商机发现到最终成交。')

        # 4.1 项目列表
        doc.add_heading('4.1 项目列表', level=2)

        doc.add_paragraph('点击左侧菜单"项目管理"进入项目列表。')

        if 'project_list' in self.screenshots:
            doc.add_picture(self.screenshots['project_list'], width=Inches(6))
            doc.add_paragraph('图4-1 项目列表页面', style='Caption')

        doc.add_paragraph('项目列表显示以下信息：')
        items = [
            '项目名称：点击可查看详情',
            '关联客户：项目所属的客户',
            '项目阶段：当前所处的销售阶段',
            '预计金额：项目预估成交金额',
            '负责人：项目负责的销售人员',
            '更新时间：最近一次更新时间'
        ]
        for item in items:
            doc.add_paragraph('• ' + item)

        # 4.2 创建项目
        doc.add_heading('4.2 创建项目', level=2)

        doc.add_paragraph('创建新项目步骤：')
        steps = [
            '点击"创建项目"按钮',
            '填写项目名称（必填）',
            '选择关联客户（必填）',
            '设置项目类型、阶段',
            '填写预计金额、预计成交日期',
            '点击"保存"完成创建'
        ]
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f'{i}. {step}')

        # 项目阶段说明
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('项目阶段说明：')
        run.bold = True

        stages = [
            ('商机发现', '初步接触客户，了解需求'),
            ('需求确认', '深入了解客户需求，确认项目可行性'),
            ('方案制定', '制定解决方案，准备报价'),
            ('商务谈判', '报价、谈判、合同协商'),
            ('签约成交', '签订合同，项目成交'),
        ]

        table = doc.add_table(rows=len(stages)+1, cols=2)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = '阶段'
        headers[1].text = '说明'

        for i, (stage, desc) in enumerate(stages, 1):
            row = table.rows[i].cells
            row[0].text = stage
            row[1].text = desc

        # 4.3 项目跟进
        doc.add_heading('4.3 项目跟进', level=2)

        doc.add_paragraph('及时记录项目跟进情况，有助于团队协作和复盘分析。')

        doc.add_paragraph('添加跟进记录：')
        steps = [
            '进入项目详情页',
            '点击"添加跟进记录"',
            '选择跟进类型（电话、拜访、邮件等）',
            '填写跟进内容和下一步计划',
            '点击"保存"'
        ]
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f'{i}. {step}')

        doc.add_page_break()

    def _add_chapter_quotation(self, doc):
        """第五章：报价管理"""
        doc.add_heading('第五章 报价管理', level=1)

        doc.add_paragraph('报价管理模块用于创建、管理和审批报价单。')

        # 5.1 报价列表
        doc.add_heading('5.1 报价列表', level=2)

        doc.add_paragraph('点击左侧菜单"报价管理"进入报价列表。')

        if 'quotation_list' in self.screenshots:
            doc.add_picture(self.screenshots['quotation_list'], width=Inches(6))
            doc.add_paragraph('图5-1 报价列表页面', style='Caption')

        # 5.2 创建报价
        doc.add_heading('5.2 创建报价', level=2)

        doc.add_paragraph('创建报价单步骤：')
        steps = [
            '点击"新建报价"按钮',
            '选择关联项目（可选）和客户',
            '添加报价明细（产品、数量、单价）',
            '设置有效期、付款条款等',
            '预览报价单内容',
            '点击"保存"或"提交审批"'
        ]
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f'{i}. {step}')

        # 5.3 发起批价
        doc.add_heading('5.3 发起批价', level=2)

        doc.add_paragraph('当报价需要特殊折扣或超出权限时，需要发起批价申请。')

        doc.add_paragraph('发起批价步骤：')
        steps = [
            '创建报价单并保存为草稿',
            '点击"发起批价"按钮',
            '填写批价原因和说明',
            '选择审批人（系统会根据金额自动匹配）',
            '提交等待审批',
            '审批通过后可正式发送给客户'
        ]
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f'{i}. {step}')

        # 批价状态说明
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('报价状态说明：')
        run.bold = True

        statuses = [
            ('草稿', '报价单已创建但未提交'),
            ('待审批', '已提交批价申请，等待审批'),
            ('已批准', '批价申请已通过'),
            ('已驳回', '批价申请被驳回，需修改后重新提交'),
            ('已发送', '报价单已发送给客户'),
            ('已过期', '报价单已超过有效期'),
        ]

        table = doc.add_table(rows=len(statuses)+1, cols=2)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = '状态'
        headers[1].text = '说明'

        for i, (status, desc) in enumerate(statuses, 1):
            row = table.rows[i].cells
            row[0].text = status
            row[1].text = desc

        doc.add_page_break()

    def _add_chapter_expense(self, doc):
        """第六章：报销管理"""
        doc.add_heading('第六章 报销管理', level=1)

        doc.add_paragraph('报销管理模块用于提交和审批各类费用报销申请。')

        # 6.1 报销列表
        doc.add_heading('6.1 报销列表', level=2)

        doc.add_paragraph('点击左侧菜单"报销管理"进入报销列表。')

        if 'expense_list' in self.screenshots:
            doc.add_picture(self.screenshots['expense_list'], width=Inches(6))
            doc.add_paragraph('图6-1 报销列表页面', style='Caption')

        doc.add_paragraph('报销列表显示您提交的所有报销单，包括：')
        items = [
            '报销单号：系统自动生成的唯一编号',
            '报销金额：本次报销的总金额',
            '关联客户：费用发生的客户（如有）',
            '审批状态：当前审批进度',
            '创建时间：提交报销的时间'
        ]
        for item in items:
            doc.add_paragraph('• ' + item)

        # 6.2 新建报销
        doc.add_heading('6.2 新建报销', level=2)

        doc.add_paragraph('新建报销单步骤：')
        steps = [
            '点击"新建报销单"按钮',
            '选择费用类型（差旅费、招待费、办公费等）',
            '填写费用明细（日期、金额、说明）',
            '上传发票或收据照片',
            '关联客户或项目（可选）',
            '点击"提交"发起审批'
        ]
        for i, step in enumerate(steps, 1):
            doc.add_paragraph(f'{i}. {step}')

        # 费用类型说明
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('常见费用类型：')
        run.bold = True

        expense_types = [
            ('差旅费', '出差产生的交通、住宿费用'),
            ('招待费', '客户招待、商务宴请费用'),
            ('交通费', '市内交通、打车费用'),
            ('通讯费', '电话费、网络费'),
            ('办公费', '办公用品采购费用'),
        ]

        table = doc.add_table(rows=len(expense_types)+1, cols=2)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = '类型'
        headers[1].text = '说明'

        for i, (etype, desc) in enumerate(expense_types, 1):
            row = table.rows[i].cells
            row[0].text = etype
            row[1].text = desc

        # 6.3 审批流程
        doc.add_heading('6.3 审批流程', level=2)

        doc.add_paragraph('报销单提交后将进入审批流程：')

        doc.add_paragraph('审批流程：')
        flow = [
            '提交报销 → 直属上级审批',
            '金额超过一定额度 → 部门经理审批',
            '审批通过 → 财务复核',
            '财务复核通过 → 等待打款'
        ]
        for step in flow:
            doc.add_paragraph('• ' + step)

        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('注意事项：')
        run.bold = True

        notes = [
            '报销申请提交后可在"我的报销"中查看审批进度',
            '如审批被驳回，可修改后重新提交',
            '发票必须清晰可辨认，否则可能被退回',
            '报销金额需与发票金额一致'
        ]
        for note in notes:
            doc.add_paragraph('• ' + note)

        doc.add_page_break()

    def _add_appendix(self, doc):
        """添加附录"""
        doc.add_heading('附录', level=1)

        # 常见问题
        doc.add_heading('常见问题', level=2)

        faqs = [
            ('Q: 忘记密码怎么办？', 'A: 请联系系统管理员重置密码。'),
            ('Q: 如何修改个人信息？', 'A: 点击右上角用户头像，选择"个人设置"进行修改。'),
            ('Q: 报销单审批需要多长时间？', 'A: 一般情况下1-3个工作日内完成审批。'),
            ('Q: 如何导出客户数据？', 'A: 在客户列表页面点击"导出"按钮即可导出Excel文件。'),
            ('Q: 系统支持哪些浏览器？', 'A: 推荐使用 Chrome、Edge 或 Firefox 最新版本。'),
        ]

        for q, a in faqs:
            p = doc.add_paragraph()
            run = p.add_run(q)
            run.bold = True
            doc.add_paragraph(a)
            doc.add_paragraph()

        # 联系方式
        doc.add_heading('技术支持', level=2)

        doc.add_paragraph('如在使用过程中遇到问题，请联系：')
        doc.add_paragraph('• 系统管理员：admin@company.com')
        doc.add_paragraph('• IT支持热线：内线 8888')

    async def run(self):
        """运行生成器"""
        # 1. 截取截图
        success = await self.capture_all_screenshots()
        if not success:
            print("截图失败，无法继续生成文档")
            return None

        # 2. 生成 Word 文档
        output_path = self.generate_word_document()

        return output_path


async def main():
    """主函数"""
    generator = SalesTutorialGenerator(
        base_url="http://localhost:5011",
        username="admin",
        password="1505562299AaBb"
    )

    output_path = await generator.run()

    if output_path:
        print("\n" + "=" * 60)
        print("销售操作手册生成完成！")
        print(f"文件位置: {output_path}")
        print("=" * 60)

        # 打开文件所在目录
        os.system(f'open "{os.path.dirname(output_path)}"')


if __name__ == '__main__':
    asyncio.run(main())
