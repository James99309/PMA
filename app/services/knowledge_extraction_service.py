# -*- coding: utf-8 -*-
"""
知识库文档文本提取服务

支持: PDF, DOCX, XLSX/CSV, TXT/MD, HTML
"""
import io
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str, mime_type: str = '') -> Optional[str]:
    """
    根据文件类型提取纯文本内容。

    Args:
        file_bytes: 文件二进制内容
        filename: 文件名（用于判断扩展名）
        mime_type: MIME 类型

    Returns:
        提取的纯文本，失败返回 None
    """
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == '.pdf' or 'pdf' in (mime_type or ''):
            return _extract_pdf(file_bytes)
        elif ext in ('.docx',) or 'wordprocessingml' in (mime_type or ''):
            return _extract_docx(file_bytes)
        elif ext in ('.xlsx', '.xls') or 'spreadsheet' in (mime_type or ''):
            return _extract_xlsx(file_bytes)
        elif ext == '.csv' or 'csv' in (mime_type or ''):
            return _extract_csv(file_bytes)
        elif ext in ('.html', '.htm') or 'html' in (mime_type or ''):
            return _extract_html(file_bytes)
        elif ext in ('.txt', '.md', '.log', '.json', '.xml', '.yaml', '.yml'):
            return _extract_text_plain(file_bytes)
        else:
            logger.warning(f"不支持的文件类型: {filename} ({mime_type})")
            return None
    except Exception as e:
        logger.error(f"提取文本失败 [{filename}]: {e}")
        return None


# ── PDF ──────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> Optional[str]:
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[第{i}页]\n{text.strip()}")
    return '\n\n'.join(pages) if pages else None


# ── DOCX ─────────────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes) -> Optional[str]:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(' | '.join(cells))

    return '\n\n'.join(paragraphs) if paragraphs else None


# ── XLSX ─────────────────────────────────────────────────────────

def _extract_xlsx(file_bytes: bytes) -> Optional[str]:
    import pandas as pd

    sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, dtype=str)
    parts = []
    for sheet_name, df in sheets.items():
        df = df.fillna('')
        if df.empty:
            continue
        header = ' | '.join(str(c) for c in df.columns)
        rows = []
        for _, row in df.iterrows():
            row_text = ' | '.join(str(v) for v in row.values if str(v).strip())
            if row_text.strip():
                rows.append(row_text)
        if rows:
            parts.append(f"[工作表: {sheet_name}]\n{header}\n" + '\n'.join(rows))
    return '\n\n'.join(parts) if parts else None


# ── CSV ──────────────────────────────────────────────────────────

def _extract_csv(file_bytes: bytes) -> Optional[str]:
    import pandas as pd

    # 尝试多种编码
    for enc in ('utf-8', 'gbk', 'gb2312', 'latin1'):
        try:
            df = pd.read_csv(io.BytesIO(file_bytes), encoding=enc, dtype=str)
            break
        except (UnicodeDecodeError, Exception):
            continue
    else:
        return None

    df = df.fillna('')
    if df.empty:
        return None

    header = ' | '.join(str(c) for c in df.columns)
    rows = []
    for _, row in df.iterrows():
        row_text = ' | '.join(str(v) for v in row.values if str(v).strip())
        if row_text.strip():
            rows.append(row_text)
    return f"{header}\n" + '\n'.join(rows) if rows else None


# ── HTML ─────────────────────────────────────────────────────────

def _extract_html(file_bytes: bytes) -> Optional[str]:
    from bs4 import BeautifulSoup

    # 尝试多种编码
    for enc in ('utf-8', 'gbk', 'latin1'):
        try:
            html_text = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        return None

    soup = BeautifulSoup(html_text, 'html.parser')
    # 移除 script/style
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    text = soup.get_text(separator='\n', strip=True)
    return text if text.strip() else None


# ── 纯文本 ───────────────────────────────────────────────────────

def _extract_text_plain(file_bytes: bytes) -> Optional[str]:
    for enc in ('utf-8', 'gbk', 'gb2312', 'latin1'):
        try:
            text = file_bytes.decode(enc)
            return text if text.strip() else None
        except UnicodeDecodeError:
            continue
    return None
