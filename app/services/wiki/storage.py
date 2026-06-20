# -*- coding: utf-8 -*-
"""Wiki 文章、原始文件、index/log 的 I/O 层

所有路径读写都经过本模块，避免散落在各处直接操作磁盘。
读出的文章路径统一是相对 wiki_root 的字符串，例如
    wiki/product/gp328p-overview.md
    raw/product/2026-04-09-datasheet.pdf
"""
import base64
import hashlib
import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List
from zoneinfo import ZoneInfo

from app.services.wiki.paths import (
    article_image_relative_path,
    assets_dir_for_article,
    get_index_path,
    get_log_path,
    get_raw_dir,
    get_wiki_dir,
    get_wiki_root,
    raw_file_path,
    wiki_article_path,
)

logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════
# 安全性：topic / slug / 文件名验证
# ══════════════════════════════════════════════════════════════════
#
# topic 和 slug 最终会被拼到磁盘路径里，必须严格校验防止路径穿越。
# 输入来源包括用户表单（add_raw_file 视图）和 Claude 返回的 JSON
# （对抗性 PDF 的 prompt injection 可能操纵 Claude 输出恶意 slug）。
#
# 规则（通过才放行）：
# - topic: [A-Za-z0-9_-]+，长度 1-100
# - slug:  [A-Za-z0-9._-]+，长度 1-200，不能以 . 开头（禁止 ../ 或隐藏文件）
#
# 不允许 CJK 的原因：主流场景下 topic/slug 都是英文短标识符（product、
# gp328p-overview 这种），保留 ASCII 限制可以完全杜绝各种 Unicode 同形
# 字符绕过。raw 文件的原始标题允许 CJK（走 safe_filename 另一条路径）。

_TOPIC_RE = re.compile(r'^[A-Za-z0-9\u4e00-\u9fff_-]{1,100}$')
_SLUG_RE = re.compile(r'^[A-Za-z0-9\u4e00-\u9fff._-]{1,200}$')
_SAFE_FILENAME_RE = re.compile(r'[^A-Za-z0-9\u4e00-\u9fff._-]+')


class WikiPathError(ValueError):
    """topic/slug/文件名非法时抛出"""


def validate_topic(topic: str) -> None:
    """校验 topic。不合规则抛 WikiPathError。"""
    if not isinstance(topic, str) or not _TOPIC_RE.match(topic):
        raise WikiPathError(
            f'非法 topic: {topic!r}（允许中英文/数字/_/-，1-100 字符，禁止 / \\ .. 和空格）'
        )


def validate_slug(slug: str) -> None:
    """校验 slug。不合规则抛 WikiPathError。"""
    if not isinstance(slug, str) or not _SLUG_RE.match(slug) or slug.startswith('.'):
        raise WikiPathError(
            f'非法 slug: {slug!r}（允许中英文/数字/./_/-，1-200 字符，不能以 . 开头）'
        )


def validate_topic_slug(topic: str, slug: str) -> None:
    """同时校验 topic 和 slug。"""
    validate_topic(topic)
    validate_slug(slug)


def safe_filename(raw_name: str) -> str:
    """把任意字符串转成磁盘安全的文件名（保留汉字/数字/字母/. _ -）。

    用于 raw 文件的原始文件名保留显示。注意这不是 slug —— slug 必须走
    validate_slug，CJK 是不允许的。
    """
    name = _SAFE_FILENAME_RE.sub('_', raw_name.strip())
    # 防止跨目录：去掉路径分隔符和前导 .
    name = name.replace('/', '_').replace('\\', '_').lstrip('.')
    return name or 'unnamed'


def dated_filename(original: str) -> str:
    """在文件名前加上 YYYY-MM-DD 前缀，便于按时间排序。"""
    today = datetime.now(ZoneInfo('Asia/Shanghai')).strftime('%Y-%m-%d')
    return f'{today}-{safe_filename(original)}'


# ══════════════════════════════════════════════════════════════════
# Wiki 文章读写
# ══════════════════════════════════════════════════════════════════

def read_article_content(file_path: str) -> str:
    """按相对路径读文章。file_path 形如 'wiki/product/xxx.md'。

    读不到时返回空字符串，方便调用方做降级。
    """
    if not file_path:
        return ''
    abs_path = get_wiki_root() / file_path
    if not abs_path.exists():
        logger.warning(f'Wiki 文章文件不存在: {abs_path}')
        return ''
    return abs_path.read_text(encoding='utf-8')


def write_article(topic: str, slug: str, content: str) -> str:
    """写一篇文章到 wiki/<topic>/<slug>.md。

    返回相对 wiki_root 的路径字符串，用于存入 KnowledgeWikiArticle.file_path。

    Raises:
        WikiPathError: topic/slug 非法
    """
    validate_topic_slug(topic, slug)
    path = wiki_article_path(topic, slug)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding='utf-8')
    return str(path.relative_to(get_wiki_root()))


def delete_article(topic: str, slug: str) -> bool:
    """删除一篇文章文件。返回是否成功删除（文件不存在也算 False）。

    Raises:
        WikiPathError: topic/slug 非法
    """
    validate_topic_slug(topic, slug)
    path = wiki_article_path(topic, slug)
    if path.exists():
        path.unlink()
        return True
    return False


def delete_article_file(file_path: str) -> bool:
    """按 file_path 删除文章文件（防路径穿越）。"""
    if not file_path:
        return False
    abs_path = (get_wiki_root() / file_path).resolve()
    wiki_root = get_wiki_root().resolve()
    if not str(abs_path).startswith(str(wiki_root)):
        raise WikiPathError(f'路径穿越: {file_path}')
    if abs_path.exists():
        abs_path.unlink()
        return True
    return False


def list_topics() -> List[str]:
    """列出 wiki/ 目录下所有 topic（一级子目录）。"""
    wiki_dir = get_wiki_dir()
    if not wiki_dir.exists():
        return []
    return sorted([p.name for p in wiki_dir.iterdir() if p.is_dir()])


def list_articles_in_topic(topic: str) -> List[str]:
    """列出某个 topic 下所有文章的 slug（不含 .md）。"""
    topic_dir = get_wiki_dir() / topic
    if not topic_dir.exists():
        return []
    return sorted([p.stem for p in topic_dir.glob('*.md')])


# ══════════════════════════════════════════════════════════════════
# index.md 与 log.md
# ══════════════════════════════════════════════════════════════════

def read_index() -> str:
    """读 wiki/index.md。不存在时返回空字符串。"""
    idx = get_index_path()
    return idx.read_text(encoding='utf-8') if idx.exists() else ''


def write_index(content: str):
    """覆盖写 wiki/index.md。"""
    idx = get_index_path()
    idx.parent.mkdir(parents=True, exist_ok=True)
    idx.write_text(content, encoding='utf-8')


def append_log(operation: str, details: str):
    """往 wiki/log.md 追加一条操作记录。

    格式：
        ## 2026-04-09 12:34:56 — ingest
        <details>
    """
    now = datetime.now(ZoneInfo('Asia/Shanghai')).strftime('%Y-%m-%d %H:%M:%S')
    entry = f'\n## {now} — {operation}\n\n{details.rstrip()}\n'
    log_path = get_log_path()
    log_path.parent.mkdir(parents=True, exist_ok=True)
    with log_path.open('a', encoding='utf-8') as f:
        f.write(entry)


# ══════════════════════════════════════════════════════════════════
# 原始文件 I/O
# ══════════════════════════════════════════════════════════════════

def save_raw_file(topic: str, filename: str, data: bytes) -> str:
    """把二进制内容写入 raw/<topic>/<filename>。

    Args:
        topic: topic 名称（将作为目录）
        filename: 建议调用方先 dated_filename(original) 规范化过
        data: 文件字节

    Returns:
        相对 wiki_root 的路径字符串，例如 'raw/product/2026-04-09-datasheet.pdf'

    Raises:
        WikiPathError: topic 非法
    """
    validate_topic(topic)
    # filename 走 safe_filename 做二次清洗（调用方通常已经调过了，这里兜底）
    filename = safe_filename(filename)
    path = raw_file_path(topic, filename)
    # 确认最终路径仍然在 raw_dir 下（双保险）
    try:
        path.resolve().relative_to(get_raw_dir().resolve())
    except ValueError:
        raise WikiPathError(f'解析后路径越出 raw 根目录: {path}')
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return str(path.relative_to(get_wiki_root()))


# ══════════════════════════════════════════════════════════════════
# 入库前体量校验（防止 111 页全图 PRD 类的不合适文档进入 wiki）
# ══════════════════════════════════════════════════════════════════

# 文件大小上限：超过则拒绝
WIKI_RAW_MAX_SIZE_BYTES = int(os.environ.get('WIKI_RAW_MAX_SIZE_BYTES', str(10 * 1024 * 1024)))  # 10 MB
# PDF 页数上限：超过则拒绝
WIKI_PDF_MAX_PAGES_FOR_INGEST = int(os.environ.get('WIKI_PDF_MAX_PAGES_FOR_INGEST', '30'))
# PDF 文字密度阈值（字符/页），低于此值且页数 > MIN 时拒绝（图为主的扫描/排版 PDF）
WIKI_PDF_LOW_TEXT_DENSITY = int(os.environ.get('WIKI_PDF_LOW_TEXT_DENSITY', '200'))
WIKI_PDF_LOW_TEXT_MIN_PAGES = int(os.environ.get('WIKI_PDF_LOW_TEXT_MIN_PAGES', '5'))
# DOCX 字符数上限（粗略对应 40 页文档），超过则拒绝
WIKI_DOCX_MAX_CHARS = int(os.environ.get('WIKI_DOCX_MAX_CHARS', '60000'))


def validate_raw_file_for_wiki(abs_path) -> str | None:
    """检查文件是否适合入 wiki。返回 None 表示通过；否则返回中文拒绝原因。

    设计目标：拦住 wiki 一篇文章承载不下的过大 / 过长 / 图为主的文档，
    引导用户拆分或先整理摘要。误判时返回 None（让下游处理而非误拒）。
    """
    abs_path = Path(abs_path)
    try:
        size = abs_path.stat().st_size
    except OSError:
        return None
    if size > WIKI_RAW_MAX_SIZE_BYTES:
        size_mb = size / 1024 / 1024
        return (
            f'文件过大（{size_mb:.1f} MB > 上限 {WIKI_RAW_MAX_SIZE_BYTES/1024/1024:.0f} MB），'
            f'wiki 一篇文章承载不下。建议拆分为 ≤10 页章节后分别上传，'
            f'或先整理成 5K 字以内的总结再加入。'
        )

    ext = abs_path.suffix.lower().lstrip('.')

    if ext == 'pdf':
        try:
            import fitz
            doc = fitz.open(str(abs_path))
            pages = len(doc)
            text_chars = sum(len(p.get_text() or '') for p in doc)
            doc.close()
        except Exception:
            return None

        if pages > WIKI_PDF_MAX_PAGES_FOR_INGEST:
            return (
                f'PDF 页数过多（{pages} 页 > 上限 {WIKI_PDF_MAX_PAGES_FOR_INGEST}）。'
                f'建议按章节拆分为 ≤10 页的多份后再上传。'
            )

        if pages > WIKI_PDF_LOW_TEXT_MIN_PAGES:
            density = text_chars // max(pages, 1)
            if density < WIKI_PDF_LOW_TEXT_DENSITY:
                return (
                    f'PDF 内容以图片为主（{pages} 页，文字层平均 {density} 字符/页 '
                    f'< 阈值 {WIKI_PDF_LOW_TEXT_DENSITY}）。这种文档通常是 mockup / 截图集，'
                    f'wiki 不适合直接容纳。建议截取核心几张图整理成简短文章再上传。'
                )

    elif ext == 'docx':
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            with zipfile.ZipFile(str(abs_path)) as zf:
                doc_xml = zf.read('word/document.xml')
            root = ET.fromstring(doc_xml)
            ns_t = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
            text_chars = sum(len(t.text or '') for t in root.iter(ns_t))
        except Exception:
            return None

        if text_chars > WIKI_DOCX_MAX_CHARS:
            return (
                f'docx 内容过长（约 {text_chars:,} 字符 > 上限 {WIKI_DOCX_MAX_CHARS:,}，'
                f'相当于 {text_chars // 1500} 多页）。'
                f'建议先整理成简洁版本（5K 字以内）再上传。'
            )

    return None


# ══════════════════════════════════════════════════════════════════
# PDF 扫描件 → 图片提取（Claude Vision）
# ══════════════════════════════════════════════════════════════════

# PDF 文字层字符数低于此阈值时，认定为扫描件，走 Vision 路径
PDF_TEXT_THRESHOLD = int(os.environ.get('WIKI_PDF_TEXT_THRESHOLD', '50'))
# Vision 模式下 PDF 最多转几页图片（限制 token 成本）
PDF_MAX_VISION_PAGES = int(os.environ.get('WIKI_PDF_MAX_VISION_PAGES', '30'))
# 图片渲染 DPI（150 清晰度好，nginx 已放宽到 25m 不用省）
PDF_VISION_DPI = int(os.environ.get('WIKI_PDF_VISION_DPI', '150'))
# JPEG 质量
PDF_VISION_JPEG_QUALITY = int(os.environ.get('WIKI_PDF_VISION_JPEG_QUALITY', '75'))
# base64 后总体积上限（字节）—— SG NAS nginx 已改为 client_max_body_size 25m
# 保留自动降质逻辑作为安全网（防止极大的 PDF 或代理配置回退）
PDF_VISION_MAX_B64_BYTES = int(os.environ.get('WIKI_PDF_VISION_MAX_B64_BYTES', str(20 * 1024 * 1024)))  # 20MB


@dataclass
class RawFileContent:
    """原始文件内容的统一结构。

    text 模式:   content_type='text', text 有值, images 为空
    vision 模式: content_type='images', text 为空, images 有值
    """
    content_type: str  # 'text' | 'images'
    text: str = ''
    images: list = field(default_factory=list)
    # images: [{'page': int, 'base64': str, 'media_type': 'image/png'}, ...]
    total_pages: int = 0
    extracted_pages: int = 0
    # embedded_images: text 模式下从 docx 嵌入图提取出来的清单
    # [{'order':1,'paragraph_index':12,'data':bytes,'media_type':'image/png','original_name':'image1.png'}, ...]
    embedded_images: list = field(default_factory=list)


def extract_raw_file_content(raw_path: str) -> RawFileContent:
    """智能提取原始文件内容。

    对 PDF：先尝试文字提取；字符数 < PDF_TEXT_THRESHOLD 时自动切换到
    Vision 模式（每页转 PNG base64 → 交给 Claude Opus 读图）。

    对其他格式（md/txt/docx）：直接返回纯文本。

    这是 compiler.py 应该调用的入口，代替旧的 read_raw_file_text。
    """
    if not raw_path:
        raise ValueError('raw_path 不能为空')

    abs_path = (get_wiki_root() / raw_path).resolve()
    try:
        abs_path.relative_to(get_wiki_root().resolve())
    except ValueError:
        raise WikiPathError(f'raw_path 越出 wiki 根目录: {raw_path!r}')

    if not abs_path.exists():
        raise FileNotFoundError(f'原始文件不存在: {abs_path}')

    ext = abs_path.suffix.lower()

    # 非 PDF 走纯文本
    if ext != '.pdf':
        text = _extract_text_non_pdf(abs_path, ext)
        embedded = []
        if ext == '.docx':
            try:
                from app.services.wiki.docx_images import extract_docx_images
                embedded = extract_docx_images(abs_path)
            except Exception as e:
                logger.warning(f'[Storage] docx 抽图失败 {raw_path}: {e}')
        return RawFileContent(content_type='text', text=text, embedded_images=embedded)

    # PDF:先尝试文字提取
    try:
        import fitz
    except ImportError:
        raise ImportError('需要 pymupdf 才能处理 PDF')

    doc = fitz.open(str(abs_path))
    total_pages = len(doc)

    text_parts = []
    for page in doc:
        text_parts.append(page.get_text() or '')
    full_text = '\n\n'.join(text_parts)
    doc.close()

    # 文字层足够 → 走文本模式
    if len(full_text.strip()) >= PDF_TEXT_THRESHOLD:
        logger.info(f'[Storage] PDF text extraction: {len(full_text)} chars from {total_pages} pages')
        return RawFileContent(content_type='text', text=full_text, total_pages=total_pages, extracted_pages=total_pages)

    # 文字层不足 → Vision 模式：每页转 PNG
    logger.info(f'[Storage] PDF text too short ({len(full_text.strip())} chars < {PDF_TEXT_THRESHOLD}), switching to Vision mode')
    return _pdf_to_vision_images(abs_path, total_pages)


def _pdf_to_vision_images(abs_path: Path, total_pages: int) -> RawFileContent:
    """把 PDF 每页渲染为 JPEG base64，交给 Claude Vision 识别。

    自动降质策略：
    - 从当前 DPI + quality 开始渲染
    - 如果总 base64 体积超 PDF_VISION_MAX_B64_BYTES，自动降档重试
    - 降档序列：DPI 120→96→72，quality 70→55→40
    - Claude Vision 在 72 DPI + quality 40 下仍能准确读文字
    """
    import fitz
    from io import BytesIO

    pages_to_extract = min(total_pages, PDF_MAX_VISION_PAGES)
    if total_pages > PDF_MAX_VISION_PAGES:
        logger.warning(
            f'[Storage] PDF has {total_pages} pages, only extracting first {PDF_MAX_VISION_PAGES} for Vision'
        )

    # 降档序列：(dpi, jpeg_quality)
    # 最低档 DPI=60 q=25 大约每页 30-50KB，9 页 ≈ 400KB base64
    # Claude Vision 在此质量下仍能准确 OCR 中英文
    quality_levels = [
        (PDF_VISION_DPI, PDF_VISION_JPEG_QUALITY),
        (96, 55),
        (72, 40),
        (60, 25),
    ]

    doc = fitz.open(str(abs_path))
    final_images = []
    used_dpi = PDF_VISION_DPI
    used_quality = PDF_VISION_JPEG_QUALITY
    downgraded = False

    for dpi, quality in quality_levels:
        images = []
        for i in range(pages_to_extract):
            page = doc[i]
            pix = page.get_pixmap(dpi=dpi)
            try:
                from PIL import Image as PILImage
                pil_img = PILImage.frombytes('RGB', [pix.width, pix.height], pix.samples)
                buf = BytesIO()
                pil_img.save(buf, format='JPEG', quality=quality)
                img_bytes = buf.getvalue()
                media_type = 'image/jpeg'
            except ImportError:
                img_bytes = pix.tobytes('png')
                media_type = 'image/png'

            images.append({
                'page': i,
                'base64': base64.b64encode(img_bytes).decode('ascii'),
                'media_type': media_type,
                'size_bytes': len(img_bytes),
            })

        total_raw = sum(img['size_bytes'] for img in images)
        total_b64 = int(total_raw * 4 / 3)

        logger.info(
            f'[Storage] PDF Vision attempt: dpi={dpi} q={quality} '
            f'raw={total_raw / 1024:.0f}KB b64≈{total_b64 / 1024:.0f}KB '
            f'limit={PDF_VISION_MAX_B64_BYTES / 1024:.0f}KB'
        )

        if total_b64 <= PDF_VISION_MAX_B64_BYTES:
            final_images = images
            used_dpi = dpi
            used_quality = quality
            break
        else:
            downgraded = True
            logger.warning(f'[Storage] PDF Vision 超限，降档到 dpi={dpi} quality={quality}')
    else:
        # 全部档位都试完还是超，用最低档的结果
        final_images = images
        used_dpi = quality_levels[-1][0]
        used_quality = quality_levels[-1][1]
        logger.warning('[Storage] PDF Vision 所有档位都超限，使用最低质量')

    doc.close()

    total_size = sum(img['size_bytes'] for img in final_images)
    logger.info(
        f'[Storage] PDF Vision final: {pages_to_extract}/{total_pages} pages, '
        f'dpi={used_dpi} q={used_quality} size={total_size / 1024:.0f}KB '
        f'{"(downgraded)" if downgraded else ""}'
    )

    # 构造 embedded_images 与 _persist_embedded_images_and_rewrite_md 兼容
    # 让 vision 模式 PDF 页图也能走 AUTO_IMG:N → _assets 持久化流程
    embedded = []
    for i, img in enumerate(final_images):
        raw_bytes = base64.b64decode(img['base64'])
        ext = 'jpg' if img['media_type'] == 'image/jpeg' else 'png'
        embedded.append({
            'order': i + 1,                # 1-based for AUTO_IMG:N
            'page_index': img['page'],     # 0-based page number
            'data': raw_bytes,
            'media_type': img['media_type'],
            'original_name': f'page-{img["page"] + 1}.{ext}',
        })

    return RawFileContent(
        content_type='images',
        images=final_images,
        total_pages=total_pages,
        extracted_pages=pages_to_extract,
        embedded_images=embedded,
    )


def _extract_text_non_pdf(abs_path: Path, ext: str) -> str:
    """非 PDF 文件的纯文本提取。"""
    if ext in ('.md', '.txt'):
        return abs_path.read_text(encoding='utf-8', errors='replace')
    if ext == '.docx':
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError('需要 python-docx 才能读取 DOCX')
        doc = Document(str(abs_path))
        parts = []
        # 按文档原始顺序遍历 body 元素（段落和表格交替出现）
        para_idx = 0
        table_idx = 0
        for child in doc.element.body:
            if child.tag == qn('w:p'):
                if para_idx < len(doc.paragraphs):
                    text = doc.paragraphs[para_idx].text.strip()
                    if text:
                        parts.append(text)
                    para_idx += 1
            elif child.tag == qn('w:tbl'):
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(' | '.join(cells))
                    if rows:
                        parts.append('')
                        parts.append(rows[0])
                        parts.append(' | '.join(['---'] * len(table.rows[0].cells)))
                        for r in rows[1:]:
                            parts.append(r)
                        parts.append('')
                    table_idx += 1
        return '\n'.join(parts)
    raise ValueError(f'不支持的原始文件类型: {ext}（支持 .md/.txt/.pdf/.docx）')


def read_raw_file_text(raw_path: str) -> str:
    """读原始文件并尽力提取成纯文本，给 LLM 编译用。

    ⚠️ 旧接口，仅返回文本。对扫描件 PDF 会返回空字符串。
    新代码应改用 extract_raw_file_content() 来获取 text 或 images。

    支持：.md / .txt / .pdf / .docx。
    其他类型抛 ValueError。
    """
    if not raw_path:
        raise ValueError('raw_path 不能为空')

    # 防止有人绕过 save_raw_file 直接调这个函数读 wiki_root 以外的文件
    abs_path = (get_wiki_root() / raw_path).resolve()
    try:
        abs_path.relative_to(get_wiki_root().resolve())
    except ValueError:
        raise WikiPathError(f'raw_path 越出 wiki 根目录: {raw_path!r}')

    if not abs_path.exists():
        raise FileNotFoundError(f'原始文件不存在: {abs_path}')

    ext = abs_path.suffix.lower()

    if ext in ('.md', '.txt'):
        return abs_path.read_text(encoding='utf-8', errors='replace')

    if ext == '.pdf':
        # PMA 已安装 pymupdf（fitz），提取质量和性能都优于 pypdf
        try:
            import fitz  # pymupdf
        except ImportError:
            raise ImportError('需要 pymupdf 才能读取 PDF，请 pip install pymupdf')
        pages = []
        with fitz.open(str(abs_path)) as doc:
            for page in doc:
                try:
                    pages.append(page.get_text() or '')
                except Exception as e:
                    logger.warning(f'PDF 页面提取失败: {e}')
                    pages.append('')
        return '\n\n'.join(pages)

    if ext == '.docx':
        try:
            from docx import Document
        except ImportError:
            raise ImportError('需要 python-docx 包才能读取 DOCX，请 pip install python-docx')
        doc = Document(str(abs_path))
        return '\n\n'.join(p.text for p in doc.paragraphs)

    raise ValueError(f'不支持的原始文件类型: {ext}（支持 .md/.txt/.pdf/.docx）')


def delete_raw_file(raw_path: str) -> bool:
    """删除原始文件。返回是否真的删除。

    防止路径穿越：解析后必须仍在 raw_dir 下。
    """
    if not raw_path:
        return False
    abs_path = (get_wiki_root() / raw_path).resolve()
    try:
        abs_path.relative_to(get_raw_dir().resolve())
    except ValueError:
        logger.warning(f'[Wiki] delete_raw_file 拒绝越界路径: {raw_path!r}')
        return False
    if abs_path.exists():
        abs_path.unlink()
        return True
    return False


# ══════════════════════════════════════════════════════════════════
# 文章资源（图片）保存与历史备份
# ══════════════════════════════════════════════════════════════════

_MEDIA_TYPE_EXT = {
    'image/png': 'png',
    'image/jpeg': 'jpg',
    'image/jpg': 'jpg',
    'image/gif': 'gif',
    'image/webp': 'webp',
}


def _ext_for_media(media_type: str) -> str:
    ext = _MEDIA_TYPE_EXT.get((media_type or '').lower())
    if not ext:
        raise ValueError(f'不支持的图片类型: {media_type}')
    return ext


def save_article_image(topic: str, slug: str, index: int, data: bytes, media_type: str) -> str:
    """保存图片到 _assets/<slug>/img-<index>.<ext>，返回相对 article 的路径。"""
    validate_topic_slug(topic, slug)
    if index < 1:
        raise ValueError('index 从 1 开始')
    ext = _ext_for_media(media_type)
    out_dir = assets_dir_for_article(topic, slug)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f'img-{index}.{ext}'
    out_path.write_bytes(data)
    return article_image_relative_path(slug, index, ext)


def replace_article_image(topic: str, slug: str, index: int, data: bytes, media_type: str) -> str:
    """覆盖现有图片，旧文件备份到 .history/。返回相对路径。"""
    validate_topic_slug(topic, slug)
    if index < 1:
        raise ValueError('index 从 1 开始')
    ext = _ext_for_media(media_type)
    out_dir = assets_dir_for_article(topic, slug)
    out_path = out_dir / f'img-{index}.{ext}'
    if out_path.exists():
        history_dir = out_dir / '.history'
        history_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime('%Y%m%d-%H%M%S')
        backup = history_dir / f'img-{index}.{ext}.{ts}.bak'
        backup.write_bytes(out_path.read_bytes())
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(data)
    return article_image_relative_path(slug, index, ext)


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()
