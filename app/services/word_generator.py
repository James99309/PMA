#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档模板生成服务
用于使用Word/Excel模板生成批价单、结算单和报价单

策略：
- Word: 使用 python-docx 直接操作 Word 文档，替换文本并动态添加表格行
- Excel: 使用 openpyxl 操作 Excel 文档，替换单元格并动态添加行
"""

import os
import subprocess
import platform
import tempfile
from datetime import datetime, timedelta
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import logging
import openpyxl
from openpyxl.styles import Alignment, Font, Border, Side
from config import Config

logger = logging.getLogger(__name__)


class WordGenerator:
    """Word文档生成器"""

    def __init__(self):
        self.template_dir = os.path.join(os.path.dirname(__file__), '..', 'templates', 'word')

    def _get_template_path(self, template_name):
        """获取模板文件路径"""
        return os.path.join(self.template_dir, template_name)

    def _get_currency_symbol(self, currency=None):
        """获取货币符号"""
        from app.utils.dictionary_helpers import get_currency_symbol
        return get_currency_symbol(currency or Config.DEFAULT_CURRENCY)

    def _format_discount_rate(self, rate):
        """格式化折扣率为百分比，保留2位小数"""
        return f"{rate * 100:.2f}%"

    def _get_project_type_label(self, project_type):
        """获取项目类型标签"""
        type_map = {
            'channel_follow': '渠道跟进',
            'sales_key': '销售重点',
            'sales_opportunity': '销售机会',
        }
        return type_map.get(project_type, project_type or '未知')

    def _get_status_label(self, status):
        """获取状态标签"""
        status_map = {
            'draft': '草稿',
            'pending': '审批中',
            'approved': '已批准',
            'rejected': '已拒绝',
        }
        return status_map.get(status, status or '未知')

    def _replace_text_in_cell(self, cell, old_text, new_text):
        """替换单元格中的文本，保持格式"""
        for paragraph in cell.paragraphs:
            if old_text in paragraph.text:
                # 保持原有格式，只替换文本
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

    def _replace_text_in_table(self, table, replacements):
        """替换表格中的所有文本"""
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    if old_text in cell.text:
                        self._replace_text_in_cell(cell, old_text, new_text)

    def _copy_row(self, table, row_idx):
        """复制表格行"""
        row = table.rows[row_idx]
        tr = row._tr
        new_tr = deepcopy(tr)
        table._tbl.append(new_tr)
        return table.rows[-1]

    def _set_cell_text(self, cell, text, font_size=9, align='center'):
        """设置单元格文本

        Args:
            cell: 单元格对象
            text: 文本内容
            font_size: 字体大小，默认9
            align: 对齐方式，'left'/'center'/'right'，默认'center'
        """
        cell.text = str(text)
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
            for run in paragraph.runs:
                run.font.size = Pt(font_size)

    def generate_pricing_order_word(self, pricing_order, include_notes=False):
        """
        生成批价单Word文档

        Args:
            pricing_order: PricingOrder 对象

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        try:
            template_path = self._get_template_path('pricing_order_template.docx')
            doc = Document(template_path)

            # 准备基本数据
            currency = pricing_order.currency or Config.DEFAULT_CURRENCY
            currency_symbol = self._get_currency_symbol(currency)

            # 计算有效期（默认审批日期起60天）
            valid_from = pricing_order.approved_at or pricing_order.created_at
            valid_to = valid_from + timedelta(days=60) if valid_from else None

            # 计算市场价总计
            market_total = sum(d.market_price * d.quantity for d in pricing_order.pricing_details)

            # 表格0: 基本信息
            table0 = doc.tables[0]
            replacements0 = {
                'PO202512-003': pricing_order.order_number,
                '2025年12月18日': (pricing_order.approved_at or pricing_order.created_at).strftime('%Y年%m月%d日') if (pricing_order.approved_at or pricing_order.created_at) else '',
                '上海东方枢纽': pricing_order.project.project_name if pricing_order.project else '',
                '渠道跟进': self._get_project_type_label(pricing_order.approval_flow_type),
                '已批准': self._get_status_label(pricing_order.status),
                '45%': self._format_discount_rate(pricing_order.pricing_total_discount_rate or 1.0),
            }
            self._replace_text_in_table(table0, replacements0)

            # 表格1: 有效期
            table1 = doc.tables[1]
            valid_from_str = valid_from.strftime('%Y年%m月%d日') if valid_from else ''
            valid_to_str = valid_to.strftime('%Y年%m月%d日') if valid_to else ''
            replacements1 = {
                '2025年12月18日 至 2026年 2月18日': f'{valid_from_str} 至 {valid_to_str}',
                '【上海东方枢纽】': f'【{pricing_order.project.project_name if pricing_order.project else ""}】',
            }
            self._replace_text_in_table(table1, replacements1)

            # 表格2: 渠道信息
            table2 = doc.tables[2]
            dealer_name = pricing_order.dealer.company_name if pricing_order.dealer else '（厂商直签）'
            distributor_name = pricing_order.distributor.company_name if pricing_order.distributor else '（无分销商）'
            project_manager = pricing_order.project.owner.real_name if pricing_order.project and pricing_order.project.owner else ''
            sales_manager = ''
            if pricing_order.project and hasattr(pricing_order.project, 'vendor_sales_manager') and pricing_order.project.vendor_sales_manager:
                sales_manager = pricing_order.project.vendor_sales_manager.real_name

            replacements2 = {
                '上海瑞康通信科技有限公司': dealer_name,
                '上海淳泊信息科技有限公司': distributor_name,
            }
            self._replace_text_in_table(table2, replacements2)

            # 表格3: 明细表格 - 动态处理
            detail_table = doc.tables[3]

            # 删除示例数据行（保留表头）
            while len(detail_table.rows) > 1:
                tr = detail_table.rows[-1]._tr
                detail_table._tbl.remove(tr)

            # 添加实际明细数据
            for idx, detail in enumerate(pricing_order.pricing_details, 1):
                row = detail_table.add_row()
                cells = row.cells

                # 设置各列数据
                self._set_cell_text(cells[0], idx)  # 序号
                self._set_cell_text(cells[1], detail.product_name or '')  # 产品名称
                self._set_cell_text(cells[2], detail.product_model or '')  # 型号
                # 规格参数（截断长文本）
                desc = detail.product_desc or ''
                if len(desc) > 40:
                    desc = desc[:40] + '...'
                self._set_cell_text(cells[3], desc)
                self._set_cell_text(cells[4], detail.brand or '和源')  # 品牌
                self._set_cell_text(cells[5], detail.quantity)  # 数量
                self._set_cell_text(cells[6], f'{currency_symbol}{detail.market_price:,.2f}')  # 零售单价
                self._set_cell_text(cells[7], self._format_discount_rate(detail.discount_rate))  # 折扣率
                self._set_cell_text(cells[8], f'{currency_symbol}{detail.unit_price:,.2f}')  # 批准单价
                self._set_cell_text(cells[9], f'{currency_symbol}{detail.total_price:,.2f}')  # 总价（数据库字段）
                self._set_cell_text(cells[10], detail.product_mn or '')  # 产品编码

            # 表格4: 总金额
            table4 = doc.tables[4]
            replacements4 = {
                '¥ 63,196.20': f'{currency_symbol} {pricing_order.pricing_total_amount:,.2f}',
            }
            self._replace_text_in_table(table4, replacements4)

            # 备注（可选）：在文档末尾追加备注段落
            if include_notes and pricing_order.notes:
                doc.add_paragraph()
                notes_para = doc.add_paragraph()
                notes_para.add_run('备注：').bold = True
                notes_para.add_run(pricing_order.notes)

            # 保存到内存
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            # 生成文件名
            project_name = pricing_order.project.project_name if pricing_order.project else "未知项目"
            safe_project_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_', '（', '）', '【', '】')).rstrip()
            filename = f"{pricing_order.order_number} & {safe_project_name}.docx"

            logger.info(f"成功生成批价单Word文档: {filename}")

            return {
                'content': output.getvalue(),
                'filename': filename
            }

        except Exception as e:
            logger.error(f"生成批价单Word文档失败: {str(e)}")
            raise

    def generate_pricing_order_word_v2(self, pricing_order, include_notes=False):
        """生成批价单Word文档（优化版样式，从零构建，不依赖 .docx 模板）"""
        from docx.shared import RGBColor
        from docx.enum.table import WD_ALIGN_VERTICAL

        # ── 颜色常量 ────────────────────────────────────────────────────
        C_DARK_BLUE  = RGBColor(0x1F, 0x3A, 0x5F)
        C_GOLD       = RGBColor(0xB8, 0x95, 0x3F)
        C_GRAY       = RGBColor(0x59, 0x59, 0x59)
        C_TEXT       = RGBColor(0x1A, 0x1A, 0x1A)
        C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
        C_GREEN      = RGBColor(0x1F, 0x7A, 0x4A)
        C_LIGHT_BLUE = RGBColor(0xA8, 0xBE, 0xD8)

        # ── 辅助函数 ─────────────────────────────────────────────────────

        def set_cell_bg(cell, hex_color):
            tcPr = cell._tc.get_or_add_tcPr()
            for s in tcPr.findall(qn('w:shd')):
                tcPr.remove(s)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        def set_cell_w(cell, dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            for w in tcPr.findall(qn('w:tcW')):
                tcPr.remove(w)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(dxa))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        def set_cell_valign(cell, val='center'):
            tcPr = cell._tc.get_or_add_tcPr()
            for v in tcPr.findall(qn('w:vAlign')):
                tcPr.remove(v)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), val)
            tcPr.append(vAlign)

        def set_spacing(para, before_pt=0, after_pt=0):
            pPr = para._p.get_or_add_pPr()
            for s in pPr.findall(qn('w:spacing')):
                pPr.remove(s)
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:before'), str(int(before_pt * 20)))
            sp.set(qn('w:after'), str(int(after_pt * 20)))
            sp.set(qn('w:line'), '276')
            sp.set(qn('w:lineRule'), 'auto')
            pPr.append(sp)

        FONT_NAME  = 'Microsoft YaHei'
        LATIN_FONT = 'Arial'   # 拉丁字符专用（LibreOffice 可找到，Word 与 YaHei 拉丁效果一致）
        MONEY_FONT = 'Arial'   # 金额数字专用字体

        def set_run_font(run, fn=FONT_NAME):
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            if fn == FONT_NAME:
                # 普通文字: 拉丁用 Arial（LibreOffice 可渲染），CJK 用微软雅黑
                rFonts.set(qn('w:ascii'),    LATIN_FONT)
                rFonts.set(qn('w:hAnsi'),    LATIN_FONT)
                rFonts.set(qn('w:eastAsia'), FONT_NAME)
                rFonts.set(qn('w:cs'),       FONT_NAME)
            else:
                # 金额/英文专用字体: 全部使用指定字体
                for attr in ('w:ascii', 'w:eastAsia', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), fn)

        def add_run(para, text, size_pt, bold=False, color=None, fn=FONT_NAME):
            run = para.add_run(str(text))
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            set_run_font(run, fn)
            return run

        def set_tbl_borders(table, color='CCCCCC', sz=4):
            tblPr = table._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                table._tbl.insert(0, tblPr)
            borders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), str(sz))
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), color)
                borders.append(b)
            existing = tblPr.find(qn('w:tblBorders'))
            if existing is not None:
                tblPr.remove(existing)
            tblPr.append(borders)
            # 移除 tblLook（禁止 Word 自动叠加条纹底色）
            look = tblPr.find(qn('w:tblLook'))
            if look is not None:
                tblPr.remove(look)

        def set_tbl_width(table, dxa):
            tblPr = table._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                table._tbl.insert(0, tblPr)
            for w in tblPr.findall(qn('w:tblW')):
                tblPr.remove(w)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(dxa))
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)

        def amount_to_chinese(amount):
            digits = '零壹贰叁肆伍陆柒捌玖'
            amount = round(float(amount or 0), 2)
            int_part = int(amount)
            dec = round((amount - int_part) * 100)
            jiao, fen = dec // 10, dec % 10

            def _int_cn(n):
                if n == 0:
                    return '零'
                wan_u = ['', '万', '亿']
                unit_u = ['', '拾', '佰', '仟']
                groups = []
                while n > 0:
                    groups.append(n % 10000)
                    n //= 10000
                result = ''
                for gi in range(len(groups) - 1, -1, -1):
                    g = groups[gi]
                    if g == 0:
                        if result and result[-1] != '零':
                            result += '零'
                        continue
                    g_str, need_zero = '', False
                    for ui in range(3, -1, -1):
                        d = (g // (10 ** ui)) % 10
                        if d:
                            if need_zero:
                                g_str += '零'
                                need_zero = False
                            g_str += digits[d] + unit_u[ui]
                        elif g_str:
                            need_zero = True
                    result += g_str + wan_u[gi]
                return result

            int_str = _int_cn(int_part) + '元'
            if jiao == 0 and fen == 0:
                dec_str = '整'
            elif jiao == 0:
                dec_str = '零' + digits[fen] + '分'
            elif fen == 0:
                dec_str = digits[jiao] + '角'
            else:
                dec_str = digits[jiao] + '角' + digits[fen] + '分'
            return '人民币 ' + int_str + dec_str

        def add_section_heading(doc, cn_text, en_text):
            p = doc.add_paragraph()
            set_spacing(p, before_pt=8, after_pt=3)
            add_run(p, '■ ', 11, color=C_GOLD)
            add_run(p, cn_text, 12, bold=True, color=C_DARK_BLUE)
            add_run(p, f'   {en_text}', 8, color=C_GRAY)
            return p

        # ── 数据准备 ─────────────────────────────────────────────────────
        try:
            import os as _os
            is_ovs    = Config.IS_OVS or _os.environ.get('WORD_RENDER_LOCALE') == 'ovs'
            date_fmt  = '%d %b %Y' if is_ovs else '%Y年%m月%d日'

            currency      = pricing_order.currency or Config.DEFAULT_CURRENCY
            curr_sym      = self._get_currency_symbol(currency)
            project       = pricing_order.project
            project_name  = project.project_name if project else ''

            valid_from    = pricing_order.approved_at or pricing_order.created_at
            valid_to      = valid_from + timedelta(days=60) if valid_from else None
            vf_str        = valid_from.strftime(date_fmt) if valid_from else ''
            vt_str        = valid_to.strftime(date_fmt)   if valid_to   else ''
            valid_days    = (valid_to - valid_from).days   if (valid_from and valid_to) else 0
            if is_ovs:
                valid_str = f'{vf_str}  to  {vt_str}    ({valid_days} days)' if vf_str else ''
                scope_str = (f'This pricing applies exclusively to the [{project_name}] project. '
                             f'Dealers must quote this order number when placing orders through designated channels.') if project_name else ''
            else:
                valid_str = f'{vf_str}  至  {vt_str}    (共 {valid_days} 天)' if vf_str else ''
                scope_str = f'本批价仅适用于【{project_name}】项目，经销商通过指定分销渠道下单时参照执行。' if project_name else ''

            dealer_name   = pricing_order.dealer.company_name      if pricing_order.dealer      else ('(Direct)'  if is_ovs else '（厂商直签）')
            distrib_name  = pricing_order.distributor.company_name if pricing_order.distributor else ('(None)'    if is_ovs else '（无分销商）')
            proj_mgr      = (project.owner.real_name if project and project.owner else '') or ''
            sales_mgr     = ''
            if project and getattr(project, 'vendor_sales_manager', None):
                sales_mgr = project.vendor_sales_manager.real_name or ''

            if is_ovs:
                type_map = {'channel_follow': 'Channel Follow', 'sales_key': 'Key Account',
                            'sales_opportunity': 'Sales Opportunity', 'sales_focus': 'Focus Sales'}
                _status_en = {'draft': 'Draft', 'pending': 'Pending', 'approved': 'Approved', 'rejected': 'Rejected'}
                status_label = _status_en.get(pricing_order.status, pricing_order.status or 'Unknown')
            else:
                type_map = {'channel_follow': '渠道跟进', 'sales_key': '销售重点',
                            'sales_opportunity': '销售机会', 'sales_focus': '重点销售'}
                status_label = self._get_status_label(pricing_order.status)
            flow_label    = type_map.get(pricing_order.approval_flow_type, pricing_order.approval_flow_type or '')
            status_color  = C_GREEN if pricing_order.status == 'approved' else C_DARK_BLUE
            status_prefix = '✔ ' if pricing_order.status == 'approved' else ''

            total_rate    = pricing_order.pricing_total_discount_rate or 1.0
            discount_str  = f'{int(round(total_rate * 100))} %'
            total_amount  = pricing_order.pricing_total_amount or 0
            total_qty     = sum(d.quantity for d in pricing_order.pricing_details)

            details       = list(pricing_order.pricing_details)
            has_item_note = include_notes and any(getattr(d, 'item_note', None) for d in details)
            has_notes     = include_notes and bool(pricing_order.notes and str(pricing_order.notes).strip())

            TABLE_W = 15038  # A4横向 29.7cm - 左右边距各1.6cm ≈ 26.5cm → 15038 dxa

            # ── 创建文档 ─────────────────────────────────────────────────
            doc = Document()
            section = doc.sections[0]
            section.page_width    = Cm(29.7)
            section.page_height   = Cm(21.0)
            section.left_margin   = Cm(1.59)
            section.right_margin  = Cm(1.59)
            section.top_margin    = Cm(1.59)
            section.bottom_margin = Cm(1.59)

            # 默认段落间距清零
            doc.styles['Normal'].paragraph_format.space_before = Pt(0)
            doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

            # 设置文档级默认字体（Microsoft YaHei），匹配原版模板
            styles_el = doc.styles.element
            doc_defaults = styles_el.find(qn('w:docDefaults'))
            if doc_defaults is None:
                doc_defaults = OxmlElement('w:docDefaults')
                styles_el.insert(0, doc_defaults)
            rPr_default = doc_defaults.find(qn('w:rPrDefault'))
            if rPr_default is None:
                rPr_default = OxmlElement('w:rPrDefault')
                doc_defaults.append(rPr_default)
            rPr_def = rPr_default.find(qn('w:rPr'))
            if rPr_def is None:
                rPr_def = OxmlElement('w:rPr')
                rPr_default.append(rPr_def)
            rFonts_def = rPr_def.find(qn('w:rFonts'))
            if rFonts_def is None:
                rFonts_def = OxmlElement('w:rFonts')
                rPr_def.insert(0, rFonts_def)
            rFonts_def.set(qn('w:ascii'),    LATIN_FONT)   # 拉丁字符用 Arial
            rFonts_def.set(qn('w:hAnsi'),    LATIN_FONT)
            rFonts_def.set(qn('w:eastAsia'), FONT_NAME)    # CJK 用微软雅黑
            rFonts_def.set(qn('w:cs'),       FONT_NAME)
            for tag, val in (('w:sz', '20'), ('w:szCs', '20')):
                el = rPr_def.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr_def.append(el)
                el.set(qn('w:val'), val)

            # ── 标题区 ───────────────────────────────────────────────────
            p0 = doc.add_paragraph()
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p0, 0, 2)
            add_run(p0, 'APPROVED  PRICING  SHEET' if is_ovs else '批  价  确  认  单', 22, bold=True, color=C_DARK_BLUE)

            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p1, 0, 2)
            add_run(p1, '' if is_ovs else 'APPROVED PRICING SHEET', 10, color=C_GRAY)

            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p2, 0, 6)
            add_run(p2, '◆  ◆  ◆', 8, color=C_GOLD)

            # ── Table 0: 基本信息（4行×6列）───────────────────────────────
            t0 = doc.add_table(rows=4, cols=6)
            set_tbl_borders(t0, color='CCCCCC', sz=4)
            set_tbl_width(t0, TABLE_W)

            LBL_W, VAL_W = 1400, 3613

            def fill_info_row(row_idx, triples):
                """triples: [(label, value, val_color, val_bold), ...]"""
                row = t0.rows[row_idx]
                for i, (lbl, val, vc, vb) in enumerate(triples):
                    lc = row.cells[i * 2]
                    vc_ = row.cells[i * 2 + 1]
                    set_cell_bg(lc, 'E8EEF5')
                    set_cell_w(lc, LBL_W)
                    set_cell_w(vc_, VAL_W)
                    set_cell_valign(lc)
                    set_cell_valign(vc_)
                    lp = lc.paragraphs[0]; set_spacing(lp, 2, 2)
                    add_run(lp, lbl, 9, bold=True, color=C_DARK_BLUE)
                    vp = vc_.paragraphs[0]; set_spacing(vp, 2, 2)
                    add_run(vp, val, 10, bold=vb, color=vc)

            created_date = (pricing_order.approved_at or pricing_order.created_at)
            date_str = created_date.strftime(date_fmt) if created_date else ''

            if is_ovs:
                _L0 = [('Order No.', pricing_order.order_number, C_GOLD, True),
                        ('Date',     date_str,                    C_TEXT, False),
                        ('Status',   status_prefix + status_label, status_color, True)]
                _L1 = [('Project',    project_name, C_TEXT, True),
                        ('Type',      flow_label,   C_TEXT, False),
                        ('Disc. Rate', discount_str, C_GOLD, True)]
                _lbl_valid = 'Valid Period'
                _lbl_scope = 'Scope'
            else:
                _L0 = [('批价单号', pricing_order.order_number, C_GOLD, True),
                        ('批价日期', date_str,                   C_TEXT, False),
                        ('批价状态', status_prefix + status_label, status_color, True)]
                _L1 = [('项目名称', project_name, C_TEXT, True),
                        ('项目类型', flow_label,   C_TEXT, False),
                        ('总折扣率', discount_str, C_GOLD, True)]
                _lbl_valid = '价格有效期'
                _lbl_scope = '适用范围'

            fill_info_row(0, _L0)
            fill_info_row(1, _L1)

            # 有效期行（合并 col 1-5）
            lc2 = t0.rows[2].cells[0]
            set_cell_bg(lc2, 'E8EEF5'); set_cell_w(lc2, LBL_W); set_cell_valign(lc2)
            lp2 = lc2.paragraphs[0]; set_spacing(lp2, 2, 2)
            add_run(lp2, _lbl_valid, 9, bold=True, color=C_DARK_BLUE)
            vc2 = t0.cell(2, 1).merge(t0.cell(2, 5))
            set_cell_valign(vc2)
            vp2 = vc2.paragraphs[0]; set_spacing(vp2, 2, 2)
            add_run(vp2, valid_str, 10, color=C_TEXT)

            # 适用范围行（合并 col 1-5）
            lc3 = t0.rows[3].cells[0]
            set_cell_bg(lc3, 'E8EEF5'); set_cell_w(lc3, LBL_W); set_cell_valign(lc3)
            lp3 = lc3.paragraphs[0]; set_spacing(lp3, 2, 2)
            add_run(lp3, _lbl_scope, 9, bold=True, color=C_DARK_BLUE)
            vc3 = t0.cell(3, 1).merge(t0.cell(3, 5))
            set_cell_valign(vc3)
            vp3 = vc3.paragraphs[0]; set_spacing(vp3, 2, 2)
            add_run(vp3, scope_str, 10, color=C_TEXT)

            sp1 = doc.add_paragraph(); set_spacing(sp1, 0, 0)

            # ── Table 1: 供应商 + 渠道信息（1行×2列）──────────────────────
            t1 = doc.add_table(rows=1, cols=2)
            set_tbl_borders(t1, color='DDDDDD', sz=4)
            set_tbl_width(t1, TABLE_W)
            half = TABLE_W // 2

            def fill_info_card(cell, title_cn, title_en, fields):
                set_cell_bg(cell, 'FBFCFD')
                set_cell_valign(cell, 'top')
                set_cell_w(cell, half)
                # 标题
                p = cell.paragraphs[0]; set_spacing(p, 3, 3)
                add_run(p, title_cn, 11, bold=True, color=C_DARK_BLUE)
                add_run(p, f'  {title_en}', 8, color=C_GRAY)
                # 字段
                for lbl, val in fields:
                    fp = cell.add_paragraph(); set_spacing(fp, 1, 1)
                    add_run(fp, lbl, 9, color=C_GRAY)
                    add_run(fp, val, 10, color=C_TEXT)

            if is_ovs:
                fill_info_card(t1.rows[0].cells[0], 'SUPPLIER', '', [
                    ('Company :  ', 'Evertac Solutions Singaproe Pte Ltd.'),
                    ('Website :  ', 'http://www.evertacsolutions.com'),
                ])
                fill_info_card(t1.rows[0].cells[1], 'CHANNEL', '', [
                    ('Dealer :  ',      dealer_name),
                    ('Distributor :  ', distrib_name),
                    ('Project Mgr :  ', proj_mgr),
                    ('Sales Mgr :  ',   sales_mgr),
                ])
            else:
                fill_info_card(t1.rows[0].cells[0], '供应商信息', 'SUPPLIER', [
                    ('企业名称 :  ', '和源通信(上海)股份有限公司'),
                    ('办公地址 :  ', '上海市普陀区武威路88号19楼6楼'),
                    ('联系电话 :  ', '021-62596028'),
                    ('官方网址 :  ', 'http://www.evertac.net'),
                ])
                fill_info_card(t1.rows[0].cells[1], '渠道信息', 'CHANNEL', [
                    ('经 销 商 :  ', dealer_name),
                    ('分 销 商 :  ', distrib_name),
                    ('项目负责人 :  ', proj_mgr),
                    ('销售负责人 :  ', sales_mgr),
                ])

            # ── 明细表格 ─────────────────────────────────────────────────
            if is_ovs:
                add_section_heading(doc, 'APPROVED PRICE DETAILS', '')
            else:
                add_section_heading(doc, '批准价格明细', 'APPROVED PRICE DETAILS')

            # 列宽：有行备注时缩减型号和小计列
            if has_item_note:
                col_w = [540, 1500, 1200, 2600, 700, 1100, 620, 1380, 800, 1400, 1798, 1400]
                if is_ovs:
                    headers = ['No.', 'Product Name', 'Model', 'Specifications', 'Brand', 'Code',
                               'Qty', 'List Price', 'Disc.', 'Unit Price', 'Amount', 'Remarks']
                else:
                    headers = ['序号', '产品名称', '型号', '规格参数', '品牌', '产品编码',
                               '数量', '零售单价', '折扣率', '批准单价', '小计金额', '备注']
            else:
                col_w = [540, 1500, 1500, 3200, 700, 1100, 620, 1380, 800, 1400, 2298]
                if is_ovs:
                    headers = ['No.', 'Product Name', 'Model', 'Specifications', 'Brand', 'Code',
                               'Qty', 'List Price', 'Disc.', 'Unit Price', 'Amount']
                else:
                    headers = ['序号', '产品名称', '型号', '规格参数', '品牌', '产品编码',
                               '数量', '零售单价', '折扣率', '批准单价', '小计金额']
            n_cols = len(headers)

            t2 = doc.add_table(rows=1 + len(details) + 1, cols=n_cols)
            set_tbl_borders(t2, color='CCCCCC', sz=4)
            set_tbl_width(t2, TABLE_W)

            # 表头
            hdr_aligns = [WD_ALIGN_PARAGRAPH.CENTER] * n_cols
            hdr_aligns[3] = WD_ALIGN_PARAGRAPH.LEFT  # 规格参数左对齐
            for ci, (hdr, w) in enumerate(zip(headers, col_w)):
                cell = t2.rows[0].cells[ci]
                set_cell_bg(cell, '1F3A5F')
                set_cell_w(cell, w)
                set_cell_valign(cell)
                p = cell.paragraphs[0]
                p.alignment = hdr_aligns[ci]
                set_spacing(p, 2, 2)
                add_run(p, hdr, 9, bold=True, color=C_WHITE)

            # 数据行对齐方式
            data_aligns = [
                WD_ALIGN_PARAGRAPH.CENTER,  # 序号
                WD_ALIGN_PARAGRAPH.LEFT,    # 产品名称
                WD_ALIGN_PARAGRAPH.LEFT,    # 型号
                WD_ALIGN_PARAGRAPH.LEFT,    # 规格参数
                WD_ALIGN_PARAGRAPH.CENTER,  # 品牌
                WD_ALIGN_PARAGRAPH.CENTER,  # 产品编码
                WD_ALIGN_PARAGRAPH.CENTER,  # 数量
                WD_ALIGN_PARAGRAPH.RIGHT,   # 零售单价
                WD_ALIGN_PARAGRAPH.CENTER,  # 折扣率
                WD_ALIGN_PARAGRAPH.RIGHT,   # 批准单价
                WD_ALIGN_PARAGRAPH.RIGHT,   # 小计金额
            ]
            if has_item_note:
                data_aligns.append(WD_ALIGN_PARAGRAPH.LEFT)  # 备注

            for di, detail in enumerate(details):
                row_bg = 'F7F9FC' if di % 2 == 1 else None
                disc_pct = f'{int(round(detail.discount_rate * 100))} %'
                values = [
                    (str(di + 1),                                              9,  False, C_TEXT,       FONT_NAME),
                    (detail.product_name or '',                                9,  True,  C_TEXT,       FONT_NAME),
                    (detail.product_model or '',                               9,  False, C_TEXT,       FONT_NAME),
                    (detail.product_desc or '',                                8,  False, C_GRAY,       FONT_NAME),
                    (detail.brand or '',                                       9,  False, C_TEXT,       FONT_NAME),
                    (getattr(detail, 'product_mn', '') or '',                  8,  False, C_TEXT,       FONT_NAME),
                    (str(detail.quantity),                                     9,  True,  C_TEXT,       FONT_NAME),
                    (f'{curr_sym}{detail.market_price:,.2f}',                  9,  False, C_GRAY,       MONEY_FONT),
                    (disc_pct,                                                 9,  True,  C_GOLD,       FONT_NAME),
                    (f'{curr_sym}{detail.unit_price:,.2f}',                    9,  True,  C_TEXT,       MONEY_FONT),
                    (f'{curr_sym}{detail.total_price:,.2f}',                   9,  True,  C_DARK_BLUE,  MONEY_FONT),
                ]
                if has_item_note:
                    values.append((getattr(detail, 'item_note', '') or '', 8, False, C_GRAY, FONT_NAME))

                row = t2.rows[1 + di]
                for ci, (val, sz, bold, color, font) in enumerate(values):
                    cell = row.cells[ci]
                    if row_bg:
                        set_cell_bg(cell, row_bg)
                    set_cell_w(cell, col_w[ci])
                    set_cell_valign(cell)
                    p = cell.paragraphs[0]
                    p.alignment = data_aligns[ci]
                    set_spacing(p, 2, 2)
                    add_run(p, val, sz, bold=bold, color=color, fn=font)

            # 小计行（合并前6列）
            sub_ri = 1 + len(details)
            merged_w = sum(col_w[:6])
            sub_lbl = t2.cell(sub_ri, 0).merge(t2.cell(sub_ri, 5))
            set_cell_bg(sub_lbl, 'E8EEF5')
            set_cell_w(sub_lbl, merged_w)
            set_cell_valign(sub_lbl)
            p = sub_lbl.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_spacing(p, 2, 2)
            add_run(p, 'Subtotal' if is_ovs else '小      计', 10, bold=True, color=C_DARK_BLUE)

            sub_right = [
                (6,  str(int(total_qty)),                    WD_ALIGN_PARAGRAPH.CENTER, 10, True,  C_DARK_BLUE, FONT_NAME),
                (7,  '—',                                    WD_ALIGN_PARAGRAPH.CENTER, 9,  False, C_GRAY,      FONT_NAME),
                (8,  '—',                                    WD_ALIGN_PARAGRAPH.CENTER, 9,  False, C_GRAY,      FONT_NAME),
                (9,  '—',                                    WD_ALIGN_PARAGRAPH.CENTER, 9,  False, C_GRAY,      FONT_NAME),
                (10, f'{curr_sym}{total_amount:,.2f}',        WD_ALIGN_PARAGRAPH.RIGHT,  10, True,  C_DARK_BLUE, MONEY_FONT),
            ]
            if has_item_note:
                sub_right.append((11, '', WD_ALIGN_PARAGRAPH.LEFT, 9, False, C_GRAY, FONT_NAME))

            for ci, val, align, sz, bold, color, font in sub_right:
                cell = t2.cell(sub_ri, ci)
                set_cell_bg(cell, 'E8EEF5')
                set_cell_w(cell, col_w[ci])
                set_cell_valign(cell)
                p = cell.paragraphs[0]
                p.alignment = align
                set_spacing(p, 2, 2)
                add_run(p, val, sz, bold=bold, color=color, fn=font)

            sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 0)

            # ── Table 3: 批准总金额（1行×3列）────────────────────────────
            t3 = doc.add_table(rows=1, cols=3)
            set_tbl_borders(t3, color='1F3A5F', sz=6)
            set_tbl_width(t3, TABLE_W)
            t3_ws = [3500, 7500, 4038]
            t3_row = t3.rows[0]

            # 左：批准总金额标题
            c0 = t3_row.cells[0]
            set_cell_bg(c0, '1F3A5F'); set_cell_w(c0, t3_ws[0]); set_cell_valign(c0)
            p = c0.paragraphs[0]; set_spacing(p, 4, 2)
            add_run(p, 'APPROVED TOTAL' if is_ovs else '批准总金额', 11, bold=True, color=C_WHITE)
            p2 = c0.add_paragraph(); set_spacing(p2, 2, 4)
            add_run(p2, '' if is_ovs else 'APPROVED TOTAL', 7, color=C_LIGHT_BLUE)

            # 中：大写金额（中文，YaHei）
            c1 = t3_row.cells[1]
            set_cell_bg(c1, '1F3A5F'); set_cell_w(c1, t3_ws[1]); set_cell_valign(c1)
            p = c1.paragraphs[0]; set_spacing(p, 4, 2)
            add_run(p, 'Amount in Words' if is_ovs else '金额大写(RMB in Capital)', 7, color=C_LIGHT_BLUE)
            p2 = c1.add_paragraph(); set_spacing(p2, 2, 4)
            cn_amount = amount_to_chinese(total_amount) if (currency == 'CNY' and not is_ovs) else f'N/A ({currency})'
            add_run(p2, cn_amount, 12, bold=True, color=C_WHITE)

            # 右：小写金额（Arial）
            c2 = t3_row.cells[2]
            set_cell_bg(c2, '1F3A5F'); set_cell_w(c2, t3_ws[2]); set_cell_valign(c2)
            p = c2.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_spacing(p, 4, 2)
            add_run(p, 'Amount' if is_ovs else '小写金额(Amount)', 7, color=C_LIGHT_BLUE)
            p2 = c2.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_spacing(p2, 2, 4)
            add_run(p2, f'{curr_sym}{total_amount:,.2f}', 16, bold=True, color=C_WHITE, fn=MONEY_FONT)

            # ── Table 4: 价格条款（6行×2列）──────────────────────────────
            if is_ovs:
                add_section_heading(doc, 'PRICING TERMS', '')
            else:
                add_section_heading(doc, '价格条款', 'PRICING TERMS')
            t4 = doc.add_table(rows=6, cols=2)
            set_tbl_borders(t4, color='DDDDDD', sz=4)
            set_tbl_width(t4, TABLE_W)
            TERM_LBL_W = 2200

            if is_ovs:
                terms = [
                    ('Price Lock',    'Approved prices remain fixed for the validity period regardless of market fluctuations.'),
                    ('Qty Adjustment','Actual order quantities may vary within ±20% of the approved quantity; excess requires re-approval.'),
                    ('Order Process', 'Dealers must quote this Pricing Order No. when placing orders to ensure approved pricing applies.'),
                    ('Payment',       'Payment terms follow the company standard channel payment policy (as confirmed in the formal order).'),
                    ('Delivery',      'Delivery date is subject to the confirmed purchase order.'),
                    ('Extension',     'To extend the validity period, apply at least 5 working days before expiry and obtain approval.'),
                ]
            else:
                terms = [
                    ('价格锁定',  '在有效期内，上述批准价格保持不变，不受市场价格波动影响。'),
                    ('数量调整',  '实际订单数量可在批价数量的 ±20% 范围内调整，超出部分需重新审批。'),
                    ('下单方式',  '经销商下单时须注明本批价单号，以便按批准价格执行。'),
                    ('付款条款',  '按公司标准渠道付款政策执行（具体以正式订单为准）。'),
                    ('交货安排',  '以正式采购订单确认的交货日期为准。'),
                    ('有效期延期', '如需延长价格有效期，须在到期前 5 个工作日内提交申请并获得审批。'),
                ]
            for ri, (term_name, term_content) in enumerate(terms):
                lc = t4.rows[ri].cells[0]
                vc = t4.rows[ri].cells[1]
                set_cell_bg(lc, 'E8EEF5'); set_cell_w(lc, TERM_LBL_W); set_cell_valign(lc)
                set_cell_w(vc, TABLE_W - TERM_LBL_W); set_cell_valign(vc)
                lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_spacing(lp, 2, 2)
                add_run(lp, term_name, 10, bold=True, color=C_DARK_BLUE)
                vp = vc.paragraphs[0]; set_spacing(vp, 2, 2)
                add_run(vp, term_content, 10, color=C_TEXT)

            # ── 备注说明 ──────────────────────────────────────────────────
            if is_ovs:
                add_section_heading(doc, 'REMARKS', '')
            else:
                add_section_heading(doc, '备注说明', 'REMARKS')

            # 整体备注（按需显示）
            if has_notes:
                pn = doc.add_paragraph(); set_spacing(pn, 1, 3)
                add_run(pn, 'Notes: ' if is_ovs else '批价备注：', 10, bold=True, color=C_DARK_BLUE)
                add_run(pn, str(pricing_order.notes).strip(), 10, color=C_TEXT)

            # 固定条款
            if is_ovs:
                fixed = [
                    'This Pricing Sheet is for price approval purposes only and does not constitute a formal purchase contract.',
                    'A formal purchase order must be issued separately; its terms shall prevail.',
                    'Product specifications are subject to change; the latest product datasheet shall govern.',
                    'This Pricing Sheet is issued in duplicate, one copy each for the supplier and the dealer.',
                ]
            else:
                fixed = [
                    '本批价单仅作为价格审批依据，不构成正式采购合同。',
                    '正式订单须另行签订，届时以正式订单条款为准。',
                    '本批价单所列产品规格如有变更，以最新产品规格书为准。',
                    '批价单一式两份，供应商与经销商各执一份。',
                ]
            for note in fixed:
                pf = doc.add_paragraph(); set_spacing(pf, 1, 1)
                add_run(pf, note, 10, color=C_TEXT)

            # ── 保存 ─────────────────────────────────────────────────────
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            safe_name = ''.join(c for c in project_name if c.isalnum() or c in ' -_（）【】').rstrip()
            filename = f'{pricing_order.order_number} & {safe_name}.docx'
            logger.info(f'生成批价单Word(v2): {filename}')
            return {'content': output.getvalue(), 'filename': filename}

        except Exception as e:
            logger.error(f'生成批价单Word(v2)失败: {e}')
            raise

    def generate_settlement_order_word_v2(self, pricing_order, include_notes=False):
        """生成结算单Word文档（优化版样式，A4竖向，与批价单 v2 风格一致）"""
        from docx.shared import RGBColor
        from docx.enum.table import WD_ALIGN_VERTICAL
        from app.models.pricing_order import SettlementOrder

        C_DARK_BLUE  = RGBColor(0x1F, 0x3A, 0x5F)
        C_GOLD       = RGBColor(0xB8, 0x95, 0x3F)
        C_GRAY       = RGBColor(0x59, 0x59, 0x59)
        C_TEXT       = RGBColor(0x1A, 0x1A, 0x1A)
        C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
        C_GREEN      = RGBColor(0x1F, 0x7A, 0x4A)
        C_LIGHT_BLUE = RGBColor(0xA8, 0xBE, 0xD8)

        FONT_NAME  = 'Microsoft YaHei'
        LATIN_FONT = 'Arial'   # 拉丁字符专用（LibreOffice 可找到，Word 与 YaHei 拉丁效果一致）
        MONEY_FONT = 'Arial'   # 金额数字专用字体

        def set_run_font(run, fn=FONT_NAME):
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            if fn == FONT_NAME:
                # 普通文字: 拉丁用 Arial（LibreOffice 可渲染），CJK 用微软雅黑
                rFonts.set(qn('w:ascii'),    LATIN_FONT)
                rFonts.set(qn('w:hAnsi'),    LATIN_FONT)
                rFonts.set(qn('w:eastAsia'), FONT_NAME)
                rFonts.set(qn('w:cs'),       FONT_NAME)
            else:
                # 金额/英文专用字体: 全部使用指定字体
                for attr in ('w:ascii', 'w:eastAsia', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), fn)

        def set_cell_bg(cell, hex_color):
            tcPr = cell._tc.get_or_add_tcPr()
            for s in tcPr.findall(qn('w:shd')):
                tcPr.remove(s)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        def set_cell_w(cell, dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            for w in tcPr.findall(qn('w:tcW')):
                tcPr.remove(w)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(dxa)); tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        def set_cell_valign(cell, val='center'):
            tcPr = cell._tc.get_or_add_tcPr()
            for v in tcPr.findall(qn('w:vAlign')):
                tcPr.remove(v)
            el = OxmlElement('w:vAlign'); el.set(qn('w:val'), val)
            tcPr.append(el)

        def set_spacing(para, before_pt=0, after_pt=0):
            pPr = para._p.get_or_add_pPr()
            for s in pPr.findall(qn('w:spacing')):
                pPr.remove(s)
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:before'), str(int(before_pt * 20)))
            sp.set(qn('w:after'), str(int(after_pt * 20)))
            sp.set(qn('w:line'), '276'); sp.set(qn('w:lineRule'), 'auto')
            pPr.append(sp)

        def add_run(para, text, size_pt, bold=False, color=None, fn=FONT_NAME):
            run = para.add_run(str(text))
            run.font.size = Pt(size_pt); run.font.bold = bold
            if color:
                run.font.color.rgb = color
            set_run_font(run, fn)
            return run

        def set_tbl_borders(table, color='CCCCCC', sz=4):
            tblPr = table._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
            borders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
                b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
                borders.append(b)
            for el in (tblPr.find(qn('w:tblBorders')), tblPr.find(qn('w:tblLook'))):
                if el is not None:
                    tblPr.remove(el)
            tblPr.append(borders)

        def set_tbl_width(table, dxa):
            tblPr = table._tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
            for w in tblPr.findall(qn('w:tblW')):
                tblPr.remove(w)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(dxa)); tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)

        def amount_to_chinese(amount):
            digits = '零壹贰叁肆伍陆柒捌玖'
            amount = round(float(amount or 0), 2)
            int_part = int(amount)
            dec = round((amount - int_part) * 100)
            jiao, fen = dec // 10, dec % 10
            def _int_cn(n):
                if n == 0: return '零'
                wan_u = ['', '万', '亿']; unit_u = ['', '拾', '佰', '仟']
                groups = []
                while n > 0:
                    groups.append(n % 10000); n //= 10000
                result = ''
                for gi in range(len(groups)-1, -1, -1):
                    g = groups[gi]
                    if g == 0:
                        if result and result[-1] != '零': result += '零'
                        continue
                    g_str, need_zero = '', False
                    for ui in range(3, -1, -1):
                        d = (g // (10**ui)) % 10
                        if d:
                            if need_zero: g_str += '零'; need_zero = False
                            g_str += digits[d] + unit_u[ui]
                        elif g_str: need_zero = True
                    result += g_str + wan_u[gi]
                return result
            int_str = _int_cn(int_part) + '元'
            if jiao == 0 and fen == 0: dec_str = '整'
            elif jiao == 0: dec_str = '零' + digits[fen] + '分'
            elif fen == 0: dec_str = digits[jiao] + '角'
            else: dec_str = digits[jiao] + '角' + digits[fen] + '分'
            return '人民币 ' + int_str + dec_str

        def add_section_heading(doc, cn_text, en_text):
            p = doc.add_paragraph(); set_spacing(p, before_pt=8, after_pt=3)
            add_run(p, '■ ', 11, color=C_GOLD)
            add_run(p, cn_text, 12, bold=True, color=C_DARK_BLUE)
            add_run(p, f'   {en_text}', 8, color=C_GRAY)
            return p

        # ── 数据准备 ─────────────────────────────────────────────────────
        try:
            settlement_order = SettlementOrder.query.filter_by(
                pricing_order_id=pricing_order.id
            ).first()
            if not settlement_order:
                from app.services.pricing_order_service import PricingOrderService
                from app import db
                settlement_order = PricingOrderService.create_settlement_order(
                    pricing_order, pricing_order.created_by)
                PricingOrderService.create_settlement_details(pricing_order, settlement_order)
                db.session.commit()

            import os as _os
            is_ovs    = Config.IS_OVS or _os.environ.get('WORD_RENDER_LOCALE') == 'ovs'
            date_fmt  = '%d %b %Y' if is_ovs else '%Y年%m月%d日'

            currency     = pricing_order.currency or Config.DEFAULT_CURRENCY
            curr_sym     = self._get_currency_symbol(currency)
            project      = pricing_order.project
            project_name = project.project_name if project else ''

            created_date = pricing_order.approved_at or pricing_order.created_at
            date_str     = created_date.strftime(date_fmt) if created_date else ''

            distributor_name = pricing_order.distributor.company_name if pricing_order.distributor else ('(None)'   if is_ovs else '（无分销商）')
            dealer_name      = pricing_order.dealer.company_name      if pricing_order.dealer      else ('(Direct)' if is_ovs else '（厂商直签）')
            proj_mgr = (project.owner.real_name if project and project.owner else '') or ''
            sales_mgr = ''
            if project and getattr(project, 'vendor_sales_manager', None):
                sales_mgr = project.vendor_sales_manager.real_name or ''

            if is_ovs:
                type_map = {'channel_follow': 'Channel Follow', 'sales_key': 'Key Account',
                            'sales_opportunity': 'Sales Opportunity', 'sales_focus': 'Focus Sales'}
                _status_en = {'draft': 'Draft', 'pending': 'Pending', 'approved': 'Approved', 'rejected': 'Rejected'}
                status_label = _status_en.get(pricing_order.status, pricing_order.status or 'Unknown')
            else:
                type_map = {'channel_follow': '渠道跟进', 'sales_key': '销售重点',
                            'sales_opportunity': '销售机会', 'sales_focus': '重点销售'}
                status_label = self._get_status_label(pricing_order.status)
            flow_label   = type_map.get(pricing_order.approval_flow_type, pricing_order.approval_flow_type or '')
            status_color = C_GREEN if pricing_order.status == 'approved' else C_DARK_BLUE
            status_prefix = '✔ ' if pricing_order.status == 'approved' else ''

            total_rate   = pricing_order.settlement_total_discount_rate or 1.0
            discount_str = f'{int(round(total_rate * 100))} %'
            total_amount = pricing_order.settlement_total_amount or 0

            details      = list(pricing_order.settlement_details)
            total_qty    = sum(d.quantity for d in details)
            has_item_note = include_notes and any(getattr(d, 'item_note', None) for d in details)
            has_notes     = include_notes and bool(pricing_order.notes and str(pricing_order.notes).strip())

            # A4 横向可用宽度: 29.7 - 1.59*2 = 26.52cm → 15038 dxa
            TABLE_W = 15038

            # ── 创建文档（A4 横向）──────────────────────────────────────
            doc = Document()
            section = doc.sections[0]
            section.page_width    = Cm(29.7)
            section.page_height   = Cm(21.0)
            section.left_margin   = Cm(1.59)
            section.right_margin  = Cm(1.59)
            section.top_margin    = Cm(1.59)
            section.bottom_margin = Cm(1.59)
            doc.styles['Normal'].paragraph_format.space_before = Pt(0)
            doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

            # docDefaults 字体
            styles_el = doc.styles.element
            dd = styles_el.find(qn('w:docDefaults'))
            if dd is None:
                dd = OxmlElement('w:docDefaults'); styles_el.insert(0, dd)
            rPrD = dd.find(qn('w:rPrDefault'))
            if rPrD is None:
                rPrD = OxmlElement('w:rPrDefault'); dd.append(rPrD)
            rPr_d = rPrD.find(qn('w:rPr'))
            if rPr_d is None:
                rPr_d = OxmlElement('w:rPr'); rPrD.append(rPr_d)
            rF = rPr_d.find(qn('w:rFonts'))
            if rF is None:
                rF = OxmlElement('w:rFonts'); rPr_d.insert(0, rF)
            rF.set(qn('w:ascii'),    LATIN_FONT)   # 拉丁字符用 Arial
            rF.set(qn('w:hAnsi'),    LATIN_FONT)
            rF.set(qn('w:eastAsia'), FONT_NAME)    # CJK 用微软雅黑
            rF.set(qn('w:cs'),       FONT_NAME)
            for tag, val in (('w:sz', '20'), ('w:szCs', '20')):
                el = rPr_d.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag); rPr_d.append(el)
                el.set(qn('w:val'), val)

            # ── 标题区 ───────────────────────────────────────────────────
            p0 = doc.add_paragraph()
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_spacing(p0, 0, 2)
            add_run(p0, 'SETTLEMENT  SHEET' if is_ovs else '结 算 确 认 单', 22, bold=True, color=C_DARK_BLUE)

            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_spacing(p1, 0, 2)
            add_run(p1, '' if is_ovs else 'SETTLEMENT SHEET', 10, color=C_GRAY)

            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_spacing(p2, 0, 6)
            add_run(p2, '◆  ◆  ◆', 8, color=C_GOLD)

            # ── Table 0: 基本信息（4行×6列）──────────────────────────────
            # landscape 6 列: LBL=1400, VAL=3613 → 3*(1400+3613)=15039≈15038
            LBL_W, VAL_W = 1400, 3613
            t0 = doc.add_table(rows=4, cols=6)
            set_tbl_borders(t0); set_tbl_width(t0, TABLE_W)

            def fill_info_row(row_idx, triples):
                row = t0.rows[row_idx]
                for i, (lbl, val, vc, vb) in enumerate(triples):
                    lc = row.cells[i * 2]; vc_ = row.cells[i * 2 + 1]
                    set_cell_bg(lc, 'E8EEF5')
                    set_cell_w(lc, LBL_W); set_cell_w(vc_, VAL_W)
                    set_cell_valign(lc); set_cell_valign(vc_)
                    lp = lc.paragraphs[0]; set_spacing(lp, 2, 2)
                    add_run(lp, lbl, 9, bold=True, color=C_DARK_BLUE)
                    vp = vc_.paragraphs[0]; set_spacing(vp, 2, 2)
                    add_run(vp, val, 10, bold=vb, color=vc)

            if is_ovs:
                _S0 = [('Order No.',      settlement_order.order_number, C_GOLD, True),
                        ('Date',           date_str,                      C_TEXT, False),
                        ('Pricing Order',  pricing_order.order_number,    C_TEXT, False)]
                _S1 = [('Project',         project_name, C_TEXT, True),
                        ('Type',           flow_label,   C_TEXT, False),
                        ('Disc. Rate',     discount_str, C_GOLD, True)]
                _lbl_entity = 'Settle To'
                _lbl_note   = 'Note'
                _settle_note = ('This Settlement Sheet is the distributor pickup settlement voucher, '
                                'executed at distributor price, managed separately from the Pricing Sheet (dealer price).')
            else:
                _S0 = [('结算单号',   settlement_order.order_number, C_GOLD, True),
                        ('结算日期',  date_str,                      C_TEXT, False),
                        ('关联批价单', pricing_order.order_number,   C_TEXT, False)]
                _S1 = [('项目名称', project_name, C_TEXT, True),
                        ('项目类型', flow_label,   C_TEXT, False),
                        ('总折扣率', discount_str, C_GOLD, True)]
                _lbl_entity = '结算对象'
                _lbl_note   = '结算说明'
                _settle_note = '本结算单为分销商提货结算凭证，按分销价格执行，与批价单（经销商价格）分开管理。'

            fill_info_row(0, _S0)
            fill_info_row(1, _S1)

            # Row 2: 结算对象（label + 合并 col 1-5）
            lc2 = t0.rows[2].cells[0]
            set_cell_bg(lc2, 'E8EEF5'); set_cell_w(lc2, LBL_W); set_cell_valign(lc2)
            lp2 = lc2.paragraphs[0]; set_spacing(lp2, 2, 2)
            add_run(lp2, _lbl_entity, 9, bold=True, color=C_DARK_BLUE)
            vc2 = t0.cell(2, 1).merge(t0.cell(2, 5))
            set_cell_valign(vc2)
            vp2 = vc2.paragraphs[0]; set_spacing(vp2, 2, 2)
            add_run(vp2, distributor_name, 10, color=C_TEXT)

            # Row 3: 结算说明（label + 合并 col 1-5，固定说明文字）
            lc3 = t0.rows[3].cells[0]
            set_cell_bg(lc3, 'E8EEF5'); set_cell_w(lc3, LBL_W); set_cell_valign(lc3)
            lp3 = lc3.paragraphs[0]; set_spacing(lp3, 2, 2)
            add_run(lp3, _lbl_note, 9, bold=True, color=C_DARK_BLUE)
            vc3 = t0.cell(3, 1).merge(t0.cell(3, 5))
            set_cell_valign(vc3)
            vp3 = vc3.paragraphs[0]; set_spacing(vp3, 2, 2)
            add_run(vp3, _settle_note, 9, color=C_TEXT)

            sp1 = doc.add_paragraph(); set_spacing(sp1, 0, 0)

            # ── Table 1: 供应商 + 渠道信息 ────────────────────────────────
            t1 = doc.add_table(rows=1, cols=2)
            set_tbl_borders(t1, color='DDDDDD'); set_tbl_width(t1, TABLE_W)
            half = 7519  # 15038 // 2

            def fill_info_card(cell, title_cn, title_en, fields):
                set_cell_bg(cell, 'FBFCFD')
                set_cell_valign(cell, 'top'); set_cell_w(cell, half)
                p = cell.paragraphs[0]; set_spacing(p, 3, 3)
                add_run(p, title_cn, 11, bold=True, color=C_DARK_BLUE)
                add_run(p, f'  {title_en}', 8, color=C_GRAY)
                for lbl, val in fields:
                    fp = cell.add_paragraph(); set_spacing(fp, 1, 1)
                    add_run(fp, lbl, 9, color=C_GRAY)
                    add_run(fp, val, 10, color=C_TEXT)

            if is_ovs:
                fill_info_card(t1.rows[0].cells[0], 'SUPPLIER', '', [
                    ('Company :  ', 'Evertac Solutions Singaproe Pte Ltd.'),
                    ('Website :  ', 'http://www.evertacsolutions.com'),
                ])
                fill_info_card(t1.rows[0].cells[1], 'CHANNEL', '', [
                    ('Distributor :  ', distributor_name),
                    ('Dealer :  ',      dealer_name),
                    ('Project Mgr :  ', proj_mgr),
                    ('Sales Mgr :  ',   sales_mgr),
                ])
            else:
                fill_info_card(t1.rows[0].cells[0], '供应商信息', 'SUPPLIER', [
                    ('企业名称 :  ', '和源通信(上海)股份有限公司'),
                    ('办公地址 :  ', '上海市普陀区武威路88号19楼6楼'),
                    ('联系电话 :  ', '021-62596028'),
                    ('官方网址 :  ', 'http://www.evertac.net'),
                ])
                fill_info_card(t1.rows[0].cells[1], '渠道信息', 'CHANNEL', [
                    ('分 销 商 :  ', distributor_name),
                    ('经 销 商 :  ', dealer_name),
                    ('项目负责人 :  ', proj_mgr),
                    ('销售负责人 :  ', sales_mgr),
                ])

            # ── 结算明细表格 ──────────────────────────────────────────────
            if is_ovs:
                add_section_heading(doc, 'SETTLEMENT PRICE DETAILS', '')
            else:
                add_section_heading(doc, '结算价格明细', 'SETTLEMENT PRICE DETAILS')

            if has_item_note:
                col_w = [500, 1600, 1400, 2200, 720, 1280, 720, 1320, 1000, 1320, 1678, 1300]
                if is_ovs:
                    headers = ['No.', 'Product Name', 'Model', 'Specifications', 'Brand', 'Code',
                               'Qty', 'List Price', 'Disc.', 'Unit Price', 'Amount', 'Remarks']
                else:
                    headers = ['序号','产品名称','型号','规格参数','品牌','产品编码',
                               '数量','零售单价','折扣率','结算单价','结算总价','备注']
            else:
                col_w = [500, 1600, 1400, 2900, 720, 1280, 720, 1320, 1000, 1320, 2278]
                if is_ovs:
                    headers = ['No.', 'Product Name', 'Model', 'Specifications', 'Brand', 'Code',
                               'Qty', 'List Price', 'Disc.', 'Unit Price', 'Amount']
                else:
                    headers = ['序号','产品名称','型号','规格参数','品牌','产品编码',
                               '数量','零售单价','折扣率','结算单价','结算总价']
            # 确保列宽总和等于 TABLE_W（以规格参数列吸收误差）
            diff = TABLE_W - sum(col_w)
            col_w[3] += diff
            n_cols = len(headers)

            t2 = doc.add_table(rows=1 + len(details) + 1, cols=n_cols)
            set_tbl_borders(t2); set_tbl_width(t2, TABLE_W)

            hdr_aligns = [WD_ALIGN_PARAGRAPH.CENTER] * n_cols
            hdr_aligns[3] = WD_ALIGN_PARAGRAPH.LEFT
            for ci, (hdr, w) in enumerate(zip(headers, col_w)):
                cell = t2.rows[0].cells[ci]
                set_cell_bg(cell, '1F3A5F'); set_cell_w(cell, w); set_cell_valign(cell)
                p = cell.paragraphs[0]; p.alignment = hdr_aligns[ci]; set_spacing(p, 2, 2)
                add_run(p, hdr, 9, bold=True, color=C_WHITE)

            data_aligns = [
                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.LEFT,   WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
                WD_ALIGN_PARAGRAPH.RIGHT,
            ]
            if has_item_note:
                data_aligns.append(WD_ALIGN_PARAGRAPH.LEFT)

            for di, detail in enumerate(details):
                row_bg = 'F7F9FC' if di % 2 == 1 else None
                disc_pct = f'{int(round(detail.discount_rate * 100))} %'
                values = [
                    (str(di + 1),                               9,  False, C_TEXT,       FONT_NAME),
                    (detail.product_name or '',                 9,  True,  C_TEXT,       FONT_NAME),
                    (detail.product_model or '',                9,  False, C_TEXT,       FONT_NAME),
                    (detail.product_desc or '',                 8,  False, C_GRAY,       FONT_NAME),
                    (detail.brand or '',                        9,  False, C_TEXT,       FONT_NAME),
                    (getattr(detail, 'product_mn', '') or '',   8,  False, C_TEXT,       FONT_NAME),
                    (str(detail.quantity),                      9,  True,  C_TEXT,       FONT_NAME),
                    (f'{curr_sym}{detail.market_price:,.2f}',   9,  False, C_GRAY,       MONEY_FONT),
                    (disc_pct,                                  9,  True,  C_GOLD,       FONT_NAME),
                    (f'{curr_sym}{detail.unit_price:,.2f}',     9,  True,  C_TEXT,       MONEY_FONT),
                    (f'{curr_sym}{detail.total_price:,.2f}',    9,  True,  C_DARK_BLUE,  MONEY_FONT),
                ]
                if has_item_note:
                    values.append((getattr(detail, 'item_note', '') or '', 8, False, C_GRAY, FONT_NAME))

                row = t2.rows[1 + di]
                for ci, (val, sz, bold, color, font) in enumerate(values):
                    cell = row.cells[ci]
                    if row_bg: set_cell_bg(cell, row_bg)
                    set_cell_w(cell, col_w[ci]); set_cell_valign(cell)
                    p = cell.paragraphs[0]; p.alignment = data_aligns[ci]; set_spacing(p, 2, 2)
                    add_run(p, val, sz, bold=bold, color=color, fn=font)

            # 小计行
            sub_ri = 1 + len(details)
            merged_w = sum(col_w[:6])
            sub_lbl = t2.cell(sub_ri, 0).merge(t2.cell(sub_ri, 5))
            set_cell_bg(sub_lbl, 'E8EEF5'); set_cell_w(sub_lbl, merged_w); set_cell_valign(sub_lbl)
            p = sub_lbl.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_spacing(p, 2, 2)
            add_run(p, 'Subtotal' if is_ovs else '小      计', 10, bold=True, color=C_DARK_BLUE)

            sub_right = [
                (6,  str(int(total_qty)),                   WD_ALIGN_PARAGRAPH.CENTER, 10, True,  C_DARK_BLUE, FONT_NAME),
                (7,  '—',                                   WD_ALIGN_PARAGRAPH.CENTER, 9,  False, C_GRAY,      FONT_NAME),
                (8,  '—',                                   WD_ALIGN_PARAGRAPH.CENTER, 9,  False, C_GRAY,      FONT_NAME),
                (9,  '—',                                   WD_ALIGN_PARAGRAPH.CENTER, 9,  False, C_GRAY,      FONT_NAME),
                (10, f'{curr_sym}{total_amount:,.2f}',       WD_ALIGN_PARAGRAPH.RIGHT,  10, True,  C_DARK_BLUE, MONEY_FONT),
            ]
            if has_item_note:
                sub_right.append((11, '', WD_ALIGN_PARAGRAPH.LEFT, 9, False, C_GRAY, FONT_NAME))
            for ci, val, align, sz, bold, color, font in sub_right:
                cell = t2.cell(sub_ri, ci)
                set_cell_bg(cell, 'E8EEF5'); set_cell_w(cell, col_w[ci]); set_cell_valign(cell)
                p = cell.paragraphs[0]; p.alignment = align; set_spacing(p, 2, 2)
                add_run(p, val, sz, bold=bold, color=color, fn=font)

            sp2 = doc.add_paragraph(); set_spacing(sp2, 0, 0)

            # ── Table 3: 结算总金额 ───────────────────────────────────────
            t3 = doc.add_table(rows=1, cols=3)
            set_tbl_borders(t3, color='1F3A5F', sz=6); set_tbl_width(t3, TABLE_W)
            t3_ws = [3500, 7500, 4038]  # sum=15038

            c0 = t3.rows[0].cells[0]
            set_cell_bg(c0, '1F3A5F'); set_cell_w(c0, t3_ws[0]); set_cell_valign(c0)
            p = c0.paragraphs[0]; set_spacing(p, 4, 2)
            add_run(p, 'SETTLEMENT TOTAL' if is_ovs else '结算总金额', 11, bold=True, color=C_WHITE)
            p2 = c0.add_paragraph(); set_spacing(p2, 2, 4)
            add_run(p2, '' if is_ovs else 'SETTLEMENT TOTAL', 7, color=C_LIGHT_BLUE)

            c1 = t3.rows[0].cells[1]
            set_cell_bg(c1, '1F3A5F'); set_cell_w(c1, t3_ws[1]); set_cell_valign(c1)
            p = c1.paragraphs[0]; set_spacing(p, 4, 2)
            add_run(p, 'Amount in Words' if is_ovs else '金额大写(RMB in Capital)', 7, color=C_LIGHT_BLUE)
            p2 = c1.add_paragraph(); set_spacing(p2, 2, 4)
            cn_amount = amount_to_chinese(total_amount) if (currency == 'CNY' and not is_ovs) else f'N/A ({currency})'
            add_run(p2, cn_amount, 11, bold=True, color=C_WHITE)

            c2 = t3.rows[0].cells[2]
            set_cell_bg(c2, '1F3A5F'); set_cell_w(c2, t3_ws[2]); set_cell_valign(c2)
            p = c2.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_spacing(p, 4, 2)
            add_run(p, 'Amount' if is_ovs else '小写金额(Amount)', 7, color=C_LIGHT_BLUE)
            p2 = c2.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_spacing(p2, 2, 4)
            add_run(p2, f'{curr_sym}{total_amount:,.2f}', 14, bold=True, color=C_WHITE, fn=MONEY_FONT)

            # ── Table 4: 结算条款 ─────────────────────────────────────────
            if is_ovs:
                add_section_heading(doc, 'SETTLEMENT TERMS', '')
            else:
                add_section_heading(doc, '结算条款', 'SETTLEMENT TERMS')
            t4 = doc.add_table(rows=6, cols=2)
            set_tbl_borders(t4, color='DDDDDD'); set_tbl_width(t4, TABLE_W)
            TERM_LBL_W = 2200

            if is_ovs:
                terms = [
                    ('Settlement Basis', f'This Settlement Sheet is executed per the referenced Pricing Order ({pricing_order.order_number}); settlement price is the distributor pickup price.'),
                    ('Payment Method',   'Telegraphic Transfer (T/T); goods shipped upon receipt of payment.'),
                    ('Payment Due',      'Please complete payment within 7 working days of settlement confirmation.'),
                    ('Invoice Type',     'VAT invoice or standard commercial invoice as applicable.'),
                    ('Delivery Terms',   'EXW (distributor self-collect or nominated carrier).'),
                    ('Lead Time',        'Goods will be dispatched within 5 working days after payment is received.'),
                ]
            else:
                terms = [
                    ('结算依据',  f'本结算单依据关联批价单（{pricing_order.order_number}）执行，结算价格为分销商提货价。'),
                    ('付款方式',  '电汇（T/T），款到发货。'),
                    ('付款期限',  '请于结算单确认后 7 个工作日内完成付款。'),
                    ('发票类型',  '增值税专用发票（13%）。'),
                    ('交货方式',  'EXW 上海（分销商自提或指定物流）。'),
                    ('交货时间',  '款到后 5 个工作日内安排发货。'),
                ]
            for ri, (term_name, term_content) in enumerate(terms):
                lc = t4.rows[ri].cells[0]; vc = t4.rows[ri].cells[1]
                set_cell_bg(lc, 'E8EEF5'); set_cell_w(lc, TERM_LBL_W); set_cell_valign(lc)
                set_cell_w(vc, TABLE_W - TERM_LBL_W); set_cell_valign(vc)
                lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_spacing(lp, 2, 2)
                add_run(lp, term_name, 10, bold=True, color=C_DARK_BLUE)
                vp = vc.paragraphs[0]; set_spacing(vp, 2, 2)
                add_run(vp, term_content, 10, color=C_TEXT)

            # ── 备注说明 ──────────────────────────────────────────────────
            if is_ovs:
                add_section_heading(doc, 'REMARKS', '')
            else:
                add_section_heading(doc, '备注说明', 'REMARKS')

            if has_notes:
                pn = doc.add_paragraph(); set_spacing(pn, 1, 3)
                add_run(pn, 'Notes: ' if is_ovs else '结算备注：', 10, bold=True, color=C_DARK_BLUE)
                add_run(pn, str(pricing_order.notes).strip(), 10, color=C_TEXT)

            if is_ovs:
                fixed = [
                    'This Settlement Sheet is the settlement voucher between the supplier and distributor, managed separately from the Pricing Sheet (supplier–dealer).',
                    'The distributor pays the supplier at the settlement price and supplies the dealer at the pricing order price.',
                    'Please quote the Settlement Order No. when making payment for reconciliation purposes.',
                    'This Settlement Sheet is issued in duplicate, one copy each for the supplier and the distributor.',
                ]
            else:
                fixed = [
                    '本结算单为厂商与分销商之间的结算凭证，与批价单（厂商与经销商）分开管理。',
                    '分销商按本结算价格向厂商付款提货，再按批价单价格向经销商供货。',
                    '付款时请注明结算单号，以便核销。',
                    '结算单一式两份，供应商与分销商各执一份。',
                ]
            for note in fixed:
                pf = doc.add_paragraph(); set_spacing(pf, 1, 1)
                add_run(pf, note, 10, color=C_TEXT)

            # ── 保存 ─────────────────────────────────────────────────────
            output = BytesIO()
            doc.save(output); output.seek(0)

            safe_name = ''.join(c for c in project_name if c.isalnum() or c in ' -_（）【】').rstrip()
            filename = f'{settlement_order.order_number} & {safe_name}.docx'
            logger.info(f'生成结算单Word(v2): {filename}')
            return {'content': output.getvalue(), 'filename': filename}

        except Exception as e:
            logger.error(f'生成结算单Word(v2)失败: {e}')
            raise

    def generate_settlement_order_word(self, pricing_order, include_notes=False):
        """
        生成结算单Word文档

        Args:
            pricing_order: PricingOrder 对象

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        try:
            # 获取关联的结算单
            from app.models.pricing_order import SettlementOrder
            settlement_order = SettlementOrder.query.filter_by(pricing_order_id=pricing_order.id).first()

            # 如果结算单不存在，创建一个
            if not settlement_order:
                logger.warning(f"批价单 {pricing_order.order_number} 没有关联的结算单，正在创建...")
                from app.services.pricing_order_service import PricingOrderService
                from app import db

                settlement_order = PricingOrderService.create_settlement_order(
                    pricing_order,
                    pricing_order.created_by
                )
                PricingOrderService.create_settlement_details(pricing_order, settlement_order)
                db.session.commit()

            template_path = self._get_template_path('settlement_order_template.docx')
            doc = Document(template_path)

            # 准备基本数据
            currency = pricing_order.currency or Config.DEFAULT_CURRENCY
            currency_symbol = self._get_currency_symbol(currency)

            # 表格0: 基本信息
            table0 = doc.tables[0]
            replacements0 = {
                'SO202512-003': settlement_order.order_number,
                '2025年12月18日': (pricing_order.approved_at or pricing_order.created_at).strftime('%Y年%m月%d日') if (pricing_order.approved_at or pricing_order.created_at) else '',
                '上海东方枢纽': pricing_order.project.project_name if pricing_order.project else '',
                '渠道跟进': self._get_project_type_label(pricing_order.approval_flow_type),
                'PO202512-003': pricing_order.order_number,
                '40.50%': self._format_discount_rate(pricing_order.settlement_total_discount_rate or 1.0),
            }
            self._replace_text_in_table(table0, replacements0)

            # 表格1: 结算对象
            table1 = doc.tables[1]
            distributor_name = pricing_order.distributor.company_name if pricing_order.distributor else '（无分销商）'
            replacements1 = {
                '上海淳泊信息科技有限公司': distributor_name,
            }
            self._replace_text_in_table(table1, replacements1)

            # 表格2: 渠道信息
            table2 = doc.tables[2]
            dealer_name = pricing_order.dealer.company_name if pricing_order.dealer else '（厂商直签）'
            replacements2 = {
                '上海瑞康通信科技有限公司': dealer_name,
                '上海淳泊信息科技有限公司': distributor_name,
            }
            self._replace_text_in_table(table2, replacements2)

            # 表格3: 明细表格 - 动态处理
            detail_table = doc.tables[3]

            # 删除示例数据行（保留表头）
            while len(detail_table.rows) > 1:
                tr = detail_table.rows[-1]._tr
                detail_table._tbl.remove(tr)

            # 添加实际明细数据
            for idx, detail in enumerate(pricing_order.settlement_details, 1):
                row = detail_table.add_row()
                cells = row.cells

                # 设置各列数据
                self._set_cell_text(cells[0], idx)  # 序号
                self._set_cell_text(cells[1], detail.product_name or '')  # 产品名称
                self._set_cell_text(cells[2], detail.product_model or '')  # 型号
                # 规格参数（截断长文本）
                desc = detail.product_desc or ''
                if len(desc) > 40:
                    desc = desc[:40] + '...'
                self._set_cell_text(cells[3], desc)
                self._set_cell_text(cells[4], detail.brand or '和源')  # 品牌
                self._set_cell_text(cells[5], detail.quantity)  # 数量
                self._set_cell_text(cells[6], f'{currency_symbol}{detail.market_price:,.2f}')  # 零售单价
                self._set_cell_text(cells[7], self._format_discount_rate(detail.discount_rate))  # 折扣率
                self._set_cell_text(cells[8], f'{currency_symbol}{detail.unit_price:,.2f}')  # 结算单价
                self._set_cell_text(cells[9], f'{currency_symbol}{detail.total_price:,.2f}')  # 总价（数据库字段）
                self._set_cell_text(cells[10], detail.product_mn or '')  # 产品编码

            # 表格4: 总金额
            table4 = doc.tables[4]
            replacements4 = {
                '¥ 56,876.58': f'{currency_symbol} {pricing_order.settlement_total_amount:,.2f}',
            }
            self._replace_text_in_table(table4, replacements4)

            # 保存到内存
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            # 生成文件名
            project_name = pricing_order.project.project_name if pricing_order.project else "未知项目"
            safe_project_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_', '（', '）', '【', '】')).rstrip()
            filename = f"{settlement_order.order_number} & {safe_project_name}.docx"

            logger.info(f"成功生成结算单Word文档: {filename}")

            return {
                'content': output.getvalue(),
                'filename': filename
            }

        except Exception as e:
            logger.error(f"生成结算单Word文档失败: {str(e)}")
            raise

    def convert_word_to_pdf(self, word_content, filename):
        """
        将Word文档转换为PDF

        Args:
            word_content: Word文档的字节内容
            filename: 原始文件名

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                tmp_docx.write(word_content)
                tmp_docx_path = tmp_docx.name

            # PDF输出路径
            tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')

            # 优先走 Mac mini office-convert 服务（容器无需 LibreOffice）
            service_url = os.environ.get('PMA_OFFICE_CONVERT_URL', '').rstrip('/')
            if service_url:
                import requests as _http
                try:
                    resp = _http.post(
                        f"{service_url}/convert",
                        files={'file': (os.path.basename(tmp_docx_path), word_content,
                                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
                        data={'target': 'pdf'},
                        timeout=120,
                    )
                    if resp.status_code == 200:
                        with open(tmp_pdf_path, 'wb') as f:
                            f.write(resp.content)
                    else:
                        logger.warning(f"office-convert 服务返回 {resp.status_code}: {resp.text[:200]}")
                        raise Exception(f"远程转换失败 HTTP {resp.status_code}")
                except _http.exceptions.RequestException as e:
                    logger.error(f"office-convert 服务不可达: {e}")
                    raise Exception(f"远程转换服务不可达: {e}")
                # 跳过本地路径
                system = None
            else:
                system = platform.system()

            if system == "Darwin":  # macOS
                # 尝试使用 LibreOffice
                libreoffice_paths = [
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                    '/usr/local/bin/soffice',
                    'soffice'
                ]

                soffice_path = None
                for path in libreoffice_paths:
                    if os.path.exists(path) or path == 'soffice':
                        soffice_path = path
                        break

                if soffice_path:
                    cmd = [
                        soffice_path,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', os.path.dirname(tmp_docx_path),
                        tmp_docx_path
                    ]
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

                    if result.returncode != 0:
                        logger.warning(f"LibreOffice转换失败: {result.stderr}")
                        raise Exception("LibreOffice转换失败")
                else:
                    # 尝试使用 docx2pdf (需要 Microsoft Word)
                    try:
                        from docx2pdf import convert
                        convert(tmp_docx_path, tmp_pdf_path)
                    except ImportError:
                        raise Exception("未找到PDF转换工具，请安装LibreOffice或docx2pdf")

            elif system == "Linux":
                # Linux 使用 LibreOffice
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(tmp_docx_path),
                    tmp_docx_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

                if result.returncode != 0:
                    logger.error(f"LibreOffice转换失败: {result.stderr}")
                    raise Exception("LibreOffice转换失败")

            elif system == "Windows":
                # Windows 使用 docx2pdf
                try:
                    from docx2pdf import convert
                    convert(tmp_docx_path, tmp_pdf_path)
                except ImportError:
                    raise Exception("请安装docx2pdf: pip install docx2pdf")

            # 读取PDF内容
            if os.path.exists(tmp_pdf_path):
                with open(tmp_pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()

                # 清理临时文件
                os.unlink(tmp_docx_path)
                os.unlink(tmp_pdf_path)

                # 生成PDF文件名
                pdf_filename = filename.replace('.docx', '.pdf')

                return {
                    'content': pdf_content,
                    'filename': pdf_filename
                }
            else:
                raise Exception(f"PDF文件未生成: {tmp_pdf_path}")

        except subprocess.TimeoutExpired:
            logger.error("PDF转换超时")
            raise Exception("PDF转换超时")
        except Exception as e:
            logger.error(f"Word转PDF失败: {str(e)}")
            # 清理临时文件
            if 'tmp_docx_path' in locals() and os.path.exists(tmp_docx_path):
                os.unlink(tmp_docx_path)
            if 'tmp_pdf_path' in locals() and os.path.exists(tmp_pdf_path):
                os.unlink(tmp_pdf_path)
            raise

    def generate_pricing_order_pdf(self, pricing_order, include_notes=False):
        """
        使用Word模板生成批价单PDF

        Args:
            pricing_order: PricingOrder 对象

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        # 先生成Word（使用优化版生成器）
        word_result = self.generate_pricing_order_word_v2(pricing_order, include_notes=include_notes)

        # 转换为PDF
        pdf_result = self.convert_word_to_pdf(
            word_result['content'],
            word_result['filename']
        )

        return pdf_result

    def generate_settlement_order_pdf(self, pricing_order, include_notes=False):
        """
        使用Word模板生成结算单PDF

        Args:
            pricing_order: PricingOrder 对象

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        # 先生成Word
        word_result = self.generate_settlement_order_word_v2(pricing_order, include_notes=include_notes)

        # 转换为PDF
        pdf_result = self.convert_word_to_pdf(
            word_result['content'],
            word_result['filename']
        )

        return pdf_result


    def generate_quotation_word(self, quotation):
        """
        生成报价单Word文档

        Args:
            quotation: Quotation 对象

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        try:
            template_path = self._get_template_path('quotation_template.docx')
            doc = Document(template_path)

            # 准备基本数据
            currency = quotation.currency or Config.DEFAULT_CURRENCY
            currency_symbol = self._get_currency_symbol(currency)

            # 表格0: 客户信息
            table0 = doc.tables[0]
            customer_name = quotation.customer.company_name if quotation.customer else ''
            contact_name = quotation.contact.name if quotation.contact else ''
            quotation_date = quotation.created_at.strftime('%Y年%m月%d日') if quotation.created_at else ''

            replacements0 = {
                '上海市长宁区新泾镇人民政府': customer_name,
                '包惠青': contact_name,
            }
            self._replace_text_in_table(table0, replacements0)

            # 表格1: 明细表格 - 动态处理
            detail_table = doc.tables[1]

            # 删除示例数据行（保留表头）
            while len(detail_table.rows) > 1:
                tr = detail_table.rows[-1]._tr
                detail_table._tbl.remove(tr)

            # 添加实际明细数据（包括父子关系）
            idx = 0
            for detail in quotation.details:
                # 跳过配置产品（子产品会在父产品后面输出）
                if detail.is_accessory:
                    continue

                idx += 1
                row = detail_table.add_row()
                cells = row.cells

                # 设置父产品数据（序号居中，名称/型号/规格左对齐，数量居中，单价/金额右对齐）
                self._set_cell_text(cells[0], idx, align='center')  # 序号
                self._set_cell_text(cells[1], detail.product_name or '', align='left')  # 产品名称
                self._set_cell_text(cells[2], detail.product_model or '', align='left')  # 型号
                # 规格参数（截断长文本）
                desc = detail.product_desc or ''
                if len(desc) > 50:
                    desc = desc[:50] + '...'
                self._set_cell_text(cells[3], desc, align='left')  # 规格
                self._set_cell_text(cells[4], detail.quantity, align='center')  # 数量
                self._set_cell_text(cells[5], f'{currency_symbol}{detail.unit_price:,.2f}', align='right')  # 单价
                self._set_cell_text(cells[6], f'{currency_symbol}{detail.total_price:,.2f}', align='right')  # 金额

                # 输出子产品（配置产品）
                if detail.configurations:
                    for child in detail.configurations:
                        idx += 1
                        child_row = detail_table.add_row()
                        child_cells = child_row.cells

                        self._set_cell_text(child_cells[0], idx, align='center')  # 子产品也有序号
                        self._set_cell_text(child_cells[1], child.product_name or '', align='left')  # 产品名称
                        self._set_cell_text(child_cells[2], child.product_model or '', align='left')  # 型号
                        child_desc = child.product_desc or ''
                        if len(child_desc) > 50:
                            child_desc = child_desc[:50] + '...'
                        self._set_cell_text(child_cells[3], child_desc, align='left')  # 规格
                        self._set_cell_text(child_cells[4], child.quantity, align='center')  # 数量
                        self._set_cell_text(child_cells[5], f'{currency_symbol}{child.unit_price:,.2f}', align='right')  # 单价
                        self._set_cell_text(child_cells[6], f'{currency_symbol}{child.total_price:,.2f}', align='right')  # 金额

            # 表格2: 汇总金额
            table2 = doc.tables[2]
            # 计算小计（包含所有明细，含子产品）
            subtotal = quotation.amount or sum(d.total_price or 0 for d in quotation.details)
            vat_rate = 0  # 税率，默认0%
            vat_amount = subtotal * vat_rate
            total_amount = subtotal + vat_amount

            replacements2 = {
                '¥140,436': f'{currency_symbol}{subtotal:,.2f}',
                '¥0': f'{currency_symbol}{vat_amount:,.2f}',
                '¥140,436.00': f'{currency_symbol}{total_amount:,.2f}',
            }
            self._replace_text_in_table(table2, replacements2)

            # 保存到内存
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            # 生成文件名
            project_name = quotation.project.project_name if quotation.project else "未知项目"
            safe_project_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_', '（', '）', '【', '】')).rstrip()
            filename = f"{quotation.quotation_number} & {safe_project_name}.docx"

            logger.info(f"成功生成报价单Word文档: {filename}")

            return {
                'content': output.getvalue(),
                'filename': filename
            }

        except Exception as e:
            logger.error(f"生成报价单Word文档失败: {str(e)}")
            raise

    def generate_quotation_pdf(self, quotation):
        """
        使用Word模板生成报价单PDF

        Args:
            quotation: Quotation 对象

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        # 先生成Word
        word_result = self.generate_quotation_word(quotation)

        # 转换为PDF
        pdf_result = self.convert_word_to_pdf(
            word_result['content'],
            word_result['filename']
        )

        return pdf_result

    def _is_ovs_database(self):
        """检测当前是否使用OVS数据库"""
        from flask import current_app
        db_url = current_app.config.get('DATABASE_URL', '')
        db_type = os.environ.get('PMA_DB_TYPE', '') or os.environ.get('SUPABASE_DB_TYPE', '')
        return 'pqzviljbpfoqvyfulakl' in db_url or db_type == 'ovs'

    def generate_quotation_excel(self, quotation, template_type=None):
        """
        使用Excel模板生成报价单Excel文档

        Args:
            quotation: Quotation 对象
            template_type: 模板类型，可选值：
                - 'ovs': 强制使用OVS Singapore模板
                - 'sp8d': 强制使用SP8D默认模板
                - None: 根据数据库类型自动选择

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        # 强制指定模板类型
        if template_type == 'ovs':
            return self._generate_quotation_excel_ovs(quotation)
        elif template_type == 'sp8d':
            return self._generate_quotation_excel_default(quotation)

        # 自动检测：OVS数据库使用专用模板
        if self._is_ovs_database():
            return self._generate_quotation_excel_ovs(quotation)

        # 原有SP8D/默认模板逻辑
        return self._generate_quotation_excel_default(quotation)

    def _generate_quotation_excel_default(self, quotation):
        """
        使用默认Excel模板生成报价单（SP8D/本地）

        模板结构 (EVERTAC_Quotation_Template.xlsx):
        - A7: 客户名称
        - B9: 联系人
        - G6: 报价编号
        - G7: 日期
        - G8: 有效期
        - G9: 报价人
        - G10: 货币
        - 明细行从14行开始，列：A(序号) B(产品名称) C(型号) D(规格) E(数量) F(单价) G(金额) H(产品编码MN)
        - 汇总行：G列显示标签，H列显示数值
        """
        try:
            template_path = self._get_template_path('quotation_template.xlsx')
            wb = openpyxl.load_workbook(template_path)
            ws = wb.active

            # 准备基本数据
            currency = quotation.currency or Config.DEFAULT_CURRENCY
            currency_symbol = self._get_currency_symbol(currency)

            # 货币显示名称
            currency_names = {
                'CNY': 'CNY (人民币)',
                'USD': 'USD (美元)',
                'EUR': 'EUR (欧元)',
            }
            currency_display = currency_names.get(currency, currency)

            # 替换基本信息
            customer_name = quotation.customer.company_name if quotation.customer else ''
            contact_name = quotation.contact.name if quotation.contact else ''
            quotation_date = quotation.created_at.strftime('%Y-%m-%d') if quotation.created_at else ''
            valid_until = (quotation.created_at + timedelta(days=30)).strftime('%Y-%m-%d') if quotation.created_at else ''
            prepared_by = quotation.owner.real_name if quotation.owner else ''

            # 替换单元格值（根据新模板位置）
            ws['A7'] = customer_name  # 客户名称
            ws['B9'] = contact_name  # 联系人
            ws['G6'] = quotation.quotation_number  # 报价编号
            ws['G7'] = quotation_date  # 日期
            ws['G8'] = valid_until  # 有效期
            ws['G9'] = prepared_by  # 报价人
            ws['G10'] = currency_display  # 货币

            # 明细起始行
            detail_start_row = 14
            template_detail_count = 5  # 模板中有5行示例数据（14-18行）

            # 收集所有明细（包括子产品）
            all_details = []
            for detail in quotation.details:
                if detail.is_accessory and detail.parent_item_id:
                    continue
                all_details.append(detail)
                if detail.configurations:
                    for child in detail.configurations:
                        all_details.append(child)

            # 插入行前处理合并单元格（openpyxl insert_rows 不能正确移动合并区域）
            from openpyxl.styles import Alignment, Font
            from copy import copy

            insert_point = detail_start_row + template_detail_count  # 行19
            extra_rows = max(0, len(all_details) - template_detail_count)

            if extra_rows > 0:
                # 收集并解除插入点以下的所有合并单元格
                merges_to_shift = []
                for merge in list(ws.merged_cells.ranges):
                    if merge.min_row >= insert_point:
                        cell = ws.cell(row=merge.min_row, column=merge.min_col)
                        merges_to_shift.append({
                            'min_row': merge.min_row,
                            'max_row': merge.max_row,
                            'min_col': merge.min_col,
                            'max_col': merge.max_col,
                            'value': cell.value,
                            'font': copy(cell.font),
                            'alignment': copy(cell.alignment),
                        })
                        ws.unmerge_cells(str(merge))

                # 插入新行
                ws.insert_rows(insert_point, extra_rows)

                # 在偏移后的位置重新合并并恢复内容
                for m in merges_to_shift:
                    new_min = m['min_row'] + extra_rows
                    new_max = m['max_row'] + extra_rows
                    ws.merge_cells(start_row=new_min, start_column=m['min_col'],
                                   end_row=new_max, end_column=m['max_col'])
                    cell = ws.cell(row=new_min, column=m['min_col'])
                    cell.value = m['value']
                    cell.font = m['font']
                    cell.alignment = m['alignment']

            # 定义样式
            align_top = Alignment(vertical='top')
            align_top_wrap = Alignment(wrap_text=True, vertical='top')
            align_top_center = Alignment(horizontal='center', vertical='top')
            align_top_right = Alignment(horizontal='right', vertical='top')
            currency_fmt = '\\¥#,##0'

            # 填充明细数据
            for idx, detail in enumerate(all_details):
                row = detail_start_row + idx

                # 序号 - 居中顶部对齐
                ws[f'A{row}'] = idx + 1
                ws[f'A{row}'].alignment = align_top_center

                # 产品名称 - 顶部对齐
                ws[f'B{row}'] = detail.product_name or ''
                ws[f'B{row}'].alignment = align_top

                # 型号 - 顶部对齐
                ws[f'C{row}'] = detail.product_model or ''
                ws[f'C{row}'].alignment = align_top

                # 规格参数 - 不截断，保留完整内容，自动换行
                desc = detail.product_desc or ''
                ws[f'D{row}'] = desc
                ws[f'D{row}'].alignment = align_top_wrap

                # 根据内容计算行高（每行约15点，最小25点）
                line_count = desc.count('\n') + 1 if desc else 1
                max_line_length = max(len(line) for line in desc.split('\n')) if desc else 0
                wrapped_lines = max(line_count, (max_line_length // 40) + 1)
                row_height = max(25, wrapped_lines * 15)
                ws.row_dimensions[row].height = row_height

                # 数量 - 居中顶部对齐
                ws[f'E{row}'] = detail.quantity or 0
                ws[f'E{row}'].alignment = align_top_center

                # 单价 - 右对齐顶部 + 货币格式
                ws[f'F{row}'] = detail.unit_price or 0
                ws[f'F{row}'].alignment = align_top_right
                ws[f'F{row}'].number_format = currency_fmt

                # 金额公式 - 右对齐顶部 + 货币格式
                ws[f'G{row}'] = f'=E{row}*F{row}'
                ws[f'G{row}'].alignment = align_top_right
                ws[f'G{row}'].number_format = currency_fmt

                # 产品编码MN - 顶部对齐
                mn_code = detail.configured_mn or detail.product_mn or ''
                ws[f'H{row}'] = mn_code
                ws[f'H{row}'].alignment = align_top

            # 清除多余的模板行
            if len(all_details) < template_detail_count:
                for idx in range(len(all_details), template_detail_count):
                    row = detail_start_row + idx
                    for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
                        ws[f'{col}{row}'] = ''

            # 计算汇总行位置（明细结束后空一行）
            detail_end_row = detail_start_row + max(len(all_details), template_detail_count) - 1
            summary_row = detail_end_row + 2  # 空一行后的汇总行

            # 清除汇总区域可能残留的模板内容
            for r in range(summary_row, summary_row + 3):
                for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
                    ws[f'{col}{r}'] = None

            # 设置汇总行
            summary_fmt = '\\¥#,##0.00'
            ws[f'G{summary_row}'] = '小计 Subtotal'
            ws[f'G{summary_row}'].alignment = align_top_right
            ws[f'G{summary_row}'].font = Font(bold=True)
            ws[f'H{summary_row}'] = f'=SUM(G{detail_start_row}:G{detail_end_row})'
            ws[f'H{summary_row}'].number_format = summary_fmt
            ws[f'H{summary_row}'].font = Font(bold=True)

            ws[f'G{summary_row + 1}'] = '税费 VAT (0%)'
            ws[f'G{summary_row + 1}'].alignment = align_top_right
            ws[f'H{summary_row + 1}'] = 0
            ws[f'H{summary_row + 1}'].number_format = summary_fmt

            ws[f'G{summary_row + 2}'] = '总计 TOTAL'
            ws[f'G{summary_row + 2}'].alignment = align_top_right
            ws[f'G{summary_row + 2}'].font = Font(bold=True)
            ws[f'H{summary_row + 2}'] = f'=H{summary_row}+H{summary_row + 1}'
            ws[f'H{summary_row + 2}'].number_format = summary_fmt
            ws[f'H{summary_row + 2}'].font = Font(bold=True)

            # 设置页面打印选项 - 适应一页宽度，纵向布局
            from openpyxl.worksheet.properties import PageSetupProperties

            ws.page_setup.orientation = 'portrait'  # 纵向打印
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0  # 0表示不限制高度页数

            # 设置适应页面属性
            if ws.sheet_properties.pageSetUpPr is None:
                ws.sheet_properties.pageSetUpPr = PageSetupProperties()
            ws.sheet_properties.pageSetUpPr.fitToPage = True

            ws.print_options.horizontalCentered = True

            # 设置打印区域
            last_row = summary_row + 3
            ws.print_area = f'A1:H{last_row}'

            # 保存到内存
            output = BytesIO()
            wb.save(output)
            output.seek(0)

            # 生成文件名
            project_name = quotation.project.project_name if quotation.project else "未知项目"
            safe_project_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_', '（', '）', '【', '】')).rstrip()
            filename = f"{quotation.quotation_number} & {safe_project_name}.xlsx"

            logger.info(f"成功生成报价单Excel文档: {filename}")

            return {
                'content': output.getvalue(),
                'filename': filename
            }

        except Exception as e:
            logger.error(f"生成报价单Excel文档失败: {str(e)}")
            raise

    def _generate_quotation_excel_ovs(self, quotation):
        """
        使用OVS专用Excel模板生成报价单（Singapore模板）

        模板结构 (quotation_template_ovs.xlsx):
        - D2-D4: 公司信息（已固定在模板中）
        - B7: Company Name 标签, 值填入后方单元格
        - G7: Quotation No.
        - G8: Quotation Date
        - B9: Company Address
        - G9: Payment Terms (固定: Net 30 Days)
        - G10: Shipping Terms (固定: FOB Singapore)
        - B11: Contact Person
        - G11: Validity (固定: 30 Days)
        - B12: Contact No.
        - G12: Ref No.
        - 明细从第16行开始: B(S/N) C(Item No.) D(Brand) E(Description) F(Quantity) G(Unit Price USD) H(Amount USD)
        - 汇总行固定: G38(Total before GST) G39(GST=0) G40(Total after GST)
        """
        try:
            template_path = self._get_template_path('quotation_template_ovs.xlsx')
            wb = openpyxl.load_workbook(template_path)
            ws = wb.active

            # 准备基本数据
            customer_name = quotation.customer.company_name if quotation.customer else ''
            customer_address = ''
            if quotation.customer:
                # 尝试获取客户地址
                addr_parts = []
                if hasattr(quotation.customer, 'street') and quotation.customer.street:
                    addr_parts.append(quotation.customer.street)
                if hasattr(quotation.customer, 'city') and quotation.customer.city:
                    addr_parts.append(quotation.customer.city)
                if hasattr(quotation.customer, 'country') and quotation.customer.country:
                    addr_parts.append(quotation.customer.country)
                customer_address = ', '.join(addr_parts) if addr_parts else ''

            contact_name = quotation.contact.name if quotation.contact else ''
            contact_phone = quotation.contact.phone if quotation.contact else ''

            # 日期格式 mm/dd/yyyy
            quotation_date = quotation.created_at.strftime('%m/%d/%Y') if quotation.created_at else ''
            valid_until = (quotation.created_at + timedelta(days=30)).strftime('%m/%d/%Y') if quotation.created_at else ''

            # 填充客户信息
            # 根据模板合并单元格分析：
            # - B7:C8 是 "Company Name:" 标签，值填入 D7 (D7:F8 合并)
            # - B9:C10 是 "Company Address:" 标签，值填入 D9 (D9:F10 合并)
            # - B11 是 "Contact Person:" 标签，值填入 D11 (D11:F11 合并)
            # - B12 是 "Contact No.:" 标签，值填入 D12 (D12:F12 合并)
            # - G7-G12 是右侧标签，值填入 H7-H12
            ws['D7'] = customer_name       # Company Name: 后
            ws['D9'] = customer_address    # Company Address: 后
            ws['D11'] = contact_name       # Contact Person: 后
            ws['D12'] = contact_phone      # Contact No.: 后

            # 报价信息区域
            ws['H7'] = quotation.quotation_number  # Quotation No.: 后
            ws['H8'] = quotation_date              # Quotation Date: 后
            ws['H9'] = 'Net 30 Days'               # Payment Terms: 固定值
            ws['H10'] = 'FOB Singapore'            # Shipping Terms: 固定值
            ws['H11'] = '30 Days'                  # Validity: 固定值
            ws['H12'] = quotation.quotation_number # Ref No.: 使用报价编号

            # 明细配置
            detail_start_row = 16
            template_detail_count = 10  # OVS模板有10行预置行（16-25）
            summary_row_before_gst = 38  # 固定汇总行位置

            # 收集所有明细（包括子产品）
            all_details = []
            for detail in quotation.details:
                if detail.is_accessory and detail.parent_item_id:
                    continue
                all_details.append(detail)
                if detail.configurations:
                    for child in detail.configurations:
                        all_details.append(child)

            # 如果明细数量超过模板预置行数，需要插入新行
            if len(all_details) > template_detail_count:
                extra_rows = len(all_details) - template_detail_count
                # 在最后一个模板行之后插入
                ws.insert_rows(detail_start_row + template_detail_count, extra_rows)
                # 更新汇总行位置
                summary_row_before_gst += extra_rows

            # 填充明细数据
            from openpyxl.styles import Alignment

            # 定义对齐样式
            align_top = Alignment(vertical='top')
            align_top_wrap = Alignment(wrap_text=True, vertical='top')
            align_top_center = Alignment(horizontal='center', vertical='top')
            align_top_right = Alignment(horizontal='right', vertical='top')

            for idx, detail in enumerate(all_details):
                row = detail_start_row + idx

                # S/N (序号) - 居中顶部对齐
                ws[f'B{row}'] = idx + 1
                ws[f'B{row}'].alignment = align_top_center

                # Item No. - 产品编码MN - 顶部对齐
                mn_code = detail.configured_mn or detail.product_mn or ''
                ws[f'C{row}'] = mn_code
                ws[f'C{row}'].alignment = align_top

                # Brand - 品牌 - 顶部对齐
                brand = getattr(detail, 'brand', '') or ''
                ws[f'D{row}'] = brand
                ws[f'D{row}'].alignment = align_top

                # Description - 产品描述（产品名称 + 型号 + 规格）
                desc_parts = []
                if detail.product_name:
                    desc_parts.append(detail.product_name)
                if detail.product_model:
                    desc_parts.append(f"Model: {detail.product_model}")
                if detail.product_desc:
                    # 不截断规格描述，保留完整内容
                    desc_parts.append(detail.product_desc)

                description = '\n'.join(desc_parts) if desc_parts else ''
                ws[f'E{row}'] = description
                ws[f'E{row}'].alignment = align_top_wrap

                # 根据内容计算行高（每行约15点，最小30点）
                line_count = description.count('\n') + 1
                # 估算每行字符数，E列宽度约50字符
                max_line_length = max(len(line) for line in description.split('\n')) if description else 0
                wrapped_lines = max(line_count, (max_line_length // 50) + 1)
                row_height = max(30, wrapped_lines * 15)
                ws.row_dimensions[row].height = row_height

                # Quantity - 居中顶部对齐
                ws[f'F{row}'] = detail.quantity or 0
                ws[f'F{row}'].alignment = align_top_center

                # Unit Price (USD) - 右对齐顶部
                ws[f'G{row}'] = detail.unit_price or 0
                ws[f'G{row}'].alignment = align_top_right

                # Amount (USD) - 公式 - 右对齐顶部
                ws[f'H{row}'] = f'=F{row}*G{row}'
                ws[f'H{row}'].alignment = align_top_right

            # 清除多余的模板预置行
            if len(all_details) < template_detail_count:
                for idx in range(len(all_details), template_detail_count):
                    row = detail_start_row + idx
                    for col in ['B', 'C', 'D', 'E', 'F', 'G', 'H']:
                        ws[f'{col}{row}'] = ''

            # 计算明细结束行
            detail_end_row = detail_start_row + max(len(all_details), 1) - 1

            # 更新汇总区域的SUM公式范围
            # H38 = SUM(H16:H{detail_end_row})
            ws[f'H{summary_row_before_gst}'] = f'=SUM(H{detail_start_row}:H{detail_end_row})'
            # H39 = GST (固定为0)
            ws[f'H{summary_row_before_gst + 1}'] = 0
            # H40 = Total after GST
            ws[f'H{summary_row_before_gst + 2}'] = f'=H{summary_row_before_gst}+H{summary_row_before_gst + 1}'

            # 设置页面打印选项 - 适应一页宽度，纵向布局
            from openpyxl.worksheet.properties import PageSetupProperties

            ws.page_setup.orientation = 'portrait'  # 纵向打印
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0  # 0表示不限制高度页数

            # 设置适应页面属性
            if ws.sheet_properties.pageSetUpPr is None:
                ws.sheet_properties.pageSetUpPr = PageSetupProperties()
            ws.sheet_properties.pageSetUpPr.fitToPage = True

            ws.print_options.horizontalCentered = True

            # 设置打印区域（OVS模板到汇总行+签名区）
            last_row = summary_row_before_gst + 12  # 留出签名区域
            ws.print_area = f'A1:H{last_row}'

            # 保存到内存
            output = BytesIO()
            wb.save(output)
            output.seek(0)

            # 生成文件名
            project_name = quotation.project.project_name if quotation.project else "Unknown Project"
            safe_project_name = "".join(c for c in project_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"{quotation.quotation_number} - {safe_project_name}.xlsx"

            logger.info(f"成功生成OVS报价单Excel文档: {filename}")

            return {
                'content': output.getvalue(),
                'filename': filename
            }

        except Exception as e:
            import traceback
            logger.error(f"生成OVS报价单Excel文档失败: {str(e)}")
            logger.error(f"异常详情:\n{traceback.format_exc()}")
            raise

    def generate_quotation_excel_pdf(self, quotation, template_type=None):
        """
        生成报价单PDF（基于Excel模板）

        先生成Excel，然后使用LibreOffice转换为PDF

        Args:
            quotation: Quotation 对象
            template_type: 模板类型 ('ovs', 'sp8d', None)

        Returns:
            dict: {'content': bytes, 'filename': str}
        """
        import subprocess
        import tempfile
        import shutil

        try:
            # 1. 先生成Excel
            excel_result = self.generate_quotation_excel(quotation, template_type=template_type)
            excel_content = excel_result['content']
            excel_filename = excel_result['filename']

            # 优先走 Mac mini office-convert 服务
            service_url = os.environ.get('PMA_OFFICE_CONVERT_URL', '').rstrip('/')
            if service_url:
                import requests as _http
                try:
                    resp = _http.post(
                        f"{service_url}/convert",
                        files={'file': (excel_filename, excel_content,
                                        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')},
                        data={'target': 'pdf'},
                        timeout=120,
                    )
                except _http.exceptions.RequestException as e:
                    logger.error(f"office-convert 服务不可达: {e}")
                    raise RuntimeError(f"远程转换服务不可达: {e}")
                if resp.status_code != 200:
                    logger.error(f"office-convert 服务返回 {resp.status_code}: {resp.text[:200]}")
                    raise RuntimeError(f"远程转换失败 HTTP {resp.status_code}")
                pdf_filename = excel_filename.replace('.xlsx', '.pdf')
                return {'content': resp.content, 'filename': pdf_filename}

            # 2. 创建临时目录（本地 fallback）
            temp_dir = tempfile.mkdtemp()
            try:
                # 保存Excel到临时文件
                excel_path = os.path.join(temp_dir, excel_filename)
                with open(excel_path, 'wb') as f:
                    f.write(excel_content)

                # 3. 使用LibreOffice转换为PDF
                # 查找LibreOffice
                soffice_paths = [
                    '/opt/homebrew/bin/soffice',  # macOS Homebrew
                    '/usr/bin/soffice',           # Linux
                    '/usr/bin/libreoffice',       # Linux alternative
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS App
                ]

                soffice_cmd = None
                for path in soffice_paths:
                    if os.path.exists(path):
                        soffice_cmd = path
                        break

                if not soffice_cmd:
                    raise RuntimeError("LibreOffice未安装，无法转换PDF (可配置 PMA_OFFICE_CONVERT_URL 改用远程服务)")

                # 执行转换
                cmd = [
                    soffice_cmd,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    excel_path
                ]

                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=60
                )

                if result.returncode != 0:
                    logger.error(f"LibreOffice转换失败: {result.stderr}")
                    raise RuntimeError(f"PDF转换失败: {result.stderr}")

                # 4. 读取生成的PDF
                pdf_filename = excel_filename.replace('.xlsx', '.pdf')
                pdf_path = os.path.join(temp_dir, pdf_filename)

                if not os.path.exists(pdf_path):
                    raise RuntimeError(f"PDF文件未生成: {pdf_path}")

                with open(pdf_path, 'rb') as f:
                    pdf_content = f.read()

                return {
                    'content': pdf_content,
                    'filename': pdf_filename
                }

            finally:
                # 清理临时目录
                shutil.rmtree(temp_dir, ignore_errors=True)

        except Exception as e:
            import traceback
            logger.error(f"Excel转PDF失败: {str(e)}")
            logger.error(f"异常详情:\n{traceback.format_exc()}")
            raise


# 创建全局实例
word_generator = WordGenerator()


# 便捷函数
def generate_pricing_order_word(pricing_order, include_notes=False):
    """生成批价单Word文档的便捷函数"""
    return word_generator.generate_pricing_order_word_v2(pricing_order, include_notes=include_notes)


def generate_settlement_order_word(pricing_order):
    """生成结算单Word文档的便捷函数"""
    return word_generator.generate_settlement_order_word_v2(pricing_order)


def generate_pricing_order_pdf_from_word(pricing_order):
    """使用Word模板生成批价单PDF的便捷函数"""
    return word_generator.generate_pricing_order_pdf(pricing_order)


def generate_settlement_order_pdf_from_word(pricing_order):
    """使用Word模板生成结算单PDF的便捷函数"""
    return word_generator.generate_settlement_order_pdf(pricing_order)


def generate_quotation_word(quotation):
    """生成报价单Word文档的便捷函数"""
    return word_generator.generate_quotation_word(quotation)


def generate_quotation_pdf_from_word(quotation):
    """使用Word模板生成报价单PDF的便捷函数"""
    return word_generator.generate_quotation_pdf(quotation)
